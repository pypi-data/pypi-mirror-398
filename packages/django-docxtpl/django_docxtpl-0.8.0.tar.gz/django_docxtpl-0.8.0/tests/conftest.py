"""Pytest configuration and fixtures."""

from io import BytesIO

import pytest
from docx import Document
from docxtpl import DocxTemplate


@pytest.fixture
def sample_context():
    """Sample context for template rendering."""
    return {"name": "World", "title": "Test Document"}


@pytest.fixture
def simple_docx_template(tmp_path):
    """Create a simple DOCX template for testing.

    The template contains Jinja2 variables: {{ name }} and {{ title }}
    """
    # Create a simple Word document with Jinja2 placeholders
    doc = Document()
    doc.add_heading("{{ title }}", level=1)
    doc.add_paragraph("Hello, {{ name }}!")
    doc.add_paragraph("This is a test document.")

    template_path = tmp_path / "test_template.docx"
    doc.save(template_path)

    return template_path


@pytest.fixture
def simple_docx_bytes(simple_docx_template):
    """Return the template as bytes."""
    return simple_docx_template.read_bytes()


@pytest.fixture
def rendered_docx_bytes(simple_docx_template, sample_context):
    """Return a rendered DOCX as bytes."""
    doc = DocxTemplate(simple_docx_template)
    doc.render(sample_context)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@pytest.fixture
def template_dir(tmp_path, simple_docx_template):
    """Create a template directory with a test template."""
    import shutil

    template_dir = tmp_path / "templates"
    template_dir.mkdir()

    dest = template_dir / "test_template.docx"
    shutil.copy(simple_docx_template, dest)

    return template_dir
