"""Tests for django_docxtpl.response module."""

from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from django.test import RequestFactory

from django_docxtpl.response import DocxTemplateResponse


@pytest.fixture
def request_factory():
    """Django request factory."""
    return RequestFactory()


@pytest.fixture
def get_request(request_factory):
    """A simple GET request."""
    return request_factory.get("/")


class TestDocxTemplateResponse:
    """Tests for DocxTemplateResponse class."""

    def test_basic_docx_response(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test basic DOCX response generation."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="test",
        )

        assert response.status_code == 200
        assert response["Content-Type"] == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert 'filename="test.docx"' in response["Content-Disposition"]

    def test_filename_gets_extension(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that filename gets correct extension automatically."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="report",
            output_format="docx",
        )

        assert 'filename="report.docx"' in response["Content-Disposition"]

    def test_content_disposition_attachment(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test Content-Disposition is attachment by default."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="test",
        )

        assert response["Content-Disposition"].startswith("attachment;")

    def test_content_disposition_inline(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test Content-Disposition can be inline."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="test",
            as_attachment=False,
        )

        assert response["Content-Disposition"].startswith("inline;")

    def test_template_not_found(self, get_request, sample_context):
        """Test that missing template raises FileNotFoundError."""
        with pytest.raises(FileNotFoundError, match="Template not found"):
            DocxTemplateResponse(
                request=get_request,
                template="/nonexistent/template.docx",
                context=sample_context,
                filename="test",
            )

    def test_template_from_template_dir(
        self, get_request, sample_context, template_dir
    ):
        """Test loading template from DOCXTPL_TEMPLATE_DIR."""
        with patch("django_docxtpl.response.get_template_dir") as mock_dir:
            mock_dir.return_value = template_dir

            response = DocxTemplateResponse(
                request=get_request,
                template="test_template.docx",
                context=sample_context,
                filename="test",
            )

            assert response.status_code == 200

    def test_empty_context(self, get_request, simple_docx_template):
        """Test response with empty context."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=None,
            filename="test",
        )

        assert response.status_code == 200

    @patch("django_docxtpl.response.convert_docx")
    def test_pdf_conversion(
        self, mock_convert, get_request, simple_docx_template, sample_context
    ):
        """Test PDF conversion is called for PDF output format."""
        mock_convert.return_value = b"%PDF-1.4 fake pdf"

        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="report",
            output_format="pdf",
        )

        assert response.status_code == 200
        assert response["Content-Type"] == "application/pdf"
        assert 'filename="report.pdf"' in response["Content-Disposition"]
        mock_convert.assert_called_once()

    @patch("django_docxtpl.response.convert_docx")
    def test_odt_conversion(
        self, mock_convert, get_request, simple_docx_template, sample_context
    ):
        """Test ODT conversion."""
        mock_convert.return_value = b"fake odt content"

        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="document",
            output_format="odt",
        )

        assert response["Content-Type"] == "application/vnd.oasis.opendocument.text"
        assert 'filename="document.odt"' in response["Content-Disposition"]

    def test_content_is_valid_docx(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that response content is valid DOCX."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="test",
        )

        # DOCX files start with PK (zip signature)
        content = response.content
        assert content[:2] == b"PK"

    def test_template_variables_rendered(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that template variables are rendered."""
        from docx import Document

        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="test",
        )

        # Parse the response content
        doc = Document(BytesIO(response.content))

        # Get all text from the document
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Check that variables were replaced
        assert "World" in full_text  # {{ name }} was replaced
        assert "Test Document" in full_text  # {{ title }} was replaced
        assert "{{" not in full_text  # No unreplaced variables


class TestResolveTemplatePath:
    """Tests for template path resolution."""

    def test_absolute_path(self, get_request, simple_docx_template, sample_context):
        """Test absolute path resolution."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,  # This is already absolute
            context=sample_context,
            filename="test",
        )

        assert response.status_code == 200

    def test_relative_path_with_template_dir(
        self, get_request, sample_context, template_dir
    ):
        """Test relative path is resolved from template directory."""
        with patch("django_docxtpl.response.get_template_dir") as mock_dir:
            mock_dir.return_value = template_dir

            response = DocxTemplateResponse(
                request=get_request,
                template="test_template.docx",
                context=sample_context,
                filename="test",
            )

            assert response.status_code == 200

    def test_path_object_accepted(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that Path objects are accepted."""
        response = DocxTemplateResponse(
            request=get_request,
            template=Path(simple_docx_template),
            context=sample_context,
            filename="test",
        )

        assert response.status_code == 200


class TestUpdateFields:
    """Tests for update_fields parameter."""

    @patch("django_docxtpl.response.convert_docx")
    def test_update_fields_passed_to_convert(
        self, mock_convert, get_request, simple_docx_template, sample_context
    ):
        """Test that update_fields is passed to convert_docx."""
        mock_convert.return_value = b"%PDF-1.4 fake pdf"

        DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="report",
            output_format="pdf",
            update_fields=True,
        )

        # Check that convert_docx was called with update_fields=True
        mock_convert.assert_called_once()
        _, kwargs = mock_convert.call_args
        assert kwargs.get("update_fields") is True

    @patch("django_docxtpl.response.convert_docx")
    def test_update_fields_false_by_default(
        self, mock_convert, get_request, simple_docx_template, sample_context
    ):
        """Test that update_fields is False by default."""
        mock_convert.return_value = b"%PDF-1.4 fake pdf"

        DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="report",
            output_format="pdf",
        )

        mock_convert.assert_called_once()
        _, kwargs = mock_convert.call_args
        assert kwargs.get("update_fields") is False

    @patch("django_docxtpl.response.update_fields_in_docx")
    def test_update_fields_with_docx_output(
        self, mock_update, get_request, simple_docx_template, sample_context
    ):
        """Test that update_fields works with DOCX output format."""
        mock_update.return_value = b"updated docx content"

        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="report",
            output_format="docx",
            update_fields=True,
        )

        # update_fields_in_docx should be called for DOCX with update_fields=True
        mock_update.assert_called_once()
        assert response.content == b"updated docx content"

    def test_update_fields_not_called_for_docx_without_flag(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that update_fields_in_docx is not called for DOCX without flag."""
        with patch("django_docxtpl.response.update_fields_in_docx") as mock_update:
            DocxTemplateResponse(
                request=get_request,
                template=simple_docx_template,
                context=sample_context,
                filename="report",
                output_format="docx",
                update_fields=False,
            )

            mock_update.assert_not_called()


class TestCallableContext:
    """Tests for callable context support in DocxTemplateResponse."""

    def test_context_as_callable(self, get_request, simple_docx_template):
        """Test that context can be callable receiving DocxTemplate and tmp_dir."""
        received_docx = None
        received_tmp_dir = None

        def context_builder(docx, tmp_dir):
            nonlocal received_docx, received_tmp_dir
            received_docx = docx
            received_tmp_dir = tmp_dir
            return {"name": "FromCallable", "title": "Callable Title"}

        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=context_builder,
            filename="test",
        )

        assert response.status_code == 200
        # Verify that the callable was called with a DocxTemplate instance
        from pathlib import Path

        from docxtpl import DocxTemplate

        assert received_docx is not None
        assert isinstance(received_docx, DocxTemplate)
        assert received_tmp_dir is not None
        assert isinstance(received_tmp_dir, Path)

    def test_context_callable_renders_correctly(
        self, get_request, simple_docx_template
    ):
        """Test that context from callable is used in rendering."""
        from io import BytesIO
        from zipfile import ZipFile

        def context_builder(docx, tmp_dir):
            return {"name": "CallableName", "title": "CallableTitle"}

        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=context_builder,
            filename="test",
        )

        # Extract document.xml to verify content was rendered
        docx_content = BytesIO(response.content)
        with ZipFile(docx_content) as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
            assert "CallableName" in doc_xml
            assert "CallableTitle" in doc_xml

    def test_context_dict_still_works(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that regular dict context still works (backward compatibility)."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="test",
        )

        assert response.status_code == 200

    def test_context_callable_with_pdf_conversion(
        self, get_request, simple_docx_template
    ):
        """Test callable context works with PDF conversion."""
        with patch("django_docxtpl.response.convert_docx") as mock_convert:
            mock_convert.return_value = b"%PDF fake"

            def context_builder(docx, tmp_dir):
                return {"name": "Test", "title": "Title"}

            response = DocxTemplateResponse(
                request=get_request,
                template=simple_docx_template,
                context=context_builder,
                filename="test",
                output_format="pdf",
            )

            assert response.status_code == 200
            mock_convert.assert_called_once()


class TestJinjaEnv:
    """Tests for custom Jinja2 Environment support."""

    def test_jinja_env_with_custom_filter(self, get_request, tmp_path):
        """Test DocxTemplateResponse with custom jinja_env filter."""
        from io import BytesIO
        from zipfile import ZipFile

        from docx import Document
        from jinja2 import Environment

        # Create a template with custom filter syntax
        doc = Document()
        doc.add_paragraph("Value: {{ value|milers }}")
        template_path = tmp_path / "filter_template.docx"
        doc.save(template_path)

        # Create custom jinja_env with filter
        def format_milers(value):
            return f"{value:,.0f}".replace(",", ".")

        jinja_env = Environment(autoescape=True)
        jinja_env.filters["milers"] = format_milers

        response = DocxTemplateResponse(
            request=get_request,
            template=template_path,
            context={"value": 1234567},
            filename="test",
            jinja_env=jinja_env,
        )

        assert response.status_code == 200

        # Verify the filter was applied
        docx_content = BytesIO(response.content)
        with ZipFile(docx_content) as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
            assert "1.234.567" in doc_xml

    def test_jinja_env_none_uses_default(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that jinja_env=None uses default Jinja2 Environment."""
        response = DocxTemplateResponse(
            request=get_request,
            template=simple_docx_template,
            context=sample_context,
            filename="test",
            jinja_env=None,
        )

        assert response.status_code == 200


class TestAutoescape:
    """Tests for autoescape parameter in DocxTemplateResponse."""

    def test_autoescape_default_is_false(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that autoescape is False by default."""
        with patch("django_docxtpl.response.DocxTemplate") as mock_docx:
            mock_instance = MagicMock()
            mock_docx.return_value = mock_instance
            mock_instance.save = MagicMock(side_effect=lambda buf: buf.write(b"PK"))

            DocxTemplateResponse(
                request=get_request,
                template=simple_docx_template,
                context=sample_context,
                filename="test",
            )

            mock_instance.render.assert_called_once()
            _, kwargs = mock_instance.render.call_args
            assert kwargs.get("autoescape") is False

    def test_autoescape_true_passed_to_render(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that autoescape=True is passed to doc.render()."""
        with patch("django_docxtpl.response.DocxTemplate") as mock_docx:
            mock_instance = MagicMock()
            mock_docx.return_value = mock_instance
            mock_instance.save = MagicMock(side_effect=lambda buf: buf.write(b"PK"))

            DocxTemplateResponse(
                request=get_request,
                template=simple_docx_template,
                context=sample_context,
                filename="test",
                autoescape=True,
            )

            mock_instance.render.assert_called_once()
            _, kwargs = mock_instance.render.call_args
            assert kwargs.get("autoescape") is True

    def test_autoescape_false_explicitly(
        self, get_request, simple_docx_template, sample_context
    ):
        """Test that autoescape=False is passed explicitly."""
        with patch("django_docxtpl.response.DocxTemplate") as mock_docx:
            mock_instance = MagicMock()
            mock_docx.return_value = mock_instance
            mock_instance.save = MagicMock(side_effect=lambda buf: buf.write(b"PK"))

            DocxTemplateResponse(
                request=get_request,
                template=simple_docx_template,
                context=sample_context,
                filename="test",
                autoescape=False,
            )

            mock_instance.render.assert_called_once()
            _, kwargs = mock_instance.render.call_args
            assert kwargs.get("autoescape") is False
