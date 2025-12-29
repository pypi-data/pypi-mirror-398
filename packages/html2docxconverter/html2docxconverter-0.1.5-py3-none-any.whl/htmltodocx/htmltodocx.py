from __future__ import annotations

from dataclasses import replace
from typing import TYPE_CHECKING
from bs4 import BeautifulSoup, Tag, NavigableString
from docx.shared import Cm

from htmltodocx.config import base_config
from htmltodocx.htmlimage import add_image_to_docx
from htmltodocx.htmllist import HTMLListHelper
from htmltodocx.htmltable import (
    apply_row_height,
    TableCellStyleApplier,
    TableStyleApplier,
)
from htmltodocx.parse_style import (
    css_length_to_inches, StyleSchema, parse_style_selector, apply_styles_to_text,
    ParserStyle,
)
from htmltodocx.schemas import ConfigSchema

if TYPE_CHECKING:
    from docx.table import _Cell
    from bs4 import PageElement
    from docx.text.paragraph import Paragraph
    from typing import Optional
    from docx.document import Document


class HTMLtoDocx:
    def __init__(
            self,
            document: Document,
            config: ConfigSchema = base_config,
            message_start: str | None = None,
            layout_table_class: str = None,
    ):
        self.doc = document
        self.config = config
        self.parser_style = ParserStyle(config)
        self.message_start = message_start
        self.first_paragraph = True
        self.layout_table_class = layout_table_class
        self.list_helper = HTMLListHelper(self.doc)
        self.list_num_counter = 1
        self.abstract_num_counter = 1

    def parse_html(self, html: str) -> Document:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup.contents:
            if tag:
                self.handle_tag(tag)
        return self.doc

    def handle_tag(
            self,
            tag: NavigableString | Tag | PageElement,
            paragraph: Paragraph | None = None,
            with_indent: bool = False,
            table_cell: _Cell | None = None,
    ) -> None:
        if tag.name == "p":
            self.handle_paragraph(tag, paragraph, with_indent)
        elif tag.name == "ul" or tag.name == "ol":
            self.handle_list_items(tag, paragraph, table_cell)
        elif tag.name == "table":
            self.handle_table(tag, table_cell)
        elif tag.name in ['figure', 'img']:
            self.handle_image(tag, paragraph)
        elif tag.text and paragraph:
            self.add_text_and_style(tag, paragraph)

    def handle_image(self, tag: Tag, paragraph: Paragraph) -> None:
        if not paragraph:
            paragraph = self.doc.add_paragraph()
        self.parser_style.apply_position_image(tag, paragraph)
        add_image_to_docx(tag, paragraph)

    def handle_paragraph(
            self, tag: list[Tag] | Tag, paragraph: Optional[Paragraph] = None, with_indent: bool = False
    ) -> None:
        if not paragraph:
            paragraph = self.doc.add_paragraph()

        if with_indent:
            paragraph.paragraph_format.first_line_indent = Cm(1.25)

        self.parser_style.apply_position_paragraph(tag, paragraph)
        if self.first_paragraph and self.message_start:
            first = paragraph.add_run(self.message_start)
            first.underline = True
            self.first_paragraph = False
        self.add_text_and_style(tag, paragraph)

    def handle_list_items(
            self, tag: Tag, paragraph: Paragraph, table_cell: _Cell | None = None
    ) -> None:
        is_ordered = tag.name == "ol"
        num_id = self.list_helper.add_numbering_definition(is_ordered)
        not_first = False

        if table_cell:
            for li in tag.find_all("li"):
                if not not_first:
                    not_first = True
                else:
                    paragraph = table_cell.add_paragraph()
                item = li.find() or li
                self.handle_tag(item, paragraph)
                self.list_helper.add_list_item(paragraph, is_ordered, num_id)
                if not not_first:
                    not_first = True
        else:
            if not paragraph:
                paragraph = self.doc.add_paragraph()
            for li in tag.find_all("li"):
                if not_first:
                    paragraph = self.doc.add_paragraph()
                self.handle_tag(li, paragraph)
                self.list_helper.add_list_item(paragraph, is_ordered, num_id)
                if not not_first:
                    not_first = True
        self.list_num_counter += 1

    def handle_table(self, tag: Tag, table_in_table: _Cell | None = None) -> None:
        tbody = tag.find("tbody", recursive=False)
        rows = tbody.find_all("tr", recursive=False)
        max_cols = 0
        for row_tag in rows:
            count = 0
            for col in row_tag.find_all(["td", "th"], recursive=False):
                count += int(col.get("colspan", 1))
            max_cols = max(max_cols, count)

        if table_in_table:
            table = table_in_table.add_table(rows=len(rows), cols=max_cols)
        else:
            table = self.doc.add_table(rows=len(rows), cols=max_cols)

        layout_table = False
        if self.layout_table_class and tag and tag.get('class'):
            layout_table = True if self.layout_table_class in tag.get("class") else False

        table_style = TableStyleApplier(tag, table).apply()
        cell_style = TableCellStyleApplier(
            layout_table=layout_table,
            table_style=table_style,
            parser_style=self.parser_style,
            config=self.config,
        )
        occupancy = [[False] * max_cols for _ in range(len(rows))]
        for row_idx, row_tag in enumerate(rows):
            col_tags = row_tag.find_all(["td", "th"], recursive=False)
            col_ptr = 0
            height = table_style.get("height")
            height_cell_array = []
            for cell_tag in col_tags:
                while col_ptr < max_cols and occupancy[row_idx][col_ptr]:
                    col_ptr += 1

                rowspan = int(cell_tag.get("rowspan", 1))
                colspan = int(cell_tag.get("colspan", 1))

                start_cell = table.cell(row_idx, col_ptr)
                for r in range(row_idx, row_idx + rowspan):
                    for c in range(col_ptr, col_ptr + colspan):
                        if r < len(rows) and c < max_cols:
                            occupancy[r][c] = True

                if colspan > 1:
                    end_cell = table.cell(row_idx, col_ptr + colspan - 1)
                    start_cell = start_cell.merge(end_cell)

                if rowspan > 1:
                    end_cell = table.cell(row_idx + rowspan - 1, col_ptr)
                    start_cell = start_cell.merge(end_cell)

                height_cell = cell_style.apply(
                    cell_tag=cell_tag,
                    cell=start_cell,
                )
                if height_cell:
                    height_cell_array.append(int(css_length_to_inches(height_cell) * 1440))
                paragraph = start_cell.paragraphs[0]
                cell_paragraph_created = False
                for item in cell_tag.contents:
                    if isinstance(item, NavigableString) and not item.strip():
                        continue
                    if isinstance(item, Tag):
                        if item.name == "p":
                            para = (
                                start_cell.add_paragraph()
                                if cell_paragraph_created
                                else start_cell.paragraphs[0]
                            )
                            self.handle_paragraph(item, para)
                        elif item.name in ["figure", "table", "ul", "ol", "img"]:
                            self.handle_tag(item, paragraph, table_cell=start_cell)
                        else:
                            para = (
                                start_cell.add_paragraph()
                                if cell_paragraph_created
                                else start_cell.paragraphs[0]
                            )
                            self.add_text_and_style(item, para)
                    elif isinstance(item, NavigableString):
                        para = (
                            start_cell.add_paragraph()
                            if cell_paragraph_created
                            else start_cell.paragraphs[0]
                        )
                        para.add_run(item.strip())

                    cell_paragraph_created = True

                col_ptr += colspan

            if height:
                apply_row_height(table.rows[row_idx], height)

            if height_cell_array and not height:
                apply_row_height(table.rows[row_idx], max(height_cell_array))

    def add_text_and_style(
            self,
            tags: list[Tag | NavigableString] | NavigableString | Tag,
            paragraph: Paragraph,
    ) -> None:
        for tag in tags:
            if tag.name in ["p", "ul", "ol", "table", "figure"]:
                self.handle_tag(tag, paragraph)
            else:
                self._parse_styled_text(tag, paragraph)

    def _parse_styled_text(self, tag: Tag | NavigableString | PageElement, paragraph: Paragraph,
                           styles: Optional[StyleSchema] = None):

        current_styles = replace(styles) if styles is not None else StyleSchema()
        if isinstance(tag, Tag):
            current_styles = parse_style_selector(tag, current_styles)
            for child in tag.children:
                self._parse_styled_text(child, paragraph, current_styles)

        elif isinstance(tag, NavigableString):
            text = str(tag)
            if text:
                run = paragraph.add_run(text)
                if styles:
                    apply_styles_to_text(run, styles)
