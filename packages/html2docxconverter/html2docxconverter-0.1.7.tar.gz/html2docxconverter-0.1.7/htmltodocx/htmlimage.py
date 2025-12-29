from __future__ import annotations

import base64
import re
from io import BytesIO
from typing import TYPE_CHECKING
from docx.shared import Inches

from htmltodocx.parse_style import css_length_to_inches, parse_style_css

if TYPE_CHECKING:
    from bs4.element import Tag
    from docx.text.paragraph import Paragraph


def add_image_to_docx(tag: Tag, paragraph: Paragraph) -> None:
    if tag.name == "img":
        img = tag
    else:
        img = tag.find("img")
    if not img:
        return
    image_data, mime_type = parse_data_uri(img.get("src"))
    style = parse_style_css(tag.get("style", ""))
    img_width_px = int(img.get("width", 0))
    img_height_px = int(img.get("height", 0))

    if img_width_px and img_height_px:
        aspect_ratio = img_width_px / img_height_px
    else:
        css_aspect_ratio = style.get("aspect-ratio", "").split("/")
        if len(css_aspect_ratio) == 2:
            aspect_ratio = float(css_aspect_ratio[0]) / float(css_aspect_ratio[1])
        else:
            aspect_ratio = None

    target_width = None
    if "width" in style:
        width_css = style["width"]
        if "%" in width_css:
            page_width_inches = 6.5
            width_percent = float(width_css.strip("%")) / 100
            target_width = page_width_inches * width_percent
        else:
            target_width = css_length_to_inches(width_css)

    image_stream = BytesIO(image_data)
    run = paragraph.add_run()
    if target_width and aspect_ratio:
        target_height = target_width / aspect_ratio
        run.add_picture(image_stream, width=Inches(target_width), height=Inches(target_height))
    elif target_width:
        run.add_picture(image_stream, width=Inches(target_width))
    else:
        run.add_picture(image_stream)



def parse_data_uri(data_uri: str) -> tuple[bytes, str]:
    if not data_uri.startswith("data:image/"):
        raise ValueError("Некорректный data URI для изображения")
    match = re.match(r"data:image/(\w+);base64,(.*)", data_uri)
    if not match:
        raise ValueError("Некорректный формат data URI")

    mime_type = match.group(1)
    base64_data = match.group(2)
    image_data = base64.b64decode(base64_data)

    return image_data, mime_type
