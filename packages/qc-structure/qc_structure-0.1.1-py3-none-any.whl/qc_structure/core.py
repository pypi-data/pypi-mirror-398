"""
exam-qc-core: コアモジュール

ExamDocumentクラスとIDマッピングロジックを提供します。
"""

import json
from typing import List, Dict, Any
from pathlib import Path
from docx import Document

from .models import EnhancedParagraph
from .utils import iter_all_paragraphs, get_para_id


class IDMapper:
    """
    構造情報とWord段落を紐付けるマッパー（内部用）。
    """
    
    def __init__(self, structure_json_path: str):
        """
        初期化。
        
        Args:
            structure_json_path: structure.jsonのパス
        """
        with open(structure_json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        self.structures = data.get('questions', [])
        
        # start_para_id -> structure_item のマップを作成
        self.anchor_map: Dict[str, Dict[str, Any]] = {}
        for item in self.structures:
            para_id = item.get('start_para_id')
            if para_id:
                self.anchor_map[para_id] = item
    
    def get_structure_by_anchor(self, para_id: str) -> Dict[str, Any]:
        """
        アンカーID（paraId）から構造情報を取得する。
        
        Args:
            para_id: Word内部段落ID
            
        Returns:
            構造情報（存在しない場合はNone）
        """
        return self.anchor_map.get(para_id)


class ExamDocument:
    """
    試験文書全体を表すクラス。
    """
    
    def __init__(self, docx_path: str, mapper: IDMapper):
        """
        初期化。
        
        Args:
            docx_path: Wordファイルパス
            mapper: IDMapperインスタンス
        """
        self.docx_path = docx_path
        self.document = Document(docx_path)
        self.mapper = mapper
        self.paragraphs: List[EnhancedParagraph] = []
        
        self._load_paragraphs()
    
    def _load_paragraphs(self):
        """
        全段落をロードし、構造情報を紐付ける（内部メソッド）。
        """
        for raw_para, in_table in iter_all_paragraphs(self.document):
            enhanced = EnhancedParagraph(raw_para)
            
            # Word内部IDを取得
            para_id = get_para_id(raw_para)
            
            if para_id:
                # アンカーベースでマッピング
                structure = self.mapper.get_structure_by_anchor(para_id)
                if structure:
                    enhanced.pid = structure.get('start_pid')
                    enhanced.role = structure.get('role')
                    enhanced.structure = structure
            
            self.paragraphs.append(enhanced)


def load_exam_document(docx_path: str, structure_json_path: str) -> ExamDocument:
    """
    試験文書をロードする（エントリーポイント）。
    
    Args:
        docx_path: Wordファイルパス
        structure_json_path: 構造情報JSONパス
        
    Returns:
        初期化済みのExamDocumentインスタンス
    """
    mapper = IDMapper(structure_json_path)
    return ExamDocument(docx_path, mapper)
