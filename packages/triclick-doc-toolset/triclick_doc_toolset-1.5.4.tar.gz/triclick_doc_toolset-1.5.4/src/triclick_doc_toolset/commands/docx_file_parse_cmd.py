from __future__ import annotations

from typing import List, Dict
from dataclasses import asdict
import time
import logging

# 为拼接多行标题引入 docx 读取
import re
from docx import Document

# 框架依赖
from triclick_doc_toolset.framework import Command, Context, CommandRegistry

# 解析与标题/标签工具
from triclick_doc_toolset.common.word import (
    extract_docx_content_with_metadata,
    assemble_title_lines,
)
from triclick_doc_toolset.common.word.title_util import normalize_docx_text, is_toc_like_paragraph, get_title_util_cache_stats, clear_title_util_caches
from triclick_doc_toolset.common.utils import load_footnote_patterns
from triclick_doc_toolset.common.rules import (
    apply_normalize_duplicate_labels_rule,
)



class DocxFileParseCommand(Command):
    """
    框架命令子类：按“标题-表格-脚注/段落”分块（仅元数据）。
    仅使用 `extract_docx_content_with_metadata` 进行解析。
    """

    def is_satisfied(self, context: Context) -> bool:
        return context.has_document()

    def execute(self, context: Context) -> Context:
        # 解析输入路径（支持文件或文件夹），仅处理 .docx
        paths = context.resolve_document_paths(patterns=["*.docx"])
        if not paths:
            context.add_error("No DOCX files resolved from context")
            return context

        # 使用通用工具进行标题组装与标签归一化

        # 将解析得到的 TableItem 列表转为结构化字典，并写入 Context.sections
        parsed_sections: List[Dict] = []
        for p in paths:
            _start = time.perf_counter()
            doc = Document(str(p))
            footnote_patterns = load_footnote_patterns()
            sections = extract_docx_content_with_metadata(str(p), doc=doc)
            for sec in sections:
                try:
                    if (getattr(sec, "type", "") or "").strip().lower() == "figure":
                        continue
                except Exception:
                    pass
                d = asdict(sec)
                first_indices = sorted(set(int(x) for x in (sec.title_indices or [])))
                if first_indices:
                    i0 = first_indices[0]
                    if 0 <= i0 < len(doc.paragraphs):
                        para0 = doc.paragraphs[i0]
                        style0 = (getattr(getattr(para0, "style", None), "name", "") or "")
                        if (
                            is_toc_like_paragraph(para0.text or "", style0)
                        ):
                            continue
                        # 额外扫描前若干行以规避误判目录
                        toc_hit = False
                        for j in first_indices[:3]:
                            if j < 0 or j >= len(doc.paragraphs):
                                continue
                            pj = doc.paragraphs[j]
                            styj = (getattr(getattr(pj, "style", None), "name", "") or "")
                            if is_toc_like_paragraph(pj.text or "", styj):
                                toc_hit = True
                                break
                        if toc_hit:
                            continue
                # 组装标题（按索引序拼接非空行）
                foot_lines: List[tuple[int, str]] = []
                for i in first_indices:
                    if i < 0 or i >= len(doc.paragraphs):
                        continue
                    t = normalize_docx_text(doc.paragraphs[i].text)
                    if not t:
                        continue
                    s = re.sub(r"[\u00A0\u2007\u202F]+", " ", t)
                    s = re.sub(r"\s+", " ", s)
                    if any(p.search(s) for p in footnote_patterns):
                        foot_lines.append((i, t))
                full_title = assemble_title_lines(str(p), first_indices)
                if full_title:
                    d["title"] = full_title
                    # 若存在规范化后的 label，则将标题前缀替换为该 label，确保一致性
                    label = (d.get("label") or "").strip()
                    if label and re.match(r"(?i)^(table|listing|figure)\s+", label):
                        try:
                            d["title"] = re.sub(
                                r"(?i)^(?:[A-Za-z]\.|\d+\.)?\s*(table|listing|figure)\s+\S+",
                                label,
                                d["title"]
                            )
                        except Exception:
                            # 容错：保持原标题
                            pass
                if foot_lines:
                    existing = (d.get("footnote") or "").strip()
                    added = "\n".join([t for _, t in foot_lines]).strip()
                    d["footnote"] = (added if not existing else (existing + ("\n" if not existing.endswith("\n") else "") + added))
                    fidx = [int(i) for i in (d.get("footnote_indices") or []) if isinstance(i, int)]
                    fidx.extend(i for i, _ in foot_lines)
                    d["footnote_indices"] = sorted(set(fidx))
                d["source_file"] = str(p)
                parsed_sections.append(d)
            _elapsed = time.perf_counter() - _start
            stats = get_title_util_cache_stats()
            context.processing_summary.setdefault("title_util_cache_stats", []).append({"file": str(p), **stats})
            clear_title_util_caches(str(p))
            logging.getLogger(__name__).info(f"[{self.name}] 文件处理 {str(p)} 耗时 {_elapsed:.2f}s")

        # 规则：重复标签归一化（通过 self.rules 控制）
        # 查找名为 "normalize_duplicate_labels" 的规则并应用
        rule = next((r for r in self.rules if (r.get("name") == "normalize_duplicate_labels" and r.get("enabled") is True)), None)
        if rule:
            params = rule.get("params") or {}
            suffix_fmt = str(params.get("suffix_format", ".{n}"))
            apply_normalize_duplicate_labels_rule(parsed_sections, suffix_fmt)

        # 更新上下文
        context.doc_type = "docx"
        context.sections = parsed_sections
        context.processing_summary["title_table_footnote_partition"] = {
            "files_processed": len(paths),
            "sections_extracted": len(parsed_sections),
            "mode": "metadata_only",
        }
        return context


# 注册到命令注册表，便于 Pipeline 通过 YAML 创建
CommandRegistry.register("DocxFileParseCommand", DocxFileParseCommand)
