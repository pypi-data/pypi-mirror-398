from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, cast, Dict
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

from triclick_doc_toolset.common import TableItem
from triclick_doc_toolset.common.word.title_util import (
    normalize_docx_text,
    extract_leading_enumerator,
    get_paragraph_font_info,
    is_structure_separator,
)
from triclick_doc_toolset.common.utils import (
    load_title_patterns,
    load_label_patterns,
    load_footnote_patterns,
)

TITLE_PATTERNS: List[re.Pattern] = load_title_patterns()
LABEL_PATTERNS: List[re.Pattern] = load_label_patterns()
FOOTNOTE_PATTERNS: List[re.Pattern] = load_footnote_patterns()


def extract_title_label(text: str) -> Optional[str]:
    """
    从标题文本中提取简洁标签（例如："Table 2.3" 或 "Listing 1"）。
    仅返回前缀 + 编号，不包含后续描述。
    支持从配置的多个正则中按序匹配，取首个命中。
    """
    text = (text or "").strip()
    for pat in LABEL_PATTERNS:
        # 允许在任意位置提取（用于脚注中的 "same as table 2.2.1" 等）
        m = pat.search(text)
        if m:
            label = m.group(0).strip()
            label = re.sub(r'\s+', ' ', label)
            # 统一大小写（前缀首字母大写，大小写不敏感）
            label = re.sub(r'(?i)^(table|listing|figure)', lambda m: m.group(1).capitalize(), label)
            return label
    return None


# 使用公共工具 extract_leading_enumerator 以统一行为


def _iter_block_items(doc: DocxDocument):
    """
    迭代文档正文的块级元素（段落与表格），保持与文档中出现的真实顺序。

    返回值:
    - 当遇到段落时，产出 ('paragraph', Paragraph)
    - 当遇到表格时，产出 ('table', Table)

    说明:
    - 直接遍历底层 body.iterchildren()，避免仅用 doc.paragraphs 或 doc.tables 导致顺序错乱。
    - 该生成器用于在标题、表格、脚注之间建立基于相邻关系的解析。
    """
    body = doc._element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield 'paragraph', Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield 'table', Table(child, doc)


def extract_docx_content_with_metadata(docx_path: str, doc: Optional[DocxDocument] = None) -> List[TableItem]:
    """
    基于文档结构模式解析文档：
    模式1: title + table + footnote
    模式2: title + footnote (无表格)
    
    使用结构化方法而非枚举式匹配：
    1. 识别标题段落
    2. 收集连续的同字体标题段落
    3. 遇到表格则标记为模式1，否则为模式2
    4. 表格后或标题后遇到字体变化/空行则开始收集脚注
    """
    doc = doc or Document(docx_path)
    content: List[TableItem] = []
    para_idx = -1
    tbl_idx = -1

    current_section: Optional[TableItem] = None
    title_font_info = None  # 标题的字体信息
    title_start_index = None  # 标题开始的段落索引
    section_state = "seeking_title"  # seeking_title, collecting_title, after_table, collecting_footnote

    def compute_title_level(label: Optional[str]) -> int:
        if not label:
            return 0
        parts = label.split(None, 1)
        number_part = parts[1] if len(parts) == 2 else parts[0]
        return number_part.count('.')

    for kind, item in _iter_block_items(doc):
        if kind == 'paragraph':
            par = cast(Paragraph, item)
            text = normalize_docx_text(par.text)

            # 无论是否为空，段落索引均递增，以保证与 Document.paragraphs 的索引对齐
            para_idx += 1
            p_idx = para_idx

            # 空段落作为结构分隔符参与状态转换，但不记录到标题/脚注文本
            if not text:
                if section_state == "collecting_title" and current_section:
                    section_state = "collecting_footnote"
                continue

            # 检查是否为新标题：支持行首枚举（数字./字母.）后接 Table/Listing/Figure
            is_title = any(pat.match(text) for pat in TITLE_PATTERNS)
            if not is_title:
                enum = extract_leading_enumerator(text)
                if enum:
                    stripped = re.sub(r"^\s*([0-9]+|[A-Za-z])\.?\s*", "", text)
                    is_title = any(pat.match(stripped) for pat in TITLE_PATTERNS)
            
            if is_title:
                # 保存前一个section
                if current_section:
                    content.append(current_section)
                
                # 创建新section
                label = extract_title_label(text)
                # 若行首存在“数字.”或“字母.”，并且标签为 Table，则将枚举并入 label：Table 4 -> Table 4.X
                enum = extract_leading_enumerator(text)
                if enum and label and re.match(r"(?i)^table\s+", label):
                    label = f"{label}.{enum}"
                # 重复标签归一化不在解析阶段完成，统一由命令层工具处理
                section = None
                if label:
                    parts = label.split(None, 1)
                    section = (parts[1] if len(parts) == 2 else parts[0]).strip()
                
                current_section = TableItem(
                    label=label,
                    section=section,
                    title=text,
                    level=compute_title_level(label),
                    title_indices=[p_idx]
                )
                
                # 记录标题字体信息和开始位置
                title_font_info = get_paragraph_font_info(par)
                title_start_index = p_idx
                section_state = "collecting_title"
                continue

            # 处理非标题段落
            if current_section:
                if section_state == "collecting_title":
                    # 检查是否应该继续收集标题还是转为脚注
                    if title_font_info and not is_structure_separator(par, title_font_info, p_idx, title_start_index):
                        # 继续收集标题
                        current_section.title_indices.append(p_idx)
                        # 解析阶段不再拼接标题文本，统一在命令层按索引组装
                    else:
                        # 转为脚注收集
                        section_state = "collecting_footnote"
                        current_section.add_footnote_index(p_idx)
                        current_section.append_footnote_text(text)
                
                elif section_state == "after_table":
                    # 表格后的段落都是脚注
                    current_section.add_footnote_index(p_idx)
                    current_section.append_footnote_text(text)
                
                elif section_state == "collecting_footnote":
                    # 继续收集脚注
                    current_section.add_footnote_index(p_idx)
                    current_section.append_footnote_text(text)

        elif kind == 'table':
            cast(Table, item)
            tbl_idx += 1
            if current_section:
                current_section.set_table(tbl_idx)
                section_state = "after_table"  # 表格后的内容都是脚注

    # 保存最后一个section
    if current_section:
        content.append(current_section)

    return content