from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List
from docx import Document
from docx.document import Document as DocxDocument
from copy import deepcopy
import logging

from triclick_doc_toolset.common import TableItem
from triclick_doc_toolset.common.utils import slugify_text


def _append_paragraph_by_index(src_doc: DocxDocument, dst_doc: DocxDocument, para_idx: int):
    """
    将源文档指定索引的段落以底层 XML 深拷贝追加到目标文档末尾，完整保留样式与版式。

    参数:
    - src_doc: 源 Document。
    - dst_doc: 目标 Document。
    - para_idx: 段落在 `src_doc.paragraphs` 中的索引。

    注意: 不做越界检查；直接操作 `_element.body`。
    """
    dst_doc._element.body.append(deepcopy(src_doc.paragraphs[para_idx]._p))


def _append_table_by_index(src_doc: DocxDocument, dst_doc: DocxDocument, tbl_idx: int):
    # 使用底层 XML 深拷贝表格，保留所有样式/布局
    dst_doc._element.body.append(deepcopy(src_doc.tables[tbl_idx]._element))


def _get_section_label(section: TableItem, src_doc: DocxDocument) -> str:
    """获取用于文件命名的标签，优先使用 TableItem 的类型感知方法。"""
    try:
        return section.get_section_label()
    except Exception:
        return section.label or 'Table'


def _unique_local_path(base_stem: str, label: str, output_dir: str, used_paths: set) -> str:
    """
    生成唯一输出路径，避免文件名冲突。

    参数:
    - base_stem: 基础文件名（不含扩展名）。
    - label: 标签文本。
    - output_dir: 输出目录。
    - used_paths: 已使用路径集合。

    返回: 唯一的文件路径。
    """
    safe_label = slugify_text(label)
    fname = f"{base_stem}.{safe_label}.docx"
    path = os.path.join(output_dir, fname)
    counter = 2
    while path in used_paths or os.path.exists(path):
        fname = f"{base_stem}#{safe_label}_{counter}.docx"
        path = os.path.join(output_dir, fname)
        counter += 1
    used_paths.add(path)
    return path


def split_docx_into_tables_with_copy(original_doc_path: str, docx_content: List[TableItem], output_dir: str) -> List[TableItem]:
    """
    使用元数据（索引）从源文档复制元素到新文档，按起始索引排序。
    每个新文档：标题段落 + 表格 + 脚注段落。
    为确保样式完全一致，使用源文档作为模板并深拷贝段落/表格的XML。
    支持每表格独立脚注：优先读取 TableItem.tables 中的 footnote_indices。
    另外：保留原文档的节属性 sectPr（页边距/页面设置），避免清空正文后导致样式呈现差异。
    命名规则：源文件名 + 提取的标题标签（如 "Table x.x.x"）。
    保存后将输出路径写回到对应的 TableItem.local_path

    参数:
    - original_doc_path: 源 DOCX 路径
    - docx_content: 解析得到的 TableItem 列表
    - output_dir: 输出目录（必填）。

    返回: 更新后的 docx_content（包含每个表格的 local_path）。
    """
    if not output_dir:
        raise ValueError("output_dir 参数必须通过 pipeline params 指定，且不可为空")
    base_stem = Path(original_doc_path).stem

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    original_doc = Document(original_doc_path)
    # 复制原文档的节属性（必须放在 body 的最后）
    original_sect_pr = deepcopy(original_doc._element.body.sectPr)
    used_paths: set = set()
    
    for section in docx_content:
        # 为所有 section 生成独立文档：标题 + [可选表格] + 脚注
        section_label = _get_section_label(section, original_doc)
        # 使用空白文档作为模板，避免每次加载并清空原文档正文的 O(N) 成本
        new_doc = Document()
        body = new_doc._element.body
        # 清除默认空段落（空白模板包含一个默认段落，删除成本为 O(1)）
        for child in list(body.iterchildren()):
            body.remove(child)

        # 标题
        for para_idx in sorted(section.title_indices):
            _append_paragraph_by_index(original_doc, new_doc, para_idx)

        # 表格（若存在）
        if section.table_index is not None:
            _append_table_by_index(original_doc, new_doc, section.table_index)

        # 脚注/正文段落：严格使用该 section 的 footnote_indices（不额外插入空行）
        for para_idx in sorted(section.footnote_indices):
            _append_paragraph_by_index(original_doc, new_doc, para_idx)

        # 重新附加节属性到文档结尾，保证页面设置和段落间距呈现一致
        if original_sect_pr is not None:
            body.append(deepcopy(original_sect_pr))

        local_path = _unique_local_path(base_stem, section_label, output_dir, used_paths)
        new_doc.save(local_path)
        # 将输出路径写回所属 section
        section.local_path = local_path
        logging.getLogger(__name__).info(f"Saved: {local_path}")
    
    # 规则应用移至命令层（DocxFilePartitionCommand），此处仅返回结果
    return docx_content