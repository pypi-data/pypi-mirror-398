import logging
import os
from typing import Dict, Any
import docx
import docx.shared
from converter.base.converter_base import BaseConverter

logger = logging.getLogger(__name__)

class DOCXWriter(BaseConverter):
    """
    Converter for generating DOCX files from text/markdown.
    Uses python-docx to create formatted documents.
    """

    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        """
        Convert text, markdown, or html file to DOCX.
        
        Args:
            input_path: Path to the input file.
            output_path: Path to the output DOCX file.
            **kwargs: Additional arguments.
            
        Returns:
            True if conversion was successful, False otherwise.
        """
        try:
            if not self.validate_input(input_path):
                logger.error(f"Invalid input file: {input_path}")
                return False

            self._ensure_output_directory(output_path)
            logger.info(f"Converting {input_path} to DOCX {output_path}")

            ext = os.path.splitext(input_path)[1].lower()
            
            if ext == '.md':
                return self._convert_markdown(input_path, output_path)
            elif ext == '.html':
                return self._convert_html(input_path, output_path)
            else:
                return self._convert_text(input_path, output_path)

        except Exception as e:
            logger.error(f"Error converting to DOCX: {e}")
            return False

    def _convert_text(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert plain text to DOCX."""
        doc = docx.Document()
        encoding = kwargs.get('encoding', 'utf-8')
        try:
            with open(input_path, 'r', encoding=encoding) as f:
                text = f.read()
        except UnicodeDecodeError:
             with open(input_path, 'r', encoding='utf-8') as f:
                text = f.read()

        for para in text.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.save(output_path)
        return True

    def _convert_markdown(self, input_path: str, output_path: str) -> bool:
        """Convert Markdown to DOCX preserving basic formatting."""
        try:
            import markdown
        except ImportError:
            logger.error("Markdown not installed.")
            return False

        encoding = kwargs.get('encoding', 'utf-8')
        try:
            with open(input_path, 'r', encoding=encoding) as f:
                md_text = f.read()
        except UnicodeDecodeError:
             with open(input_path, 'r', encoding='utf-8') as f:
                md_text = f.read()

        html = markdown.markdown(md_text)
        return self._convert_html_content(html, output_path)

    def _convert_html(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert HTML file to DOCX."""
        encoding = kwargs.get('encoding', 'utf-8')
        try:
            with open(input_path, 'r', encoding=encoding) as f:
                html_content = f.read()
        except UnicodeDecodeError:
             with open(input_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
        return self._convert_html_content(html_content, output_path)

    def _convert_html_content(self, html_content: str, output_path: str) -> bool:
        """Shared logic to convert HTML string to DOCX."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("BeautifulSoup not installed.")
            return False

        soup = BeautifulSoup(html_content, 'html.parser')
        doc = docx.Document()

        # If there's a body, use it, otherwise use the whole soup
        root = soup.body if soup.body else soup

        for element in root.find_all(recursive=False):
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)
            
            elif element.name == 'p':
                p = doc.add_paragraph()
                self._process_inline_elements(p, element)
                
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Bullet')
                    self._process_inline_elements(p, li)
                    
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style='List Number')
                    self._process_inline_elements(p, li)

            elif element.name == 'img':
                src = element.get('src')
                if src and os.path.exists(src):
                    try:
                        # Add image. Defaulting width to something reasonable if not provided
                        # (python-docx handles sizing automatically if not specified, usually fit to page)
                        # We could parse width/height attributes potentially.
                        doc.add_picture(src, width=docx.shared.Inches(6)) 
                    except Exception as e:
                        logger.warning(f"Failed to add image {src} to DOCX: {e}")
                        p = doc.add_paragraph(f"[Image: {os.path.basename(src)}]")
                else:
                     logger.warning(f"Image source not found: {src}")
                     p = doc.add_paragraph(f"[Image: {src}]")

        doc.save(output_path)
        return True

    def _process_inline_elements(self, paragraph, element):
        """Helper to process inline elements like <b>, <i>, <span> within a paragraph."""
        try:
            from converter.processors.style_processor import StyleProcessor
            style_processor = StyleProcessor()
        except ImportError:
            style_processor = None

        for child in element.contents:
            if child.name:
                text = child.get_text()
                run = paragraph.add_run(text)
                
                # Tag based styles
                if child.name in ['strong', 'b']:
                    run.bold = True
                if child.name in ['em', 'i']:
                    run.italic = True
                    
                # Style attribute parsing (mostly for span)
                if child.has_attr('style') and style_processor:
                    style_str = child['style']
                    # Simple parsing: "color: #ff0000; font-weight: bold"
                    for style_part in style_str.split(';'):
                        if ':' in style_part:
                            prop, val = style_part.split(':', 1)
                            prop = prop.strip().lower()
                            val = val.strip()
                            
                            if prop == 'color':
                                rgb = style_processor.css_to_rgb(val)
                                if rgb:
                                    try:
                                        from docx.shared import RGBColor
                                        run.font.color.rgb = RGBColor(*rgb)
                                    except ImportError:
                                        pass
            else:
                paragraph.add_run(str(child))

    def validate_input(self, input_path: str) -> bool:
        """
        Validate if the input file is a valid text/md/html file.
        """
        if not os.path.exists(input_path):
            return False
        return input_path.lower().endswith(('.txt', '.md', '.html'))

    def extract_metadata(self, input_path: str) -> Dict[str, Any]:
        """Metadata extraction not applicable for writer input (it's just text)."""
        return {}
