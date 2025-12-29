import logging
import os
from typing import Dict, Any, Optional
from converter.base.converter_base import BaseConverter

logger = logging.getLogger(__name__)

class PDFConverter(BaseConverter):
    """
    Converter for PDF files.
    Extracts text from PDF files using PyPDF2.
    """

    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        """
        Convert PDF to text or markdown.
        
        Args:
            input_path: Path to the input PDF file.
            output_path: Path to the output file (.txt or .md).
            **kwargs: Additional arguments. 
                - ocr_enabled (bool): Force OCR usage.
                - ocr_fallback (bool): Use OCR if PDF is scanned (default: True).
                - ocr_lang (str): Language for OCR (default: 'auto').
            
        Returns:
            True if conversion was successful, False otherwise.
        """
        temp_pdf_path = None
        current_input_path = input_path

        try:
            if not self.validate_input(input_path):
                logger.error(f"Invalid input file: {input_path}")
                return False

            self._ensure_output_directory(output_path)

            # OCR Logic
            ocr_enabled = kwargs.get('ocr_enabled', False)
            ocr_fallback = kwargs.get('ocr_fallback', True)
            ocr_lang = kwargs.get('ocr_lang', 'auto')

            try:
                from converter.processors.ocr_processor import OCRProcessor
                ocr_processor = OCRProcessor()
                
                should_ocr = ocr_enabled
                if not should_ocr and ocr_fallback and ocr_processor.enabled:
                    if ocr_processor.is_scanned_pdf(input_path):
                        logger.info(f"PDF {input_path} seems scanned. Enabling automatic OCR.")
                        should_ocr = True

                if should_ocr and ocr_processor.enabled:
                    import uuid
                    temp_filename = f"temp_ocr_{uuid.uuid4()}.pdf"
                    temp_pdf_path = os.path.join(os.path.dirname(output_path), temp_filename)
                    
                    logger.info(f"Performing OCR on {input_path}...")
                    if ocr_processor.create_searchable_pdf(input_path, temp_pdf_path, lang=ocr_lang):
                        current_input_path = temp_pdf_path
                        logger.info(f"Using OCR'd PDF for conversion: {temp_pdf_path}")
                    else:
                        logger.warning("OCR failed, proceeding with original file.")
            except ImportError:
                logger.warning("OCRProcessor not available, skipping OCR logic.")
            except Exception as e:
                logger.warning(f"Error during OCR preparation: {e}")
            
            output_ext = os.path.splitext(output_path)[1].lower()
            
            if output_ext == '.md':
                result = self._convert_to_markdown(current_input_path, output_path)
            elif output_ext == '.docx':
                result = self._convert_to_docx(current_input_path, output_path)
            else:
                result = self._convert_to_text(current_input_path, output_path)
                
            return result

        except Exception as e:
            logger.error(f"Error converting PDF: {e}")
            return False
        finally:
            # Cleanup temp file
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path)
                    logger.info(f"Cleaned up temporary OCR file: {temp_pdf_path}")
                except OSError as e:
                    logger.warning(f"Failed to cleanup temp file {temp_pdf_path}: {e}")

    def _convert_to_docx(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to DOCX with images."""
        logger.info(f"Converting PDF {input_path} to DOCX {output_path}")
        try:
            import docx
            from converter.processors.image_processor import ImageProcessor
        except ImportError:
            logger.error("python-docx or ImageProcessor not available.")
            return False

        try:
            # Extract images first
            output_dir = os.path.dirname(output_path)
            images_dir = os.path.join(output_dir, "images")
            processor = ImageProcessor()
            extracted_images = processor.extract_images_from_pdf(input_path, images_dir)
            
            # Map images by page
            images_by_page = {}
            for img in extracted_images:
                page_num = img['page']
                if page_num not in images_by_page:
                    images_by_page[page_num] = []
                images_by_page[page_num].append(img['path'])

            doc = docx.Document()
            with open(input_path, 'rb') as f:
                reader = PdfReader(f)
                for i, page in enumerate(reader.pages):
                    page_num = i + 1
                    
                    # Add Text
                    text = page.extract_text()
                    if text:
                        for line in text.split('\n'):
                            if line.strip():
                                doc.add_paragraph(line.strip())
                    
                    # Add Images for this page
                    # Note: This appends images at the end of the page text. 
                    # Exact positioning requires complex layout analysis.
                    if page_num in images_by_page:
                        for img_path in images_by_page[page_num]:
                            try:
                                doc.add_picture(img_path, width=docx.shared.Inches(5))
                            except Exception as img_err:
                                logger.warning(f"Could not add image {img_path} to DOCX: {repr(img_err)}")

                    doc.add_page_break()
            
            doc.save(output_path)
            logger.info("PDF to DOCX conversion completed.")
            return True
        except Exception as e:
            logger.error(f"Failed to convert to DOCX: {e}")
            raise

    def _convert_to_text(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to plain text."""
        logger.info(f"Converting PDF {input_path} to text {output_path}")
        try:
            from utils.lazy_loader import lazy_import
            PyPDF2 = lazy_import('PyPDF2')
            
            text_content = []
            with open(input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            with open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write("\n\n".join(text_content))
            
            logger.info("PDF to Text conversion completed.")
            return True
        except ImportError as e:
            logger.error(f"PyPDF2 not available: {e}")
            raise ImportError("PyPDF2 is required for PDF conversion. Install with: pip install PyPDF2") from e
        except Exception as e:
            logger.error(f"Failed to convert to text: {e}")
            raise

    def _convert_to_markdown(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to Markdown with structure detection and images."""
        logger.info(f"Converting PDF {input_path} to Markdown {output_path}")
        try:
            from utils.lazy_loader import lazy_import
            from converter.processors.image_processor import ImageProcessor
            PyPDF2 = lazy_import('PyPDF2')
            
            # Extract images first
            output_dir = os.path.dirname(output_path)
            images_dir = os.path.join(output_dir, "images")
            processor = ImageProcessor()
            extracted_images = processor.extract_images_from_pdf(input_path, images_dir)
            
            # Map images by page
            images_by_page = {}
            for img in extracted_images:
                page_num = img['page']
                if page_num not in images_by_page:
                    images_by_page[page_num] = []
                images_by_page[page_num].append(img['path'])

            structured_content = []
            with open(input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                # First pass: Collect all font sizes to determine hierarchy
                all_font_sizes = set()
                pages_data = []
                
                for page in reader.pages:
                    page_content = []
                    
                    def visitor_body(text, cm, tm, fontDict, fontSize):
                        if text and text.strip():
                            page_content.append({
                                'text': text,
                                'size': fontSize
                            })
                            all_font_sizes.add(fontSize)
                            
                    page.extract_text(visitor_text=visitor_body)
                    pages_data.append(page_content)
                
                # Determine header thresholds
                sorted_sizes = sorted(list(all_font_sizes), reverse=True)
                # Simple heuristic: Top 2 sizes are headers if we have enough variance
                h1_size = sorted_sizes[0] if sorted_sizes else 0
                h2_size = sorted_sizes[1] if len(sorted_sizes) > 1 else 0
                
                logger.info(f"Detected font sizes: {sorted_sizes}")
                
                for i, page_items in enumerate(pages_data):
                    page_num = i + 1
                    
                    # Add content
                    for item in page_items:
                        text = item['text']
                        size = item['size']
                        
                        # Basic markdown formatting
                        prefix = ""
                        if size == h1_size and size > 12: # Arbitrary threshold for "big enough"
                            prefix = "# "
                        elif size == h2_size and size > 10:
                            prefix = "## "
                        
                        # Clean up text
                        text = text.strip()
                        if not text:
                            continue
                            
                        structured_content.append(f"{prefix}{text}")
                    
                    # Add Images for this page
                    if page_num in images_by_page:
                        structured_content.append("") # Spacing
                        for img_path in images_by_page[page_num]:
                            rel_path = os.path.relpath(img_path, os.path.dirname(output_path)).replace("\\", "/")
                            structured_content.append(f"![Image]({rel_path})")
                        structured_content.append("")
            
            # Join with newlines
            markdown_text = "\n\n".join(structured_content)
            
            with open(output_path, 'w', encoding='utf-8') as f_out:
                f_out.write(markdown_text)
                
            logger.info("PDF to Markdown conversion completed.")
            return True
            
        except Exception as e:
            logger.error(f"Failed to convert to markdown: {e}")
            raise

    def validate_input(self, input_path: str) -> bool:
        """
        Validate if the input file is a valid PDF.
        """
        if not os.path.exists(input_path):
            return False
            
        try:
            from utils.lazy_loader import lazy_import
            PyPDF2 = lazy_import('PyPDF2')
            
            with open(input_path, 'rb') as f:
                # Try to read the header or just open it with PdfReader
                PyPDF2.PdfReader(f)
            return True
        except Exception:
            return False

    def extract_metadata(self, input_path: str) -> Dict[str, Any]:
        """
        Extract metadata from PDF.
        """
        metadata = {}
        try:
            from utils.lazy_loader import lazy_import
            PyPDF2 = lazy_import('PyPDF2')
            
            with open(input_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                info = reader.metadata
                if info:
                    # PyPDF2 metadata keys often start with /
                    metadata['title'] = info.get('/Title', '')
                    metadata['author'] = info.get('/Author', '')
                    metadata['subject'] = info.get('/Subject', '')
                    metadata['producer'] = info.get('/Producer', '')
                    metadata['creator'] = info.get('/Creator', '')
                
                metadata['num_pages'] = len(reader.pages)
                metadata['is_encrypted'] = reader.is_encrypted
                
        except Exception as e:
            logger.warning(f"Could not extract metadata from {input_path}: {e}")
            
        return metadata
