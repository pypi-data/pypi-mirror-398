import os
import sys
import gc
from datetime import datetime
from typing import Optional

import Orange.data
from Orange.data import Table, Domain, StringVariable
from AnyQt.QtWidgets import QApplication, QFileDialog
from Orange.widgets import widget
from Orange.widgets.widget import Input, MultiInput, Output

# Import des bibliothèques Word
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import qn
if "site-packages/Orange/widgets" in os.path.dirname(os.path.abspath(__file__)).replace("\\", "/"):
    from Orange.widgets.orangecontrib.AAIT.utils.import_uic import uic

else:
    from orangecontrib.AAIT.utils.import_uic import uic

class OWDocumentGenerator(widget.OWWidget):
    """
    Widget simplifié pour générer des documents Word
    - Input 1 : Chemins (template in/out)
    - Input 2 : Remplacements texte (balise -> valeur)
    - Input 3 : Tables à insérer (nom = balise cible)
    """
    name = "Document Generator"
    description = "Generate Word documents with replacements and tables."
    category = "AAIT - TOOLBOX"
    icon = "icons/document_generator.png"
    if "site-packages/Orange/widgets" in os.path.dirname(os.path.abspath(__file__)).replace("\\", "/"):
        icon = "icons_dev/document_generator.png"

    gui = os.path.join(os.path.dirname(os.path.abspath(__file__)), "designer/owgenerate_word.ui")
    want_control_area = False
    priority = 1070

    class Inputs:
        paths = Input("Paths", Orange.data.Table)
        replacements = Input("Replacements", Orange.data.Table)
        tables = MultiInput("Tables", Orange.data.Table, filter_none=True)

    class Outputs:
        result = Output("Result", Orange.data.Table)

    def __init__(self):
        super().__init__()
        self.setFixedWidth(470)
        self.setFixedHeight(600)
        uic.loadUi(self.gui, self)
        # Stockage des données
        self.paths_table = None
        self.replacements_table = None
        self.tables_dict = {}  # Clé = nom de la table, Valeur = Table Orange
        self.doc = None
        self.basedir = None
        self.post_initialized()


        # Recherche du basedir
        self.find_basedir()

        # Paramètres par défaut pour le style
        self.default_font = "Arial"
        self.default_size = 11
        self.table_style = "Table Grid"

    def _resolve_table_style(self, preferred_name: Optional[str]):
        """Return a table style object or name that exists in the document.

        Tries the preferred name first; if not found or not a table style,
        falls back to a common built-in (like 'Table Grid' or 'Table Normal'),
        otherwise the first available table style. Returns None if none found.
        """
        if self.doc is None:
            return None

        styles = getattr(self.doc, "styles", None)
        if styles is None:
            return None

        # Try preferred name
        if preferred_name:
            try:
                st = styles[preferred_name]
                if st.type == WD_STYLE_TYPE.TABLE:
                    return st
            except Exception:
                pass

        # Try common English defaults
        for fallback in ("Table Grid", "Table Normal", "Normal Table"):
            try:
                st = styles[fallback]
                if st.type == WD_STYLE_TYPE.TABLE:
                    return st
            except Exception:
                continue

        # Pick first available table style
        try:
            for st in styles:
                try:
                    if st.type == WD_STYLE_TYPE.TABLE:
                        return st
                except Exception:
                    continue
        except Exception:
            pass

        return None

    def find_basedir(self):
        """Trouve le répertoire de base du workflow"""
        try:
            for obj in gc.get_objects():
                if isinstance(obj, Orange.widgets.widget.OWWidget):
                    env = getattr(obj, "workflowEnv", lambda: {})()
                    self.basedir = env.get("basedir", None)
                    if self.basedir:
                        break
        except Exception as e:
            print(e)
    # === Méthodes Input ===
    @Inputs.paths
    def set_paths(self, data):
        """Table avec Template_Path, Output_Name, Save_Dialog"""
        self.paths_table = data
        self.check_and_run()

    @Inputs.replacements
    def set_replacements(self, data):
        """Table avec colonnes Target et Value"""
        self.replacements_table = data
        self.check_and_run()

    @Inputs.tables
    def set_table(self, index, data):
        """Tables à insérer - le nom de la connexion = balise cible"""
        if data is not None:
            # Essayer de récupérer le nom de la table depuis les métadonnées
            table_name = self.get_table_name(data, index)
            self.tables_dict[table_name] = data
        self.check_and_run()

    @Inputs.tables.insert
    def insert_table(self, index, data):
        """Gère l'insertion d'une nouvelle table"""

        if data is not None:
            table_name = self.get_table_name(data, index)
            self.tables_dict[table_name] = data
        self.check_and_run()

    @Inputs.tables.remove
    def remove_table(self, index):
        """Gère la suppression d'une table"""
        # Trouver et supprimer la table correspondante
        key=list(self.tables_dict.keys())
        del self.tables_dict[key[index]]
        self.check_and_run()

    def get_table_name(self, data, index):
        """
        Essaie de récupérer un nom intelligent pour la table
        Priorité : nom de la table > attribut name > index par défaut
        """
        # Essayer de récupérer le nom depuis les attributs de la table
        if hasattr(data, 'name') and data.name:
            return data.name

        # Chercher un attribut 'table_name' dans les metas
        if data.domain.metas:
            for meta in data.domain.metas:
                if meta.name.lower() in ['table_name', 'nom_table', 'name']:
                    if len(data) > 0:
                        name = str(data[0][meta])
                        if name and name != "?":
                            return name

        # Nom par défaut basé sur l'index
        return f"tableau {index + 1}"

    def check_and_run(self):
        """Vérifie si les données nécessaires sont présentes et lance le traitement"""
        if self.paths_table is not None:
            self.run()

    def run(self):
        """Méthode principale de traitement"""
        self.error("")
        self.warning("")

        try:
            # 1. Charger le template
            template_path = self.get_template_path()
            if not template_path:
                self.error("Chemin du template non trouvé")
                return

            self.doc = Document(template_path)

            # 2. Effectuer les remplacements de texte
            if self.replacements_table is not None:
                self.perform_text_replacements()

            # 3. Insérer les tableaux
            if self.tables_dict:
                self.insert_all_tables()

            # 4. Sauvegarder le document
            output_path = self.save_document()

            # 5. Envoyer le résultat
            self.send_result(output_path)


        except Exception as e:
            self.error(f"Erreur : {str(e)}")
            import traceback
            print(traceback.format_exc())

    def get_template_path(self) -> Optional[str]:
        """Récupère le chemin du template depuis la table paths"""
        if self.paths_table is None or len(self.paths_table) == 0:
            return None

        # Chercher Template_Path dans différents emplacements
        template_path = None
        row = self.paths_table[0]

        # Dans les attributs
        for attr in self.paths_table.domain.attributes:
            if attr.name == "Template_Path":
                template_path = str(row[attr])
                break

        # Dans les metas si pas trouvé
        if not template_path:
            for i, meta in enumerate(self.paths_table.domain.metas):
                if meta.name == "Template_Path":
                    template_path = str(self.paths_table.metas[0, i])
                    break

        if not template_path:
            self.error("Colonne 'Template_Path' non trouvée")
            return None

        # Gérer les chemins relatifs
        if self.basedir and not os.path.isabs(template_path):
            template_path = template_path.replace("\\", os.sep).replace("/", os.sep)
            template_path = os.path.join(self.basedir, template_path)

        if not os.path.exists(template_path):
            self.error(f"Template non trouvé : {template_path}")
            return None

        return template_path

    def perform_text_replacements(self):
        """Effectue tous les remplacements de texte"""
        if self.replacements_table is None:
            return

        # Identifier les colonnes Target et Value
        target_col = None
        value_col = None

        # Chercher dans les attributs
        for attr in self.replacements_table.domain.attributes:
            if attr.name in ["Target", "Balise", "Cible"]:
                target_col = attr
            elif attr.name in ["Value", "Valeur", "Replacement", "Remplacer_par"]:
                value_col = attr

        # Chercher dans les metas si pas trouvé
        if not target_col or not value_col:
            for meta in self.replacements_table.domain.metas:
                if not target_col and meta.name in ["Target", "Balise", "Cible"]:
                    target_col = meta
                elif not value_col and meta.name in ["Value", "Valeur", "Replacement", "Remplacer_par"]:
                    value_col = meta

        if not target_col or not value_col:
            self.warning("Colonnes 'Target' et 'Value' non trouvées dans la table Replacements")
            return

        # Effectuer les remplacements
        count = 0
        for row_idx in range(len(self.replacements_table)):
            # Récupérer les valeurs
            if target_col in self.replacements_table.domain.attributes:
                target = str(self.replacements_table[row_idx][target_col])
            else:
                meta_idx = list(self.replacements_table.domain.metas).index(target_col)
                target = str(self.replacements_table.metas[row_idx, meta_idx])

            if value_col in self.replacements_table.domain.attributes:
                value = str(self.replacements_table[row_idx][value_col])
            else:
                meta_idx = list(self.replacements_table.domain.metas).index(value_col)
                value = str(self.replacements_table.metas[row_idx, meta_idx])

            # Gérer les valeurs spéciales
            if value == "date_today":
                value = datetime.today().strftime("%d/%m/%Y")
            elif value == "date_today_long":
                value = datetime.today().strftime("%d %B %Y")

            # Effectuer le remplacement
            self.replace_text_in_doc(target, value)
            count += 1


    def replace_text_in_doc(self, placeholder: str, value: str):
        """Remplace du texte dans tout le document"""
        def replace_in_paragraph(paragraph):
            full_text = "".join(run.text for run in paragraph.runs)

            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))

                # Supprimer les runs existants
                while len(paragraph.runs) > 0:
                    paragraph._element.remove(paragraph.runs[0]._element)

                # Ajouter le nouveau run avec style par défaut
                run = paragraph.add_run(full_text)
                run.font.name = self.default_font
                run.font.size = Pt(self.default_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)

        # Remplacer dans les paragraphes
        for paragraph in self.doc.paragraphs:
            replace_in_paragraph(paragraph)

        # Remplacer dans les tableaux
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

    def _extract_target_tag(self, data) -> Optional[str]:
        """Extrait une balise cible depuis les colonnes de la table si présente.

        Recherche d'abord dans les attributs puis dans les métas des noms
        [Target, Balise, Cible, Tag, Balise_Table]. Retourne la balise sans crochets.
        """
        try:
            # Attributs
            for attr in getattr(data.domain, 'attributes', []):
                if getattr(attr, 'name', None) in ["Target", "Balise", "Cible", "Tag", "Balise_Table"]:
                    if len(data) > 0:
                        tag = str(data[0][attr]).strip()
                        if tag and tag != "?":
                            return tag.strip("[]").strip()
            # Metas
            for i, meta in enumerate(getattr(data.domain, 'metas', [])):
                if getattr(meta, 'name', None) in ["Target", "Balise", "Cible", "Tag", "Balise_Table"]:
                    if len(data) > 0:
                        tag = str(data.metas[0, i]).strip()
                        if tag and tag != "?":
                            return tag.strip("[]").strip()
        except Exception:
            pass
        return None

    def insert_all_tables(self):
        """Insère tous les tableaux aux emplacements définis"""
        count = 0
        for table_name, table_data in self.tables_dict.items():
            # La clé est le nom de la balise cible
            target = str(table_name).strip()

            if self.insert_table_in_doc(target, table_data):
                count += 1  # Compteur pour l'app



    def insert_table_in_doc(self, tag: str, table_data: Table) -> bool:
        """Insère un tableau dans le document"""
        # Préparer les données
        headers = []
        data = []

        # Récupérer tous les en-têtes
        for attr in table_data.domain.attributes:
            headers.append(str(attr.name))
        for meta in table_data.domain.metas:
            headers.append(str(meta.name))

        # Récupérer toutes les données
        for i in range(len(table_data)):
            row_data = []
            # Attributs
            for attr in table_data.domain.attributes:
                value = table_data[i][attr]
                row_data.append(str(value) if value is not None else "")
            # Metas
            for j, meta in enumerate(table_data.domain.metas):
                value = table_data.metas[i, j]
                row_data.append(str(value) if value is not None else "")
            data.append(row_data)

        # Chercher et remplacer le tag (accepte [tag] ou tag)
        inserted = False
        variants = []
        if tag:
            variants.append(tag)
            if tag.startswith("[") and tag.endswith("]"):
                variants.append(tag[1:-1])
            else:
                variants.append(f"[{tag}]")

        # Ajouter une balise potentielle issue des métadonnées de la table
        try:
            meta_tag = self._extract_target_tag(table_data)
            if meta_tag:
                if meta_tag not in variants:
                    variants.append(meta_tag)
                bracketed = f"[{meta_tag}]"
                if bracketed not in variants:
                    variants.append(bracketed)
        except Exception:
            pass

        for p in self.doc.paragraphs:
            match_token = next((t for t in variants if t in p.text), None)
            if match_token:
                p.text = p.text.replace(match_token, "")

                if len(data) > 0:
                    # Créer le tableau
                    table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
                    # Apply style safely with fallbacks
                    try:
                        resolved_style = self._resolve_table_style(self.table_style)
                        if resolved_style is not None:
                            table.style = resolved_style
                    except Exception as e:
                        # If anything goes wrong, continue without setting style
                        print(f"Warning: unable to apply table style '{self.table_style}': {e}")

                    # En-têtes
                    for j, header in enumerate(headers):
                        cell = table.cell(0, j)
                        self.stylize_cell(cell, header)


                    # Données
                    for i, row_data in enumerate(data):
                        for j, value in enumerate(row_data):
                            cell = table.cell(i + 1, j)
                            self.stylize_cell(cell, value)

                    # Insérer le tableau
                    p._element.addnext(table._element)
                    inserted = True
                    break

        # Chercher dans les cellules de tableaux si non inséré
        if not inserted:
            for t in self.doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            match_token = next((mt for mt in variants if mt in p.text), None)
                            if match_token:
                                p.text = p.text.replace(match_token, "")

                                if len(data) > 0:
                                    table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
                                    # Apply style safely with fallbacks
                                    try:
                                        resolved_style = self._resolve_table_style(self.table_style)
                                        if resolved_style is not None:
                                            table.style = resolved_style
                                    except Exception as e:
                                        print(f"Warning: unable to apply table style '{self.table_style}': {e}")

                                    # En-têtes
                                    for j, header in enumerate(headers):
                                        cell_h = table.cell(0, j)
                                        self.stylize_cell(cell_h, header)

                                    # Données
                                    for i, row_data in enumerate(data):
                                        for j, value in enumerate(row_data):
                                            cell_d = table.cell(i + 1, j)
                                            self.stylize_cell(cell_d, value)

                                    # Insérer le tableau (imbriqué dans la cellule)
                                    p._element.addnext(table._element)
                                    inserted = True
                                    break
                        if inserted:
                            break
                    if inserted:
                        break
                if inserted:
                    break

        return inserted

    def stylize_cell(self, cell, text: str, bold: bool = False):
        """Stylise une cellule de tableau"""
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = bold
        run.font.size = Pt(9)
        run.font.name = self.default_font

        # Assurer l'application de la police
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), self.default_font)
        rFonts.set(qn('w:hAnsi'), self.default_font)



    def save_document(self) -> str:
        """Sauvegarde le document"""
        output_name = "output.docx"
        save_dialog = False

        if self.paths_table and len(self.paths_table) > 0:
            row = self.paths_table[0]

            # Chercher Output_Name
            for attr in self.paths_table.domain.attributes:
                if attr.name == "Output_Path":
                    output_name = str(row[attr])
                    break
            else:
                for i, meta in enumerate(self.paths_table.domain.metas):
                    if meta.name == "Output_Path":
                        output_name = str(self.paths_table.metas[0, i])
                        break

            # Chercher Save_Dialog
            for attr in self.paths_table.domain.attributes:
                if attr.name == "Save_Dialog":
                    save_dialog = str(row[attr]).lower() == "true"
                    break
            else:
                for i, meta in enumerate(self.paths_table.domain.metas):
                    if meta.name == "Save_Dialog":
                        save_dialog = str(self.paths_table.metas[0, i]).lower() == "true"
                        break

        # Chemin de sauvegarde
        if save_dialog:
            app = QApplication.instance()
            if app is None:
                app = QApplication(sys.argv)

            file_path, _ = QFileDialog.getSaveFileName(
                None,
                "Choisissez l'emplacement de sauvegarde",
                output_name,
                "Documents Word (*.docx);;Tous les fichiers (*)"
            )

            if not file_path:
                file_path = os.path.join(self.basedir or "", "output", output_name)
        else:
            file_path = os.path.join(self.basedir or "", "output", output_name)

        # Assurer l'extension .docx
        if not file_path.endswith(".docx"):
            file_path += ".docx"

        # Créer le répertoire si nécessaire
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        # Sauvegarder
        self.doc.save(file_path)

        return file_path
    def post_initialized(self):
        pass
    def send_result(self, output_path: str):
        """Envoie le résultat en sortie"""
        file_var = StringVariable("Generated_File")
        domain = Domain([], metas=[file_var])

        result_table = Table.from_list(
            domain=domain,
            rows=[[output_path]]
        )

        self.Outputs.result.send(result_table)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    widget = OWDocumentGenerator()
    widget.show()
    if hasattr(app, "exec"):
        app.exec()
    else:
        app.exec_()
