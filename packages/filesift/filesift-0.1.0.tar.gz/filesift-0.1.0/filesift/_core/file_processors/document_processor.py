from pathlib import Path
from typing import Dict, Any, Set, Optional
from openai import OpenAI
import tiktoken

from .base import BaseFileProcessor
from filesift._config.config import config_dict

class DocumentProcessor(BaseFileProcessor):
    """Processor for handling document files (PDF, DOCX, ODT)"""
    
    def __init__(self, max_tokens_for_summary: int = 2000):
        super().__init__()
        self.supported_extensions: Set[str] = {
            ".pdf", ".docx", ".odt"
        }
        
        # Initialize OpenAI client with config
        llm_api_key = config_dict["llm"]["LLM_API_KEY"]
        llm_base_url = config_dict["llm"]["LLM_BASE_URL"]
        if llm_base_url and len(llm_base_url) > 0:
            self.client = OpenAI(api_key=llm_api_key, base_url=llm_base_url)
        else:
            self.client = OpenAI(api_key=llm_api_key)
        
        self.max_tokens_for_summary = max_tokens_for_summary

        try:
            self.encoding = tiktoken.get_encoding("cl100k_base")
        except:
            self.encoding = None
        
        self.pdf_available = False
        self.docx_available = False
        self.odt_available = False
        
        try:
            import PyPDF2
            self.pdf_available = True
            self.PyPDF2 = PyPDF2
        except ImportError:
            pass
        
        try:
            from docx import Document as DocxDocument
            self.docx_available = True
            self.DocxDocument = DocxDocument
        except ImportError:
            pass
        
        try:
            import odf.opendocument
            import odf.text
            self.odt_available = True
            self.odf = odf
        except ImportError:
            pass
        
    def can_handle(self, file_path: Path) -> bool:
        ext = file_path.suffix.lower()
        if ext == ".pdf" and not self.pdf_available:
            return False
        if ext == ".docx" and not self.docx_available:
            return False
        if ext == ".odt" and not self.odt_available:
            return False
        return ext in self.supported_extensions
    
    def _truncate_content_for_summary(self, content: str) -> str:
        """Truncate content to fit within token limit for LLM summarization"""
        if self.encoding is None:
            max_chars = self.max_tokens_for_summary * 4
            if len(content) <= max_chars:
                return content
            return content[:max_chars] + "\n\n[... content truncated for summary ...]"
        
        tokens = self.encoding.encode(content)
        if len(tokens) <= self.max_tokens_for_summary:
            return content
        
        truncated_tokens = tokens[:self.max_tokens_for_summary]
        truncated_content = self.encoding.decode(truncated_tokens)
        
        return truncated_content + "\n\n[... content truncated for summary ...]"
    
    def _generate_llm_summary(self, file_path: Path, content: str, doc_type: str) -> str:
        """Generate summary using LLM"""
        file_info = self.extract_file_info(file_path)
        content_for_summary = self._truncate_content_for_summary(content)
        
        prompt = (
            f"Here is some information about a {doc_type} document file:\n{file_info}\n"
            "Analyze this document content and provide a concise, technical summary for a search index. "
            "Focus on the main topics, key information, structure, and important details. "
            "Do not include the full content or conversational filler. "
            "Start directly with the description.\n"
            f"```\n{content_for_summary}\n```"
        )
        messages = [{"role": "user", "content": prompt}]
        
        try:
            model = config_dict["models"]["MAIN_MODEL"]
            response = self.client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0
            )
            return response.choices[0].message.content
        except Exception as e:
            self.logger.warning(f"LLM summarization failed for {file_path}: {str(e)}")
            return f"{doc_type} document: {file_path.name}\n{file_info}"
    
    def process(self, file_path: Path) -> Dict[str, Any]:
        """Process a document file"""
        try:
            ext = file_path.suffix.lower()
            
            if ext == ".pdf":
                content = self._process_pdf(file_path)
            elif ext == ".docx":
                content = self._process_docx(file_path)
            elif ext == ".odt":
                content = self._process_odt(file_path)
            else:
                content = ""
            
            doc_type = self._detect_document_type(file_path)
            if content:
                summary = self._generate_llm_summary(file_path, content, doc_type)
            else:
                summary = f"{self.extract_file_info(file_path)}\n\nUnsupported document format."
            
            return {
                "content": content,
                "summary": summary,
                "file_type": "document",
                "document_type": doc_type,
                "metadata": {
                    "path": str(file_path),
                    "size": file_path.stat().st_size,
                    "modified": file_path.stat().st_mtime,
                }
            }
        except Exception as e:
            self.logger.error(f"Error processing document file {file_path}: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: Path) -> str:
        """Process PDF file and return extracted text content"""
        if not self.pdf_available:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
        
        try:
            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = self.PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages[:10]):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                
                if num_pages > 10:
                    text_parts.append(f"\n[... {num_pages - 10} more pages ...]")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            self.logger.warning(f"Error processing PDF {file_path}: {str(e)}")
            return ""
    
    def _process_docx(self, file_path: Path) -> str:
        """Process DOCX file and return extracted text content"""
        if not self.docx_available:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            doc = self.DocxDocument(file_path)
            
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(paragraphs)
            
            table_texts = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(row_cells))
                if table_rows:
                    table_texts.append("\n".join(table_rows))
            
            if table_texts:
                content += "\n\n--- Tables ---\n" + "\n\n".join(table_texts)
            
            return content
        except Exception as e:
            self.logger.warning(f"Error processing DOCX {file_path}: {str(e)}")
            return ""
    
    def _process_odt(self, file_path: Path) -> str:
        """Process ODT file and return extracted text content"""
        if not self.odt_available:
            raise ImportError("odfpy is required for ODT processing. Install with: pip install odfpy")
        
        try:
            doc = self.odf.opendocument.load(file_path)
            
            paragraphs = []
            for para in doc.getElementsByType(self.odf.text.P):
                text = ""
                for node in para.childNodes:
                    if node.nodeType == 3:
                        text += node.data
                if text.strip():
                    paragraphs.append(text.strip())
            
            return "\n".join(paragraphs)
        except Exception as e:
            self.logger.warning(f"Error processing ODT {file_path}: {str(e)}")
            return ""
    
    def _detect_document_type(self, file_path: Path) -> str:
        """Detect the type of document file"""
        ext = file_path.suffix.lower()
        doc_types = {
            ".pdf": "PDF",
            ".docx": "DOCX",
            ".odt": "ODT"
        }
        return doc_types.get(ext, "unknown")

