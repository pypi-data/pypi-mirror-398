"""This module contains the SSPDocParser class,
which is used to parse an SSP .docx file and return a list of dictionaries representing the table data.
"""

import logging
import re
import zipfile
from collections import defaultdict
from typing import Any, Dict, List, Optional

import docx
import lxml.etree as etree

logger = logging.getLogger("regscale")


class SSPDocParser:
    """
    Parses an SSP .docx file and returns a list of dictionaries representing the table data.
    """

    docx_path: str
    xml_content: bytes
    tables: List
    text: Dict

    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = docx.Document(self.docx_path)

    def parse(self) -> List[Dict]:
        """
        Parses the .docx file and returns a list of dictionaries representing the table data.
        :return: A list of dictionaries, each representing a table's data.
        :rtype: List[Dict]
        """
        self.tables = self.parse_xml_for_tables(self.docx_to_xml(self.docx_path))
        self.text = self.group_paragraphs_by_headings()
        return self.tables

    def get_figure_captions(self) -> Dict:
        """
        Fetches the figure captions from the .docx file.
        BETA: This method is still in development and may not work as expected.
        :return: A dictionary mapping image blobs to their captions.
        :rtype: Dict
        """
        doc = self.doc
        captions = {}
        # Iterate through all the paragraphs to find the figure captions
        for paragraph in doc.paragraphs:
            # Check if the paragraph contains an image
            if "graphicData" in paragraph._element.xml:
                # Image found, now find the subsequent paragraphs for the caption
                next_paragraph = paragraph._element.getnext()
                while next_paragraph is not None and next_paragraph.tag.endswith("p"):
                    text = next_paragraph.text.strip()
                    if text.startswith("Figure"):
                        # Found the caption, map it to the image blob
                        r_id = paragraph._element.xpath(".//a:blip/@r:embed")[0]
                        captions[r_id] = text
                        break
                    next_paragraph = next_paragraph.getnext()

        return captions

    def group_paragraphs_by_headings(self) -> Dict:
        """
        Groups the paragraphs in the .docx file by their headings.
        :return: A dictionary mapping headings to their paragraphs.
        :rtype: Dict
        """
        grouped_text = defaultdict(list)
        current_heading = None

        for para in self.doc.paragraphs:
            # Check if the paragraph is a heading based on its style
            if para.style.name.startswith("Heading"):
                # A new heading is encountered; update the current heading variable
                current_heading = para.text
            elif current_heading:
                # Add the paragraph text under the current heading key
                grouped_text[current_heading].append(para.text)
            # Note: Paragraphs before the first heading won't be included

        return dict(grouped_text)

    @staticmethod
    def docx_to_xml(docx_path: str) -> bytes:
        """
        Converts a .docx file to XML.
        :param str docx_path: The path to the .docx file.
        :return: The XML content of the .docx file.
        :rtype: bytes
        """
        with zipfile.ZipFile(docx_path, "r") as docx_zip:
            xml_content = docx_zip.read("word/document.xml")
        return xml_content

    @staticmethod
    def parse_checkbox_string(text: str) -> List[Dict]:
        """
        Parses a string like "Implementation Status (check all that apply): ☒ Implemented ..." and returns a list of dictionaries.

        :param str text: The text to parse.
        :return: A list of dictionaries, each representing a status option and its status.
        :rtype: List[Dict]
        """

        # Regular expression pattern to match status options change .*? to [^☒☐]* to match any character except ☒ and ☐
        pattern = r"([☒☐]\s*([^☒☐]*))(?=[☒☐]|$)"

        # Find all non-overlapping matches with the pattern
        matches = re.findall(pattern, text)

        # Process matches and create dictionaries
        dict_list = []
        for match in matches:
            status_indicator, option_text = match
            dict_list.append({option_text.strip(): status_indicator.startswith("☒")})

        return dict_list

    def parse_xml_for_tables(self, xml_content: bytes) -> List[Dict]:
        """
        Parses the XML content for tables and returns a list of dictionaries representing the table data.

        :param bytes xml_content: The XML content to parse.
        :return: A list of dictionaries, each representing a table's data.
        :rtype: List[Dict]
        """
        # Parse the XML content
        tree = etree.ElementTree(etree.fromstring(xml_content))
        root = tree.getroot()

        # Define namespace map to handle XML namespaces
        # Extended namespace support for page break detection
        namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",  # noqa
            "w14": "http://schemas.microsoft.com/office/word/2010/wordml",  # noqa
            "w15": "http://schemas.microsoft.com/office/word/2012/wordml",  # noqa
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",  # noqa
        }
        return self.parse_tables(root, namespaces)

    @staticmethod
    def extract_table_headers(table: Any, namespaces: Any) -> List:
        """
        Extracts headers from a table.

        :param Any table: The table element to extract headers from.
        :param Any namespaces: The XML namespaces.
        :return: A list of headers.
        :rtype: List
        """
        headers = []
        for cell in table.findall(".//w:tr[1]//w:tc", namespaces):
            cell_text = "".join(node.text for node in cell.findall(".//w:t", namespaces) if node.text)
            headers.append(cell_text.strip() if cell_text else "")
        return headers

    @staticmethod
    def has_page_break(para: Any, namespaces: Dict) -> bool:
        """
        Check if a paragraph contains a page break marker.

        :param Any para: The paragraph element to check.
        :param Dict namespaces: The XML namespaces.
        :return: True if a page break is found, False otherwise.
        :rtype: bool
        """
        # Check for lastRenderedPageBreak (automatic page break marker)
        page_breaks = para.findall(".//w:lastRenderedPageBreak", namespaces)
        # Check for manual page breaks (w:br with type='page')
        manual_breaks = para.findall(".//w:br", namespaces)
        manual_page_breaks = [br for br in manual_breaks if br.get("{%s}type" % namespaces.get("w", "")) == "page"]
        return bool(page_breaks or manual_page_breaks)

    @staticmethod
    def _process_run_element(child: Any, namespaces: Dict) -> Optional[str]:
        """
        Process a single element within a run and return its text representation.

        :param Any child: The child element to process.
        :param Dict namespaces: The XML namespaces.
        :return: The text representation or None if no text.
        :rtype: Optional[str]
        """
        tag = etree.QName(child).localname

        if tag == "t" and child.text:
            return child.text
        elif tag == "br":
            br_type = child.get("{%s}type" % namespaces.get("w", ""))
            if br_type == "page":
                logger.debug("Manual page break found within cell content")
            return "\n"
        elif tag == "lastRenderedPageBreak":
            logger.debug("Automatic page break marker found, continuing")
            return None
        elif tag == "tab":
            return "\t"
        return None

    @staticmethod
    def _extract_paragraph_text(para: Any, namespaces: Dict) -> str:
        """
        Extract text from a paragraph element, processing all runs.

        :param Any para: The paragraph element to process.
        :param Dict namespaces: The XML namespaces.
        :return: The text content from the paragraph.
        :rtype: str
        """
        para_parts = []
        for run in para.findall(".//w:r", namespaces):
            for child in run:
                text = SSPDocParser._process_run_element(child, namespaces)
                if text is not None:
                    para_parts.append(text)
        return "".join(para_parts)

    @staticmethod
    def fetch_cell_text(cell: Any, namespaces: Dict) -> str:
        """
        Fetches the text from a table cell, handling page breaks and line breaks.

        This method properly handles content that spans page breaks by:
        - Detecting w:lastRenderedPageBreak elements and continuing to read
        - Converting w:br elements to appropriate separators
        - Preserving paragraph structure

        :param Any cell: The cell element to fetch text from.
        :param Dict namespaces: The XML namespaces.
        :return: The text from the cell.
        :rtype: str
        """
        cell_parts = []
        page_break_count = 0

        for para in cell.findall(".//w:p", namespaces):
            if SSPDocParser.has_page_break(para, namespaces):
                page_break_count += 1
                logger.debug("Page break detected in cell, continuing to read content")

            para_text = SSPDocParser._extract_paragraph_text(para, namespaces)
            if para_text.strip():
                cell_parts.append(para_text.strip())

        if page_break_count > 0:
            logger.debug("Cell contained %d page break(s), all content extracted", page_break_count)

        return " ".join(cell_parts) + " " if cell_parts else ""

    def extract_vertical_row_data(self, table: any, namespaces: any) -> List[Dict]:
        """
        Extracts data from a table organized vertically, with the first cell as the key
        and the second cell as the value.

        :param any table: The table element to extract data from.
        :param any namespaces: The XML namespaces.
        :return: A list of dictionaries representing the table's data.
        :rtype: List[Dict]
        """
        dicts_list = []
        for row in table.findall(".//w:tr", namespaces)[1:]:
            cells = row.findall(".//w:tc", namespaces)
            if len(cells) >= 2:
                vertical_data = {
                    self.fetch_cell_text(cells[0], namespaces)
                    .strip(): self.fetch_cell_text(cells[1], namespaces)
                    .strip()
                }
                dicts_list.append(vertical_data)
        return dicts_list

    def extract_row_data(self, row: any, headers: any, namespaces: any) -> Dict:
        """
        Extracts data from a table row.

        :param any row: The row element to extract data from.
        :param any headers: The headers of the table.
        :param any namespaces: The XML namespaces.
        :return: A dictionary representing the row's data.
        :rtype: Dict
        """
        row_data = {}
        for header, cell in zip(headers, row.findall(".//w:tc", namespaces)):
            cell_text = self.fetch_cell_text(cell, namespaces)
            if "☒" in cell_text or "☐" in cell_text:
                row_data[header] = self.parse_checkbox_string(cell_text)
            else:
                row_data[header] = cell_text.strip() if cell_text else None
        return row_data

    def fetch_preceding_text(self) -> List[str]:
        """
        Fetches the text immediately preceding a table.
        """
        preceding_texts = []
        for element in self.doc.element.body:
            if element.tag.endswith("tbl"):
                para = element.getprevious()
                text = para.text if para is not None and para.tag.endswith("p") else ""
                preceding_texts.append(text)
        return preceding_texts

    def parse_tables(self, root: any, namespaces: any) -> List[Dict]:
        """
        Parses all tables in the XML root.

        :param any root: The XML root element.
        :param any namespaces: The XML namespaces.
        :return: A list of dictionaries, each representing a table's data.
        :rtype: List[Dict]
        """
        vertical_tables = [
            "Identification of Organization that Prepared this Document".lower(),
            "Identification of Cloud Service Provider".lower(),
            "System Owner Information".lower(),
            "System Information".lower(),
            "System Component Information".lower(),
            "ISSO (or Equivalent) Point of Contact".lower(),
        ]
        tables_list = []
        preceding_text = self.fetch_preceding_text()
        for i, table in enumerate(root.findall(".//w:tbl", namespaces)):
            tables_dicts_list = {}
            headers = self.extract_table_headers(table, namespaces)
            table_data = []
            tables_dicts_list["preceding_text"] = preceding_text[i] if i < len(preceding_text) else ""
            # Check if this is a vertical table:
            is_vertical = False
            # - Single header that matches vertical_tables list
            first_option = len(headers) == 1 and headers[0].lower() in vertical_tables
            # - OR two headers where first matches and second is empty/whitespace
            second_option = len(headers) == 2 and headers[0].lower() in vertical_tables and not headers[1].strip()
            if first_option or second_option:
                is_vertical = True

            if is_vertical:
                table_data = self.extract_vertical_row_data(table, namespaces)
                tables_dicts_list["table_data"] = table_data
                # tables_dicts_list.append({headers[0].lower(): table_data})
            else:
                for row in table.findall(".//w:tr", namespaces)[1:]:  # Skip header row
                    row_data = self.extract_row_data(row, headers, namespaces)
                    table_data.append(row_data)
                tables_dicts_list["table_data"] = table_data
            tables_list.append(tables_dicts_list)
        return tables_list
