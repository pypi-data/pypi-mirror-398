from typing import Dict, Any
from .base import BaseDocumentHandler

class DocxHandler(BaseDocumentHandler):
    """Handler for Word documents (.docx) using python-docx."""

    def __init__(self, file_path: str):
        super().__init__(file_path)
        self._doc = None
        self._text = ""
        self._metadata = {}

    def load(self) -> None:
        """Load and parse the DOCX file with OllamaVision for embedded images."""
        try:
            import docx
        except ImportError:
            raise RuntimeError("python-docx is not installed.")

        try:
            self._doc = docx.Document(self.file_path)
            
            # Extract text from paragraphs
            text_parts = [para.text for para in self._doc.paragraphs]
            
            # Extract text from tables
            for table in self._doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)
            
            self._text = "\n\n".join(text_parts)
            
            # Extract and OCR images using OllamaVision
            try:
                from ..services.ollama_vision import OllamaVisionService
                from PIL import Image
                from io import BytesIO
                
                image_texts = []
                for rel in self._doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            image = Image.open(BytesIO(image_data))
                            ocr_text = OllamaVisionService.get_text_from_pil_image(image)
                            if ocr_text.strip():
                                image_texts.append(ocr_text)
                        except:
                            continue
                
                if image_texts:
                    self._text += "\n\n[Embedded Images Text via OllamaVision]:\n" + "\n".join(image_texts)
                    self._metadata["image_ocr_engine"] = "ollama_vision"
            except ImportError:
                pass  # OllamaVision not available
            
            # Extract core properties
            core_props = self._doc.core_properties
            self._metadata.update({
                "author": core_props.author,
                "created": str(core_props.created),
                "modified": str(core_props.modified),
                "title": core_props.title,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "last_modified_by": core_props.last_modified_by,
            })
            # Remove None values
            self._metadata = {k: v for k, v in self._metadata.items() if v is not None}
            
        except Exception as e:
            raise RuntimeError(f"Failed to load DOCX file {self.file_path}: {e}")

    def get_text(self) -> str:
        if self._doc is None:
            self.load()
        return self._text

    def get_metadata(self) -> Dict[str, Any]:
        if self._doc is None:
            self.load()
        return self._metadata
