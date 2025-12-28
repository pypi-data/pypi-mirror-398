import os
import shutil
import html

def convert_to_html(text_content: str, output_path: str, title: str = None) -> None:
    """
    Convert text content to HTML using native Python string manipulation.
    """
    if not title:
        title = os.path.basename(output_path)
    
    # Basic Markdown-like parsing (very simple)
    # 1. Escape HTML
    safe_text = html.escape(text_content)
    
    # 2. Convert newlines to <br> or wrap paragraphs
    #    Simple approach: Split by double newline for paragraphs, single for line breaks
    paragraphs = safe_text.split('\n\n')
    html_paragraphs = []
    for p in paragraphs:
        if not p.strip(): continue
        # Handle headers (lines starting with #)
        if p.startswith('# '):
            html_paragraphs.append(f"<h1>{p[2:]}</h1>")
        elif p.startswith('## '):
            html_paragraphs.append(f"<h2>{p[3:]}</h2>")
        elif p.startswith('### '):
            html_paragraphs.append(f"<h3>{p[4:]}</h3>")
        elif p.startswith('- '):
            # List items
            items = p.split('\n')
            list_html = "<ul>"
            for item in items:
                if item.startswith('- '):
                    list_html += f"<li>{item[2:]}</li>"
                else:
                    list_html += f"{item}<br>" # continuation
            list_html += "</ul>"
            html_paragraphs.append(list_html)
        else:
            # Regular paragraph, convert single newlines to <br>
            p_content = p.replace('\n', '<br>')
            html_paragraphs.append(f"<p>{p_content}</p>")
            
    body_content = "\n".join(html_paragraphs)

    html_template = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        pre {{
            background: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
        }}
        code {{
            font-family: "Consolas", "Monaco", monospace;
            background: #f4f4f4;
            padding: 2px 5px;
            border-radius: 3px;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            margin: 0;
            padding-left: 20px;
            color: #666;
        }}
    </style>
</head>
<body>
{body_content}
</body>
</html>
    """
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_template)

def convert_to_docx(text_content: str, output_path: str) -> None:
    """
    Convert text content to DOCX using python-docx.
    """
    try:
        import docx
        from docx.shared import Pt
    except ImportError:
         raise RuntimeError("python-docx is not installed. Run 'pip install python-docx'.")

    doc = docx.Document()
    
    # Basic parsing
    lines = text_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    doc.save(output_path)
