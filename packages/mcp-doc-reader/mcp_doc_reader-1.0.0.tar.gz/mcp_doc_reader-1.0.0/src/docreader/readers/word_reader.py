"""Word document reader for .docx files."""

import os
import docx


def get_word_content(file_path: str) -> str:
    """
    Read content from a Word document.
    
    Args:
        file_path: Path to the Word file (.docx).
        
    Returns:
        Document content as a string.
        
    Note:
        .doc files are not supported by python-docx.
    """
    if not os.path.isabs(file_path):
        file_path = os.path.abspath(file_path)

    if not os.path.exists(file_path):
        return f"Error: File not found at {file_path}"

    # Check extension
    _, ext = os.path.splitext(file_path)
    if ext.lower() == '.doc':
        return "Error: .doc format is not natively supported. Please save as .docx first."

    output = []
    try:
        doc = docx.Document(file_path)
        
        output.append("\n--- Document Content ---\n")
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                output.append(para.text)
        
        # Read tables
        if doc.tables:
            output.append("\n--- Tables ---\n")
            for i, table in enumerate(doc.tables):
                output.append(f"Table {i+1}:")
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    output.append(" | ".join(row_text))
                output.append("-" * 20)
        
        return "\n".join(output)

    except Exception as e:
        return f"Error reading Word file: {e}"
