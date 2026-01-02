from docx import Document as DocxDocument
from udf_tools.editor import UDFEditor
from udf_tools.models.styles import UDFStyle

class DocxConverter:
    @staticmethod
    def to_udf(docx_path: str, udf_path: str):
        docx = DocxDocument(docx_path)
        editor = UDFEditor()
        
        # Add a default style if not present
        default_style = UDFStyle(name="hvl-default", family="Times New Roman", size=12.0)
        editor.add_style(default_style)
        
        for para in docx.paragraphs:
            # Simple conversion for now: just text and alignment
            alignment = 0
            if para.alignment:
                # Mapping docx alignment to UDF (simplified)
                # docx: 0=LEFT, 1=CENTER, 2=RIGHT, 3=JUSTIFY
                alignment = para.alignment
            
            editor.add_paragraph(para.text, alignment=alignment)
            
        # TODO: Handle tables, runs (bold/italic), images
        
        editor.save(udf_path)

    @staticmethod
    def from_udf(udf_path: str, docx_path: str):
        # TODO: Implement UDF -> DOCX
        pass
