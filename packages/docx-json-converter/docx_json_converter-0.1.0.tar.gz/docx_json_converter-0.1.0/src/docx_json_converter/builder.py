import json
import logging
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

DEFAULT_FONT_SIZE_PT = 12.0

def get_alignment_enum(alignment_str: Optional[str]):
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return alignment_map.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)

def apply_run_formatting(run, formatting: Dict[str, Any]) -> None:
    if not formatting:
        return

    # Bool fields
    if "bold" in formatting:
        run.bold = formatting["bold"]
    if "italic" in formatting:
        run.italic = formatting["italic"]
    if "underline" in formatting:
        run.underline = formatting["underline"]
    
    # Size
    font_size = formatting.get("font_size")
    if font_size is not None:
        run.font.size = Pt(font_size)
    else:
        run.font.size = Pt(DEFAULT_FONT_SIZE_PT)

def set_document_columns(doc: Document, column_count: int):
    """
    Sets the number of columns for the first section of the document.
    Uses OXML manipulation.
    """
    if column_count < 2:
        return

    section = doc.sections[0]
    sectPr = section._sectPr
    
    # Check if 'w:cols' element exists
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols_element = cols[0]
    else:
        cols_element = OxmlElement('w:cols')
        sectPr.append(cols_element)
        
    # Set the 'num' attribute
    cols_element.set(qn('w:num'), str(column_count))
    
    # Optional: Set spacing (defaulting to 720 twips = 0.5 inch is common)
    cols_element.set(qn('w:space'), "720")

def convert_json_to_docx(json_data: Dict[str, Any], output_path: str) -> bool:
    """
    Creates a .docx file from structured JSON data.
    """
    try:
        doc = Document()
        
        # Apply Metadata (Columns)
        metadata = json_data.get("metadata", {})
        if "columns" in metadata:
            set_document_columns(doc, metadata["columns"])
        
        blocks = json_data.get("blocks", [])
        
        for block in blocks:
            # Handle List Items
            block_type = block.get("block_type", "paragraph")
            
            if block_type == "list_item":
                try:
                    paragraph = doc.add_paragraph(style='List Bullet')
                except:
                    paragraph = doc.add_paragraph()
            else:
                paragraph = doc.add_paragraph()

            # Alignment
            block_formatting = block.get("formatting", {})
            alignment_str = block_formatting.get("alignment", "left")
            paragraph.alignment = get_alignment_enum(alignment_str)
            
            runs = block.get("runs", [])
            
            if runs:
                for run_data in runs:
                    text = run_data.get("text", "")
                    run = paragraph.add_run(text)
                    run_formatting = run_data.get("formatting", {})
                    apply_run_formatting(run, run_formatting)
            else:
                text = block.get("text", "")
                run = paragraph.add_run(text)
                apply_run_formatting(run, block_formatting)
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        return False
