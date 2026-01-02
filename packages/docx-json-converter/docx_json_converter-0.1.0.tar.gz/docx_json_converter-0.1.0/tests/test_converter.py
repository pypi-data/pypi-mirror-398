import unittest
import os
from docx import Document
from docx.oxml.ns import qn
from docx_json_converter.parser import convert_docx_to_json
from docx_json_converter.builder import convert_json_to_docx

class TestDocxConverter(unittest.TestCase):
    def setUp(self):
        self.test_docx = "temp_test.docx"
        self.output_docx = "temp_output.docx"
        
        # Create a basic doc
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.save(self.test_docx)

    def tearDown(self):
        if os.path.exists(self.test_docx):
            os.remove(self.test_docx)
        if os.path.exists(self.output_docx):
            os.remove(self.output_docx)

    def test_round_trip_basic(self):
        json_data = convert_docx_to_json(self.test_docx)
        self.assertEqual(len(json_data['blocks']), 1)
        
        success = convert_json_to_docx(json_data, self.output_docx)
        self.assertTrue(success)
        
        doc = Document(self.output_docx)
        self.assertEqual(len(doc.paragraphs), 1)

    def test_no_formatting_flag(self):
        json_data = convert_docx_to_json(self.test_docx, include_formatting=False)
        self.assertEqual(len(json_data['blocks']), 1)
        self.assertEqual(json_data['blocks'][0]['formatting'], {})

    def test_column_support(self):
        # Create JSON with 2 columns
        json_data = {
            "metadata": {"columns": 2},
            "blocks": [{"text": "Column Text", "block_type": "paragraph"}]
        }
        
        # Build
        convert_json_to_docx(json_data, self.output_docx)
        
        # Verify (Need to parse back or check XML)
        # 1. Check via Parser
        new_json = convert_docx_to_json(self.output_docx)
        self.assertEqual(new_json['metadata']['columns'], 2)
        
        # 2. Check XML directly
        doc = Document(self.output_docx)
        sectPr = doc.sections[0]._sectPr
        cols = sectPr.xpath('./w:cols')
        self.assertTrue(len(cols) > 0)
        self.assertEqual(cols[0].get(qn('w:num')), '2')

if __name__ == '__main__':
    unittest.main()
