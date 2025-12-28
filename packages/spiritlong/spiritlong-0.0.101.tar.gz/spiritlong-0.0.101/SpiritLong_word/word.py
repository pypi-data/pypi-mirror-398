## 
# Copyright (c) 2023 Chongqing Spiritlong Technology Co., Ltd.
# All rights reserved.
# 
# @author	arthuryang
# @brief	word工具集，只支持docx。
##

import	docx		# pip install python-docx
import	re

## 打开docx文件
def open_docx(filename):
	try:
		doc	= docx.Document(filename)
		return doc
	except Exception as ex:
		print(str(ex))
		exit		

## 保存docx文件
def save_docx(doc, filename):
	try:
		doc.save(filename)
	except Exception as ex:
		print(f"保存失败：{filename}")
		print(str(ex))
		exit

## 替换docx中两个段落之间的内容
#	match_start	前一段落的特征字符串
#	match_end	后一段落的特征字符串
#	paragraphs	要写入的段落
def save_docx(doc, match_start, match_end, paragraphs=None):
	p_start_index	= -1
	p_end_index	= -1
	# for i in range(len(doc.paragraphs)):
	# 	if re.findall(match_start, doc.paragraphs[i].text):
	# 		p_start	= i
	# 	if re.findall(match_end, doc.paragraphs[i].text):
	# 		p_end	= i
	
	for p in doc.paragraphs:
		if re.findall(match_start, p.text):
			p_start	= p
		if re.findall(match_end, p.text) and p_start:
			# 确保p_start和p_end的先后顺序
			p_end	= p 

	if p_start and p_end:
		# 都找到了，删除
		flag	= False
		for p in doc.paragraphs:
			if flag and p._element==p_end._element:
				# 结束段落
				flag	= False
			elif p._element==p_start._element:
				# 开始段落
				flag	= True
			elif flag:
				# 该删除
				delete_paragraph(p)
		# 在开始段落之后添加
		add_paragraphs(p_start, paragraphs)

## 删除段落
def delete_paragraph(p):
	# element才是实质的段落数据对象，Paragraph只是指向它的对象封装
	element	= p._element
	element.getparent().remove(element)
	p._p = p._element = None

## 添加多个段落
#	p_start		在此段落之后插入多个段落
#	paragraphs	列表成员每个是一个(runs, style)元组，其中runs列表成员是一个字典
def add_paragraphs(p_start, paragraphs):
	last_p	= p_start
	for p in paragraphs:
		new_p = docx.oxml.xmlchemy.OxmlElement("w:p")
		last_p._p.addnext(new_p)
		# 注意用new_p创建新段落之后，last_p将指向该新段落
		last_p	= docx.text.paragraph.Paragraph(new_p, last_p._parent)
		
		if p:
				
			# 添加此段落的各部分
			for r in p['runs']:
				new_run	= last_p.add_run(r['text'])
				if 'font_name' in r:
					new_run.font.name	= r['font_name']
					new_run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), new_run.font.name)
				if 'font_size' in r:
					new_run.font.size	= r['font_size']
				if 'font_color_rgb' in r:
					new_run.font.color.rgb	= r['font_color_rgb']
				if 'bold' in r:
					new_run.font.bold	= r['bold']
				# 段落中的风格也适用于全体文字
				if 'font_name' in p:
					new_run.font.name	= p['font_name']
					new_run._element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), new_run.font.name)
				if 'font_size' in p:
					new_run.font.size	= p['font_size']
				if 'font_color_rgb' in p:
					new_run.font.color.rgb	= p['font_color_rgb']
				if 'bold' in p:
					new_run.font.bold	= p['bold']
			# 添加此段落风格
			if 'style' in p:
				last_p.style	= p['style']
			if 'alignment' in p:
				last_p.alignment				= p['alignment']
			if 'left_indent' in p:
				last_p.paragraph_format.left_indent		= p['left_indent']
			if 'right_indent' in p:
				last_p.paragraph_format.right_indent		= p['right_indent']
			if 'first_line_indent' in p:
				last_p.paragraph_format.first_line_indent	= p['first_line_indent']
			if 'space_before' in p:
				last_p.paragraph_format.space_before 		= p['space_before']
			if 'space_after' in p:
				last_p.paragraph_format.space_after  		= p['space_after']
			if 'line_spacing' in p:
				last_p.paragraph_format.line_spacing 		= p['line_spacing']
			if 'line_spacing_rule' in p:
				last_p.paragraph_format.line_spacing_rule 	= p['line_spacing_rule']
	return last_p

## 返回默认段落
def new_paragraph_data(runs=[], 
		alignment=docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT, 
		first_line_indent	= 0,
		line_spacing_rule	= docx.enum.text.WD_LINE_SPACING.SINGLE,
		line_spacing		= docx.shared.Pt(10)	):
	return {
		'runs'			: runs,
		'alignment'		: alignment,
		'left_indent'		: docx.shared.Cm(0),
		'right_indent'		: docx.shared.Cm(0),
		'first_line_indent'	: docx.shared.Pt(first_line_indent),
		'space_before'		: docx.shared.Pt(0),
		'space_after'		: docx.shared.Pt(0),
		'line_spacing'		: line_spacing,
		'line_spacing_rule'	: line_spacing_rule,			
	}

## 返回默认run
def new_run_data(text="", font_name="宋体", font_size=12, color=docx.shared.RGBColor(0, 0, 0), bold=False, italic=False):
	return {
		"text"			: text,
		"bold"			: bold,
		"italic"		: italic,	
		"font_name"		: font_name,
		"font_size"		: docx.shared.Pt(font_size),
		"font_color_rgb"	: color,
	}		

## 在段落之后添加若干空行（第一个空行需要手动添加）
def add_blank_paragraphs(p_start, n):
	p	= p_start

	for i in range(n):
		paragraph	= {
			'runs'	: [{'text':""}],
			'style'	: p_start.style,
		}
		p		= add_paragraphs(p, [paragraph])
	return p

## 添加指定样式的段落，支持换行符
#	p_start			在此段落之后添加
#	style			样式名称
#	style_attributes	样式字典直接指定样式
def add_text_with_style(p_start, text, style=None, style_attributes={}):
	if not style:
		# 使用p_start的样式
		style	= p_start.style
	style_attributes.update({'style':style})
		
	# 分析text中的换行符
	paragraphs	= [dict({'runs' : [{'text':t}]}, **style_attributes) for t in text.split('\n')]
	return add_paragraphs(p_start, paragraphs)

## 寻找指定的文字（相同样式）进行替换，这些指定文字用〖〗包含
def replace_text(doc, old_text, text):
	old_text_string	= f"〖{old_text}〗"
	for p in doc.paragraphs:
		# 在每个running中寻找
		if re.findall(old_text_string, p.text):
			# 在此段落
			start	= 0
			end	= 0
			state	= 0
			for i in range(len(p.runs)):
				if state==0:
					# 还没找到〖
					if p.runs[i].text=="〖":
						state	= 1
						start	= i
				elif state==1:
					# 已经找到〖，检查到是否匹配文本
					if old_text_string==''.join([r.text for r in p.runs[start:i+1]]):
						# 文本匹配
						p.runs[start].clear()
						p.runs[start+1].text	= str(text)
						#p.runs[start].style	= p.runs[start+1].style
						for j in range(start+2, i+1):
							p.runs[j].clear()
						return p
				elif p.runs[i+1]=="〗":
					# 重新来
					state	= 0

## 在指定的文本之后添加
def append_text(doc, match_text, text):
	p	= find_paragraph(doc, match_text)
	p.text	= p.text.replace(match_text, match_text+text)
	return p
		
## 定位指定文本所在的段落
def find_paragraph(doc, match_text):
	for p in doc.paragraphs:
		if re.findall(match_text, p.text):
			return p

## 返回文档中的样式名称列表
def get_styles(doc):
	styles = doc.styles
	return [s.name for s in styles if s.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH]
		
## 删除word表格中的行
#	table	表对象
#	line	行序号，从0开始
def table_remove_row(table, line):
	if line>=len(table.rows):
		return 
	row	= table.rows[line]
	row._element.getparent().remove(row._element)

## 在表格中插入行
#	table	表对象
#	row	行序号从0开始）
def table_insert_row(table, row):
	tbl		= table._tbl
	successor	= tbl.tr_lst[row]
	tr		= tbl._new_tr()
	for gridCol in tbl.tblGrid.gridCol_lst:
		tc		= tr.add_tc()
		tc.width	= gridCol.w
	successor.addprevious(tr)

	style	= {
		'sz'	: 8,		# 大概对应1.0磅的粗细
		'val'	: 'single',	# 线条
		'color'	: '#000000',
	}
	for cell in table.rows[row].cells:
		set_cell_border(cell, top=style, bottom=style, left=style, right=style)
		# 竖直居中
		cell.vertical_alignment	= docx.enum.table.WD_ALIGN_VERTICAL.CENTER
	return table.rows[row]

## 设置单元格边框
# 参数举例
#	top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
#	bottom={"sz": 12, "color": "#00FF00", "val": "single"},
#	left={"sz": 24, "val": "dashed", "shadow": "true"},
#	right={"sz": 12, "val": "dashed"},
def set_cell_border(cell, **kwargs):
	tc	= cell._tc
	tcPr	= tc.get_or_add_tcPr()

	# check for tag existnace, if none found, then create one
	tcBorders = tcPr.first_child_found_in("w:tcBorders")
	if tcBorders is None:
		tcBorders	= docx.oxml.OxmlElement('w:tcBorders')
		tcPr.append(tcBorders)

	# list over all available tags
	for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
		edge_data = kwargs.get(edge)
		if edge_data:
			tag = f'w:{edge}'
		else:
			continue

		# check for tag existnace, if none found, then create one
		element	= tcBorders.find(docx.oxml.ns.qn(tag))
		if element is None:
			element	= docx.oxml.OxmlElement(tag)
			tcBorders.append(element)

		# looks like order of attributes is important
		for key in ["sz", "val", "color", "space", "shadow"]:
			if key in edge_data:
				element.set(docx.oxml.ns.qn('w:{}'.format(key)), str(edge_data[key]))

## 设置word表格中的单元格内容，包括字体和大小
def table_cell_fill(table, row, column, content, style=None, style_attributes={}):
	# 如果超出范围，就创建
	for i in range(row+1-len(table.rows)):
		table.add_row()
	for i in range(column+1-len(table.columns)):
		table.add_column(docx.shared.Cm(1))

	cell	= table.cell(row, column)
	p_start	= cell.paragraphs[0]
	add_text_with_style(p_start, content, style, style_attributes)
	delete_paragraph(p_start)

## 在指定段落后创建表格
def create_table(doc, paragraph, rows, columns, style='Table Grid'):
	table	= doc.add_table(rows, columns, style)
	# 将表格附着到段落之后
	paragraph._p.addnext(table._tbl)
	return table

## 合并单元格
def merge_cell(table, row_min, row_max, column_min, column_max, text, style=None, style_attributes={}):
	table_cell_fill(table, row_min, column_min, text, style, style_attributes)
	table.cell(row_min, column_min).merge(table.cell(row_max, column_max))
	table.cell(row_min, column_min).vertical_alignment	= docx.enum.table.WD_ALIGN_VERTICAL.CENTER

## 调整列宽
def set_table_column_width_and_vertical_center(table, width_list):
	for row in table.rows:
		for i in range(len(width_list)):
			row.cells[i].width	= docx.shared.Cm(width_list[i])	
		for i in range(len(table.columns)):
			# 都纵向居中
			row.cells[i].vertical_alignment	= docx.enum.table.WD_ALIGN_VERTICAL.CENTER
			
# 调试/测试代码
if __name__ == '__main__':
	pass

