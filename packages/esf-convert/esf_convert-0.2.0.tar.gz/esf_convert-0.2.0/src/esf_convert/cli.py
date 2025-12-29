#!/usr/bin/env python3
# /// script
# requires-python = ">=3.10"
# dependencies = [
#     "lxml>=5.0.0",
#     "python-docx>=1.1.0",
# ]
# ///
"""
ESF (ECHO Single Form) XML to Markdown/DOCX Converter

Usage:
    esf-convert <input.xml> [-o output.docx] [-f md|docx]

Options:
    -o, --output    Specify output filename (default: input filename with .docx extension)
    -f, --format    Output format: docx (Word document, default) or md (markdown)
"""

import sys
import re
import html
import argparse
from pathlib import Path
from lxml import etree
from docx import Document


# Mappings from numeric values to human-readable labels for enumerated fields
# These are based on the ESF form dictionaries (dic="XXXX" attributes)
ENUM_LABELS = {
    # dic="2004" - Yes/No fields
    'yes_no': {
        '0': 'No',
        '1': 'Yes',
        '-1': '-',
    },
    # dic="2007" - Yes/No for targeting groups
    'targeting': {
        '0': 'No',
        '1': 'Yes',
        '-1': '-',
    },
    # dic="2023" - Marker question responses (GAA, ADA, PMNE, AP)
    'marker_question': {
        '-1': 'Not applicable',
        '0': 'No',
        '1': 'Partially',
        '2': 'Yes',
    },
    # dic="2025" - Screening marker question responses
    'screening_marker_question': {
        '-1': 'Not applicable',
        '0': 'No',
        '1': 'Partially',
        '2': 'Yes',
    },
    # dic="2052" - Overall marker score (PM_IM)
    'marker_score': {
        '-2': 'Not applicable',
        '-1': 'Not applicable',
        '0': '0 - Not targeted',
        '1': '1 - Partially integrated',
        '2': '2 - Fully integrated',
    },
    # dic="2053" - Screening overall marker score (SM_IM)
    'screening_marker_score': {
        '-2': 'Not applicable',
        '-1': 'Not applicable',
        '0': '0 - Not targeted',
        '1': '1 - Partially integrated',
        '2': '2 - Fully integrated',
    },
    # dic="2062" - Logistics support/challenge dropdown
    'log_support': {
        '-1': '-',
        '0': 'No',
        '1': 'Yes',
    },
    # dic="2013" - Outcome indicator type (SOITYP)
    'indicator_type': {
        '-1': '-',
        '0': 'Custom indicator',
    },
    # dic="2017" - Result indicator type (RITYP)
    'result_indicator_type': {
        '-1': '-',
        '0': 'Custom indicator',
    },
}

# Map field refs to their enum type
FIELD_ENUM_TYPES = {
    # Yes/No fields (dic="2004")
    'PM_APPLICABLE': 'yes_no',
    'DPIA': 'yes_no',
    'IPY': 'yes_no',
    'SBDY': 'yes_no',
    'STUY': 'yes_no',
    'VISY': 'yes_no',
    'IEVY': 'yes_no',
    'EEVY': 'yes_no',
    'EAUDY': 'yes_no',
    'AliUnY': 'yes_no',
    'AliFaY': 'yes_no',
    'AliRcY': 'yes_no',
    'AliOtY': 'yes_no',
    'AliNaY': 'yes_no',
    'VISA1Y_2021': 'yes_no',
    'VISA2Y_2021': 'yes_no',
    'VISA3Y_2021': 'yes_no',
    'VISA4Y_2021': 'yes_no',
    'VISA5Y_2021': 'yes_no',
    'VISB1Y_2021': 'yes_no',
    'VISB2Y_2021': 'yes_no',
    # Targeting (dic="2007")
    'TGY': 'targeting',
    # Marker questions (dic="2023")
    'PM_GAA': 'marker_question',
    'PM_ADA': 'marker_question',
    'PM_PMNE': 'marker_question',
    'PM_AP': 'marker_question',
    # Screening marker questions (dic="2025")
    'SM_GAA': 'screening_marker_question',
    'SM_ADA': 'screening_marker_question',
    'SM_PMNE': 'screening_marker_question',
    'SM_AP': 'screening_marker_question',
    # Marker overall scores (dic="2052")
    'PM_IM': 'marker_score',
    # Screening marker overall scores (dic="2053")
    'SM_IM': 'screening_marker_score',
    # Logistics support (dic="2062")
    'LOG_SUP': 'log_support',
    'LOG_CHAL': 'log_support',
    # Indicator types
    'SOITYP': 'indicator_type',
    'RITYP': 'result_indicator_type',
}


def convert_enum_value(ref: str, value: str) -> str:
    """Convert a numeric enum value to its human-readable label."""
    if not value or ref not in FIELD_ENUM_TYPES:
        return value

    enum_type = FIELD_ENUM_TYPES[ref]
    labels = ENUM_LABELS.get(enum_type, {})
    return labels.get(value.strip(), value)


def clean_html_content(content: str) -> str:
    if not content or content.strip() == "." or content.strip() == "-1":
        return ""
    content = html.unescape(content)
    content = content.replace("<![CDATA[", "").replace("]]>", "")
    content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'<p[^>]*>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'</p>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1\n', content, flags=re.IGNORECASE | re.DOTALL)
    content = re.sub(r'<ul[^>]*>|</ul>|<ol[^>]*>|</ol>', '\n', content, flags=re.IGNORECASE)
    content = re.sub(r'<span[^>]*>|</span>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'<[^>]+>', '', content)
    content = re.sub(r'\n{3,}', '\n\n', content)
    content = re.sub(r'[ \t]+', ' ', content)
    content = content.replace('\xa0', ' ').replace('Â', '').strip()
    return content


def extract_section_number(label_text: str) -> tuple[str, str]:
    match = re.match(r'^(\d+(?:\.\d+)*)\.?\s*(.*)', label_text)
    if match:
        return match.group(1), match.group(2).strip()
    return "", label_text.strip()


def get_section_level(section_num: str) -> int:
    if not section_num:
        return 3
    return min(1 + len(section_num.split('.')), 4)


def format_number(value: str) -> str:
    if not value or value == '.' or value == '-1':
        return '-'
    try:
        if ',' in value and '.' in value:
            value = value.replace('.', '').replace(',', '.')
        elif ',' in value:
            value = value.replace(',', '.')
        num = float(value)
        if num == int(num):
            return f"{int(num):,}"
        return f"{num:,.2f}"
    except ValueError:
        return value


def get_element_value(parent: etree._Element, ref: str) -> str:
    """Get text value of an element by ref attribute."""
    for elem in parent.iter('element'):
        if elem.get('ref') == ref:
            return clean_html_content(elem.text or "")
    return ""


def parse_nested_items(data_elem: etree._Element, field_mapping: dict) -> list[dict]:
    """Parse items from nested row/grp structures, skipping templates."""
    items = []
    if data_elem is None:
        return items
    
    for row in data_elem.findall('.//row'):
        uid = row.get('uid', '')
        if '{{' in uid:  # Skip template rows
            continue
        
        item = {}
        # Look in grp elements or directly in row
        for grp in row.findall('.//grp'):
            for elem in grp.findall('element'):
                ref = elem.get('ref', '')
                if ref in field_mapping:
                    item[field_mapping[ref]] = clean_html_content(elem.text or "")
        
        # Also check direct element children of row
        for elem in row.findall('element'):
            ref = elem.get('ref', '')
            if ref in field_mapping:
                item[field_mapping[ref]] = clean_html_content(elem.text or "")
        
        # Only add if we got meaningful data
        if any(v for v in item.values() if v):
            items.append(item)
    
    return items


def parse_outcome_indicators(chapter: etree._Element) -> list[dict]:
    """Parse outcome indicators (SOI) from chapter 7."""
    indicators = []
    
    # Find SOI element
    soi_elem = None
    for elem in chapter.findall('element'):
        if elem.get('ref') == 'SOI':
            soi_elem = elem
            break
    
    if soi_elem is None:
        return indicators
    
    field_mapping = {
        'SOISD': 'definition',
        'SOIS': 'source',
        'SOIVB': 'baseline',
        'SOIV': 'target',
        'SOIVIr': 'progress',
        'SOIVFr': 'achieved',
        'SOICo': 'comments',
        'SOIIr': 'interim_update',
        'SOIFr': 'final_report',
    }
    
    data_elem = soi_elem.find('data')
    items = parse_nested_items(data_elem, field_mapping)
    
    return [i for i in items if i.get('definition')]


def parse_result_indicators(ri_elem: etree._Element) -> list[dict]:
    """Parse result indicators from RI element."""
    field_mapping = {
        'RISD': 'definition',
        'RIS': 'source',
        'RISFr': 'source_final',
        'RIB': 'baseline',
        'RIT': 'target',
        'RITIr': 'progress',
        'RITFr': 'achieved',
        'RIPC': 'comments',
    }
    
    data_elem = ri_elem.find('data')
    items = parse_nested_items(data_elem, field_mapping)
    
    return [i for i in items if i.get('definition')]


def parse_activities(ra_elem: etree._Element) -> list[dict]:
    """Parse activities from RA element."""
    field_mapping = {
        'RASD': 'short_description',
        'RALD': 'detailed_description',
        'RAFr': 'final_report',
    }
    
    data_elem = ra_elem.find('data')
    items = parse_nested_items(data_elem, field_mapping)
    
    return [i for i in items if i.get('short_description')]


def parse_result_chapter(chapter: etree._Element) -> dict:
    """Parse a RESULT chapter (section 7.3)."""
    result = {
        'title': '',
        'sector': '',
        'subsector': '',
        'amount': '',
        'indicators': [],
        'activities': [],
        'interim_update': '',
        'final_report': '',
    }
    
    # Get result title from RT element (search all descendants)
    for elem in chapter.iter('element'):
        ref = elem.get('ref', '')
        if ref == 'RT' and elem.text:
            result['title'] = clean_html_content(elem.text)
        elif ref == 'RIr' and elem.text:  # Interim update on activities
            result['interim_update'] = clean_html_content(elem.text)
        elif ref == 'RFr' and elem.text:  # Final report
            result['final_report'] = clean_html_content(elem.text)
    
    # Find RESULT_TAB element (search all descendants)
    result_tab = None
    for elem in chapter.iter('element'):
        if elem.get('ref') == 'RESULT_TAB':
            result_tab = elem
            break
    
    if result_tab is not None:
        data_elem = result_tab.find('data')
        if data_elem is not None:
            for row in data_elem.findall('row'):
                uid = row.get('uid', '')
                if '{{' in uid:
                    continue
                
                # Look for sector info, indicators, and activities in each row
                for elem in row.iter('element'):
                    ref = elem.get('ref', '')
                    
                    if ref == 'RS' and elem.text:
                        result['sector'] = clean_html_content(elem.text)
                    elif ref == 'RSSN' and elem.text:
                        result['subsector'] = clean_html_content(elem.text)
                    elif ref == 'RAm' and elem.text:
                        result['amount'] = format_number(elem.text)
                    elif ref == 'RI':
                        result['indicators'] = parse_result_indicators(elem)
                    elif ref == 'RA':
                        result['activities'] = parse_activities(elem)
    
    return result


def build_ref_to_label_mapping(root: etree._Element) -> tuple[dict, dict]:
    """Build mappings from element refs to their label IDs and guideline IDs."""
    ref_to_label = {}
    ref_to_guideline = {}
    for ref_elem in root.iter('ref'):
        elem_ref = ref_elem.get('ref')
        if elem_ref:
            data_elem = ref_elem.find('data')
            if data_elem is not None:
                label_id = data_elem.get('label')
                if label_id:
                    ref_to_label[elem_ref] = label_id
                guideline_id = data_elem.get('guideline')
                if guideline_id:
                    ref_to_guideline[elem_ref] = guideline_id
    return ref_to_label, ref_to_guideline


def parse_esf_xml(xml_source: Path | bytes) -> dict:
    """Parse ESF XML from a file path or bytes.

    Args:
        xml_source: Either a Path to an XML file or XML content as bytes.

    Returns:
        Parsed data dictionary with metadata, labels, chapters, and results.
    """
    parser = etree.XMLParser(recover=True, encoding='utf-8')
    if isinstance(xml_source, bytes):
        root = etree.fromstring(xml_source, parser)
    else:
        tree = etree.parse(str(xml_source), parser)
        root = tree.getroot()
    
    # Build label lookup
    labels = {}
    for label_elem in root.iter('label'):
        ref = label_elem.get('ref')
        text = label_elem.text or ""
        text = text.replace("<![CDATA[", "").replace("]]>", "").strip()
        if ref:
            labels[ref] = text
    
    ref_to_label, ref_to_guideline = build_ref_to_label_mapping(root)
    
    # Extract metadata
    metadata = {}
    vars_elem = root.find('.//vars')
    if vars_elem is not None:
        for child in vars_elem:
            text = child.text or ""
            if text and text.strip():
                metadata[child.tag] = clean_html_content(text)
    
    project = root if root.tag == 'project' else root.find('.//project')
    if project is not None:
        if project.get('version'):
            metadata['xml_version'] = project.get('version')
        if project.get('export_date'):
            metadata['export_date'] = project.get('export_date')
    
    # Find all chapters
    chapters = []
    results = []
    
    for chapter in root.iter('chapter'):
        # Skip template chapters
        if chapter.get('template') == 'TRUE':
            continue
        
        chapter_label_ref = chapter.get('label', '')
        chapter_name = chapter.get('name', '')
        chapter_group = chapter.get('group', '')
        
        # Handle RESULT chapters specially (section 7.3)
        if chapter_group == 'RESULT':
            result = parse_result_chapter(chapter)
            if result['title']:
                results.append(result)
            continue
        
        chapter_title = labels.get(chapter_label_ref, f"Chapter {chapter_name}")
        
        elements = []
        outcome_indicators = []
        
        # Check if this is chapter 7 (has SOI elements)
        if chapter_name == 'chapter7':
            outcome_indicators = parse_outcome_indicators(chapter)
        
        for elem in chapter.findall('element'):
            elem_ref = elem.get('ref', '')
            elem_content = elem.text
            
            # Skip elements we handle specially
            if elem_ref in ('SOI', 'RESULT_TAB', 'RI', 'RA'):
                continue
            
            # Skip elements with nested data
            if elem.find('data') is not None:
                continue
            
            if elem_content and '{{' in elem_content:
                continue
            
            if not elem_content or elem_content.strip() in ('', '.', '-1'):
                continue
            
            label_id = ref_to_label.get(elem_ref, elem_ref)
            elem_label = labels.get(label_id, '')
            cleaned_content = clean_html_content(elem_content)

            # Convert enumerated field values to their labels
            cleaned_content = convert_enum_value(elem_ref, cleaned_content)

            # Get guideline/tip if available
            guideline_id = ref_to_guideline.get(elem_ref, '')
            guideline_text = clean_html_content(labels.get(guideline_id, '')) if guideline_id else ''

            if cleaned_content:
                elements.append({
                    'ref': elem_ref,
                    'label': elem_label,
                    'content': cleaned_content,
                    'guideline': guideline_text,
                })
        
        if chapter_title or elements or outcome_indicators:
            chapters.append({
                'label_ref': chapter_label_ref,
                'name': chapter_name,
                'title': chapter_title,
                'elements': elements,
                'outcome_indicators': outcome_indicators,
            })
    
    return {
        'metadata': metadata,
        'labels': labels,
        'ref_to_label': ref_to_label,
        'chapters': chapters,
        'results': results,
    }


def generate_yaml_frontmatter(metadata: dict) -> str:
    lines = ['---']
    field_mappings = [
        ('edoc_ref', 'reference_number'),
        ('agreement_num', 'agreement_number'),
        ('doctyp', 'document_type'),
        ('action_type', 'action_type'),
        ('partner_type', 'partner_type'),
        ('project_cr_dt', 'creation_date'),
        ('project_submit_dt', 'submission_date'),
        ('xml_version', 'xml_version'),
        ('export_date', 'export_date'),
        ('fpa', 'fpa_version'),
        ('id', 'project_id'),
    ]
    for internal_key, yaml_key in field_mappings:
        value = metadata.get(internal_key, '')
        if value:
            if any(c in str(value) for c in [':', '"', '\n', '#', '[', ']', '{', '}']):
                value = f'"{value}"'
            lines.append(f'{yaml_key}: {value}')
    lines.append('---')
    return '\n'.join(lines)


def should_skip_element(label: str) -> bool:
    if not label:
        return False
    skip_patterns = [
        'cannot be empty', 'please select', 'bytes limit',
        'click', 'please add', 'please specify',
        'incorrect format', 'comments on this chapter',
    ]
    return any(p in label.lower() for p in skip_patterns)


def generate_markdown(data: dict) -> str:
    lines = []
    metadata = data['metadata']
    
    lines.append(generate_yaml_frontmatter(metadata))
    lines.append("")
    
    # Get title
    title = None
    for chapter in data['chapters']:
        for elem in chapter['elements']:
            if elem['ref'] == 'AT' and elem['content']:
                title = elem['content']
                break
        if title:
            break
    
    lines.append(f"# {title}" if title else "# ECHO Single Form")
    lines.append("")
    
    # Process chapters
    for chapter in data['chapters']:
        chapter_title = clean_html_content(chapter['title'])
        section_num, title_text = extract_section_number(chapter_title)
        
        if section_num:
            lines.append(f"## {section_num}. {title_text}")
        elif title_text:
            lines.append(f"## {title_text}")
        else:
            continue
        lines.append("")
        
        # Process elements
        for element in chapter['elements']:
            elem_label = element['label']
            elem_content = element['content']
            elem_ref = element['ref']
            elem_guideline = element.get('guideline', '')

            if should_skip_element(elem_label) or elem_ref == 'AT':
                continue
            if elem_ref.startswith('heading') and not elem_content:
                continue

            elem_section, elem_title = extract_section_number(elem_label)

            if elem_section:
                level = get_section_level(elem_section)
                lines.append(f"{'#' * level} {elem_section}. {elem_title}")
                lines.append("")
            elif elem_title and not elem_title.startswith('[') and not elem_title.startswith('&'):
                lines.append(f"### {elem_title}")
                lines.append("")

            # Add guideline/tip in italics if available
            if elem_guideline:
                lines.append(f"*{elem_guideline}*")
                lines.append("")

            if elem_content:
                lines.append(elem_content)
                lines.append("")
        
        # Process outcome indicators (section 7.2)
        if chapter.get('outcome_indicators'):
            lines.append("### 7.2. Outcome Indicators")
            lines.append("")

            for i, ind in enumerate(chapter['outcome_indicators'], 1):
                lines.append(f"#### Outcome Indicator {i}")
                lines.append("")
                if ind.get('definition'):
                    lines.append(f"**Definition:** {ind['definition']}")
                    lines.append("")
                if ind.get('source'):
                    lines.append(f"**Source:** {ind['source']}")
                    lines.append("")
                if ind.get('baseline'):
                    lines.append(f"**Baseline:** {ind['baseline']}")
                if ind.get('target'):
                    lines.append(f"**Target:** {ind['target']}")
                if ind.get('progress'):
                    lines.append(f"**Progress:** {ind['progress']}")
                lines.append("")
                if ind.get('comments'):
                    lines.append(f"**Comments:** {ind['comments']}")
                    lines.append("")
                if ind.get('interim_update'):
                    lines.append(f"**Interim Update:** {ind['interim_update']}")
                    lines.append("")

        # Process results (section 7.3) - output after Chapter 7's content
        if chapter['name'] == 'chapter7' and data['results']:
            lines.append("### 7.3. Results")
            lines.append("")

            for i, result in enumerate(data['results'], 1):
                lines.append(f"#### Result {i}: {result['title']}")
                lines.append("")

                if result['subsector']:
                    lines.append(f"**Subsector:** {result['subsector']}")
                if result['amount']:
                    lines.append(f"**Estimated Amount:** €{result['amount']}")
                lines.append("")

                # Result indicators
                if result['indicators']:
                    lines.append("##### Result Indicators")
                    lines.append("")

                    for j, ind in enumerate(result['indicators'], 1):
                        lines.append(f"**Indicator {j}:** {ind.get('definition', '')}")
                        lines.append("")
                        if ind.get('source'):
                            lines.append(f"- **Source:** {ind['source']}")
                        if ind.get('baseline'):
                            lines.append(f"- **Baseline:** {ind['baseline']}")
                        if ind.get('target'):
                            lines.append(f"- **Target:** {ind['target']}")
                        if ind.get('progress'):
                            lines.append(f"- **Progress:** {ind['progress']}")
                        if ind.get('comments'):
                            lines.append(f"- **Comments:** {ind['comments']}")
                        lines.append("")

                # Activities
                if result['activities']:
                    lines.append("##### Activities")
                    lines.append("")

                    for j, act in enumerate(result['activities'], 1):
                        lines.append(f"**Activity {j}:** {act.get('short_description', '')}")
                        lines.append("")
                        if act.get('detailed_description'):
                            lines.append(act['detailed_description'])
                            lines.append("")

                # Interim update on activities
                if result['interim_update']:
                    lines.append("##### Interim Update on Activities")
                    lines.append("")
                    lines.append(result['interim_update'])
                    lines.append("")

            # Section 7.4 - Final Report on Results
            has_final_reports = any(r.get('final_report') for r in data['results'])
            if has_final_reports:
                lines.append("### 7.4. Final Report on Results")
                lines.append("")

                for i, result in enumerate(data['results'], 1):
                    if result.get('final_report'):
                        lines.append(f"#### Result {i}: {result['title']}")
                        lines.append("")
                        lines.append(result['final_report'])
                        lines.append("")

    return '\n'.join(lines)


def generate_docx(data: dict, output_path: Path | None = None) -> bytes | None:
    """Generate a Word document from parsed ESF data.

    Args:
        data: Parsed ESF data dictionary
        output_path: Path to save the document. If None, returns bytes instead.

    Returns:
        None if output_path is provided, otherwise the document as bytes.
    """
    from io import BytesIO
    doc = Document()
    metadata = data['metadata']

    # Get title
    title = None
    for chapter in data['chapters']:
        for elem in chapter['elements']:
            if elem['ref'] == 'AT' and elem['content']:
                title = elem['content']
                break
        if title:
            break

    # Add title
    doc.add_heading(title if title else "ECHO Single Form", 0)

    # Add metadata as a table
    if metadata:
        field_mappings = [
            ('edoc_ref', 'Reference Number'),
            ('agreement_num', 'Agreement Number'),
            ('doctyp', 'Document Type'),
            ('action_type', 'Action Type'),
            ('partner_type', 'Partner Type'),
            ('project_cr_dt', 'Creation Date'),
            ('project_submit_dt', 'Submission Date'),
            ('export_date', 'Export Date'),
        ]
        meta_items = [(label, metadata.get(key, '')) for key, label in field_mappings if metadata.get(key)]
        if meta_items:
            table = doc.add_table(rows=len(meta_items), cols=2)
            table.style = 'Table Grid'
            for i, (label, value) in enumerate(meta_items):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            doc.add_paragraph()

    # Process chapters
    for chapter in data['chapters']:
        chapter_title = clean_html_content(chapter['title'])
        section_num, title_text = extract_section_number(chapter_title)

        if section_num:
            doc.add_heading(f"{section_num}. {title_text}", 1)
        elif title_text:
            doc.add_heading(title_text, 1)
        else:
            continue

        # Process elements
        for element in chapter['elements']:
            elem_label = element['label']
            elem_content = element['content']
            elem_ref = element['ref']
            elem_guideline = element.get('guideline', '')

            if should_skip_element(elem_label) or elem_ref == 'AT':
                continue
            if elem_ref.startswith('heading') and not elem_content:
                continue

            elem_section, elem_title = extract_section_number(elem_label)

            if elem_section:
                level = min(get_section_level(elem_section), 4)
                doc.add_heading(f"{elem_section}. {elem_title}", level)
            elif elem_title and not elem_title.startswith('[') and not elem_title.startswith('&'):
                doc.add_heading(elem_title, 2)

            # Add guideline/tip in italics if available
            if elem_guideline:
                p = doc.add_paragraph()
                p.add_run(elem_guideline).italic = True

            if elem_content:
                doc.add_paragraph(elem_content)

        # Process outcome indicators (section 7.2)
        if chapter.get('outcome_indicators'):
            doc.add_heading("7.2. Outcome Indicators", 2)

            for i, ind in enumerate(chapter['outcome_indicators'], 1):
                doc.add_heading(f"Outcome Indicator {i}", 3)

                if ind.get('definition'):
                    p = doc.add_paragraph()
                    p.add_run("Definition: ").bold = True
                    p.add_run(ind['definition'])

                if ind.get('source'):
                    p = doc.add_paragraph()
                    p.add_run("Source: ").bold = True
                    p.add_run(ind['source'])

                if ind.get('baseline'):
                    p = doc.add_paragraph()
                    p.add_run("Baseline: ").bold = True
                    p.add_run(ind['baseline'])

                if ind.get('target'):
                    p = doc.add_paragraph()
                    p.add_run("Target: ").bold = True
                    p.add_run(ind['target'])

                if ind.get('progress'):
                    p = doc.add_paragraph()
                    p.add_run("Progress: ").bold = True
                    p.add_run(ind['progress'])

                if ind.get('comments'):
                    p = doc.add_paragraph()
                    p.add_run("Comments: ").bold = True
                    p.add_run(ind['comments'])

                if ind.get('interim_update'):
                    p = doc.add_paragraph()
                    p.add_run("Interim Update: ").bold = True
                    p.add_run(ind['interim_update'])

        # Process results (section 7.3) - output after Chapter 7's content
        if chapter['name'] == 'chapter7' and data['results']:
            doc.add_heading("7.3. Results", 2)

            for i, result in enumerate(data['results'], 1):
                doc.add_heading(f"Result {i}: {result['title']}", 3)

                if result['subsector']:
                    p = doc.add_paragraph()
                    p.add_run("Subsector: ").bold = True
                    p.add_run(result['subsector'])

                if result['amount']:
                    p = doc.add_paragraph()
                    p.add_run("Estimated Amount: ").bold = True
                    p.add_run(f"€{result['amount']}")

                # Result indicators
                if result['indicators']:
                    doc.add_heading("Result Indicators", 4)

                    for j, ind in enumerate(result['indicators'], 1):
                        p = doc.add_paragraph()
                        p.add_run(f"Indicator {j}: ").bold = True
                        p.add_run(ind.get('definition', ''))

                        if ind.get('source'):
                            doc.add_paragraph(f"Source: {ind['source']}", style='List Bullet')
                        if ind.get('baseline'):
                            doc.add_paragraph(f"Baseline: {ind['baseline']}", style='List Bullet')
                        if ind.get('target'):
                            doc.add_paragraph(f"Target: {ind['target']}", style='List Bullet')
                        if ind.get('progress'):
                            doc.add_paragraph(f"Progress: {ind['progress']}", style='List Bullet')
                        if ind.get('comments'):
                            doc.add_paragraph(f"Comments: {ind['comments']}", style='List Bullet')

                # Activities
                if result['activities']:
                    doc.add_heading("Activities", 4)

                    for j, act in enumerate(result['activities'], 1):
                        p = doc.add_paragraph()
                        p.add_run(f"Activity {j}: ").bold = True
                        p.add_run(act.get('short_description', ''))

                        if act.get('detailed_description'):
                            doc.add_paragraph(act['detailed_description'])

                # Interim update on activities
                if result['interim_update']:
                    doc.add_heading("Interim Update on Activities", 4)
                    doc.add_paragraph(result['interim_update'])

            # Section 7.4 - Final Report on Results
            has_final_reports = any(r.get('final_report') for r in data['results'])
            if has_final_reports:
                doc.add_heading("7.4. Final Report on Results", 2)

                for i, result in enumerate(data['results'], 1):
                    if result.get('final_report'):
                        doc.add_heading(f"Result {i}: {result['title']}", 3)
                        doc.add_paragraph(result['final_report'])

    if output_path:
        doc.save(str(output_path))
        return None
    else:
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()


def main():
    parser = argparse.ArgumentParser(
        description='Convert ESF (ECHO Single Form) XML to Markdown or DOCX',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
    esf-convert form.xml                    # Output as Word document (default)
    esf-convert form.xml -f md              # Output as Markdown
    esf-convert form.xml -o output.md       # Auto-detects format from extension
    esf-convert form.xml -f md -o report    # Explicit format with custom name
        '''
    )
    parser.add_argument('input', type=Path, help='Input ESF XML file')
    parser.add_argument('-o', '--output', type=Path, help='Output file path')
    parser.add_argument('-f', '--format', choices=['md', 'docx'],
                        help='Output format (default: docx, or inferred from output extension)')

    args = parser.parse_args()
    input_path = args.input

    if not input_path.exists():
        print(f"Error: Input file '{input_path}' not found.")
        sys.exit(1)

    # Determine output format
    output_format = args.format
    if not output_format and args.output:
        # Infer from output extension
        ext = args.output.suffix.lower()
        if ext == '.docx':
            output_format = 'docx'
        elif ext in ('.md', '.markdown'):
            output_format = 'md'
    if not output_format:
        output_format = 'docx'

    # Determine output path
    if args.output:
        output_path = args.output
        # Add extension if missing
        if not output_path.suffix:
            output_path = output_path.with_suffix('.docx' if output_format == 'docx' else '.md')
    else:
        output_path = input_path.with_suffix('.docx' if output_format == 'docx' else '.md')

    print(f"Parsing ESF XML: {input_path}")

    try:
        data = parse_esf_xml(input_path)

        if output_format == 'docx':
            generate_docx(data, output_path)
            print(f"Word document generated: {output_path}")
        else:
            markdown = generate_markdown(data)
            output_path.write_text(markdown, encoding='utf-8')
            print(f"Markdown file generated: {output_path}")

        total_elements = sum(len(ch.get('elements', [])) for ch in data['chapters'])
        total_outcome_indicators = sum(len(ch.get('outcome_indicators', [])) for ch in data['chapters'])
        total_results = len(data['results'])
        total_result_indicators = sum(len(r.get('indicators', [])) for r in data['results'])
        total_activities = sum(len(r.get('activities', [])) for r in data['results'])

        print(f"\nSummary:")
        print(f"  - Chapters: {len(data['chapters'])}")
        print(f"  - Content elements: {total_elements}")
        print(f"  - Outcome indicators (7.2): {total_outcome_indicators}")
        print(f"  - Results (7.3): {total_results}")
        print(f"  - Result indicators: {total_result_indicators}")
        print(f"  - Activities: {total_activities}")

    except etree.XMLSyntaxError as e:
        print(f"Error parsing XML: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}")
        raise


if __name__ == "__main__":
    main()
