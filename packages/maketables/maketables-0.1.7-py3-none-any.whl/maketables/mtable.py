import os
from typing import ClassVar

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from great_tables import GT
from IPython.display import display

from .symbols import translate_symbols

# Methods
# -	make: Just return output object (gt, docx, tex, html) or display directly in notebook or as tex when rendered to pdf in Quarto
# -	save: Save output object in new file to path (docx, tex, html) add parameter replace to replace existing file otherwise error message when file exists
# -	update: Update existing file with output object (so far only docx) at specified position
#
# Note:
# - both save and update have a parameter "show" to display the output object in the notebook as gt
##
# - Handling of paths:
#     - in save: if file_name is None, use combination of default_path and label as file_name to store the file
#     - in update: if file_name is relative path and default path is specified, use default_path to update the file_path


class MTable:
    r"""
    A table creation class supporting multiple output formats.

    MTable provides a unified interface for creating tables that can be output
    as HTML (via great-tables), LaTeX, or Word documents (DOCX). It supports
    features like multi-level column headers, row grouping, and extensive
    styling customization for each output format.

    Parameters
    ----------
    df : pd.DataFrame
        The data to display in the table.
    notes : str, optional
        Notes to display at the bottom of the table. Default is empty string.
    caption : str, optional
        Table caption. Default is None.
    tab_label : str, optional
        Label for the table (used in LaTeX \\label{}). Default is None.
    rgroup_sep : str, optional
        Row group separator style. "tb" for top+bottom lines, "t" for top only,
        "b" for bottom only, "" for no lines. Default is "tb".
    rgroup_display : bool, optional
        Whether to display row group names. Default is True.
    default_paths : str, dict, or None, optional
        Default paths for saving files. Can be a string (used for all types)
        or dict mapping output types to paths. Default is None.
    tex_style : dict, optional
        LaTeX style overrides for auto-display in notebooks. Can include any parameter
        from DEFAULT_TEX_STYLE such as: first_col_width, tab_width, texlocation,
        arraystretch, tabcolsep, data_align, x_col_align, cmidrule_trim,
        first_row_addlinespace, data_addlinespace, rgroup_addlinespace,
        group_header_format,.... See DEFAULT_TEX_STYLE for full options.
    typst_style : dict, optional
        Typst style overrides for auto-display in notebooks. Can include any parameter
        from DEFAULT_TYPST_STYLE such as: first_col_width, tab_width, data_align,
        group_header_format, notes_font_size. See DEFAULT_TYPST_STYLE for full options.
    docx_style : dict, optional
        DOCX style overrides for auto-display in notebooks. Can include any parameter
        from DEFAULT_DOCX_STYLE such as: first_col_width, font_name, font_color_rgb,
        font_size_pt, notes_font_size_pt, caption_font_name, caption_font_size_pt,
        caption_align, notes_align, align_center_cells, border_*_rule_sz,
        cell_margins_dxa, table_style_name. See DEFAULT_DOCX_STYLE for full options.
    gt_style : dict, optional
        GT/HTML style overrides for auto-display in notebooks. Can include any parameter
        from DEFAULT_GT_STYLE plus all GT tab_options() parameters such as: align,
        table_width, first_col_width, table_font_size_all, source_notes_font_size,
        data_row_padding, column_labels_padding, border styles/colors/widths, etc.
        See DEFAULT_GT_STYLE for full options.

    Examples
    --------
    Basic usage:

    >>> df = pd.DataFrame({"A": [1, 2], "B": ["x", "y"]})
    >>> table = MTable(df, caption="My Table")
    >>> table  # Auto-displays in notebook

    With style overrides for auto-display:

    >>> table = MTable(
    ...     df,
    ...     tex_style={"first_col_width": "3cm"},
    ...     typst_style={"first_col_width": "2.5cm"},
    ...     docx_style={"first_col_width": "2in"},
    ...     gt_style={"table_width": "100%"},
    ... )
    """

    # Class attributes for default values
    DEFAULT_NOTES = ""
    DEFAULT_CAPTION = None
    DEFAULT_TAB_LABEL = None
    DEFAULT_RGROUP_SEP = "tb"
    DEFAULT_RGROUP_DISPLAY = True
    DEFAULT_SAVE_PATH = None  # can be string or dict
    DEFAULT_REPLACE = False
    DEFAULT_SAVE_TYPE = "html"
    ADMISSIBLE_TYPES: ClassVar[list[str]] = ["gt", "tex", "typst", "docx", "html"]
    ADMISSIBLE_SAVE_TYPES: ClassVar[list[str]] = ["tex", "typst", "docx", "html"]

    # Default TeX style (override globally via MTable.DEFAULT_TEX_STYLE.update({...})
    # or per-call via tex_style in make/save/_output_tex)
    DEFAULT_TEX_STYLE: ClassVar[dict[str, object]] = {
        # Table dimensions
        "tab_width": r"\linewidth",  # Target width for tabularx (None for normal tabular)
        "first_col_width": None,  # LaTeX length for first column (None for flexible)
        # Table placement
        "texlocation": "htbp",  # LaTeX float placement specifier
        # Row height and column separation (scoped to the table)
        "arraystretch": 1,  # float or str
        "tabcolsep": "3pt",  # TeX length
        # Alignment
        "data_align": "c",  # l|c|r for non-tabularx data columns
        "x_col_align": "center",  # left|center|right for tabularx X columns
        # Rules/spacing
        "cmidrule_trim": "lr",  # "", "l", "r", "lr"
        "first_row_addlinespace": "1ex",  # spacing before first row of each row group; None disables
        "data_addlinespace": "0.5ex",  # spacing before and after data rows; None disables
        "rgroup_addlinespace": None,  # spacing between row groups (independent of rgroup_sep); None disables
        # Row-group header formatting
        "group_header_format": r"\emph{%s}",
        # Optional LaTeX injected right after \begingroup to scope table-level commands
        # Example: r"\\small" or r"\\fontsize{9pt}{11pt}\\selectfont"
        "group_intro": None,
        # Notes intro placed at start of notes minipage (e.g., font size or prefix text)
        # Examples: r"\\footnotesize" (default), r"\\small", r"\\textbf{Note: }\\footnotesize"
        "notes_intro": r"\footnotesize",
    }

    # Default Typst style (Typst table syntax)
    DEFAULT_TYPST_STYLE: ClassVar[dict[str, object]] = {
        # Table dimensions
        "tab_width": "100%",  # Target width for table (e.g., "100%", "15cm")
        "first_col_width": None,  # Column width for first column (e.g., "3cm")
        # Row height
        "row_height": "auto",  # Row height (e.g., "auto", "10pt")
        # Alignment
        "data_align": "center",  # left | center | right for data columns
        # Spacing
        "first_row_addlinespace": None,  # spacing before first row of each row group
        "data_addlinespace": None,  # spacing between data rows
        "rgroup_addlinespace": None,  # spacing between row groups
        # Row-group header formatting
        "group_header_format": "%s",  # plain formatting for group headers
        # Notes font size
        "notes_font_size": "10pt",
    }

    # Shared defaults (override per subclass if needed)
    DEFAULT_LABELS: ClassVar[dict[str, str]] = {}
    
    # Simple default DOCX styling. Users can tweak this globally or per instance.
    DEFAULT_DOCX_STYLE: ClassVar[dict[str, object]] = {
        "font_name": "Times New Roman",
        "font_color_rgb": (0, 0, 0),
        "font_size_pt": 11,  # body and header
        "notes_font_size_pt": 9,  # notes row
        # Caption-specific defaults
        "caption_font_name": "Times New Roman",
        "caption_font_size_pt": 11,
        "caption_align": "center",  # left|center|right|justify
        "notes_align": "justify",  # left|center|right|justify
        "align_center_cells": True,  # center all cells except first column
        # borders (Word size units; 4=thin, 8=thick)
        "border_top_rule_sz": 8,  # top rule above first header row
        "border_header_rule_sz": 4,  # bottom rule under last header row
        "border_bottom_rule_sz": 8,  # bottom rule under last data row
        "border_group_rule_sz": 4,  # lines above/below row group labels
        # table cell margins (dxa; 20 dxa = 1 pt)
        "cell_margins_dxa": {"left": 0, "right": 0, "top": 60, "bottom": 60},
        # optional table style name in Word (None => 'Table Grid')
        "table_style_name": None,
        # prevent page breaks within tables
        "prevent_page_breaks": True,
        # first column width (in inches, cm, pt, or None for auto)
        "first_col_width": None,  # e.g., "2.5in", "6cm", "180pt"
    }
    # Default GT styling - dictionary passed directly to GT.tab_options(**dict)
    # This allows users to specify ANY GT styling parameter, not just predefined ones
    # (override globally via MTable.DEFAULT_GT_STYLE.update({...}) or per instance via gt_style={...})
    DEFAULT_GT_STYLE: ClassVar[dict[str, object]] = {
        # Special parameters handled separately (not passed to tab_options)
        "align": "center",  # left | center | right (used for cols_align)
        "table_width": None,  # e.g., "100%" or None for auto-width
        "first_col_width": None,  # e.g., "150px", "3cm", "20%" or None for auto-width
        "table_font_size_all": "16px",  # Sets font size for all table elements except notes; individual font sizes override this
        "source_notes_font_size": "10px",
        # Standard tab_options parameters
        "data_row_padding": "2px",
        "column_labels_padding": "2px",
        # Caption spacing (GT expects single values for padding, not CSS shorthand)
        "heading_padding": "6px",  # Space between caption and table
        # Column label borders
        "column_labels_border_top_style": "solid",
        "column_labels_border_top_color": "black",
        "column_labels_border_top_width": "2px",  # Top rule (toprule) - thickest
        "column_labels_border_bottom_style": "solid",
        "column_labels_border_bottom_color": "black",
        "column_labels_border_bottom_width": "0.25px",  # Header midrule - medium
        "column_labels_vlines_color": "white",
        "column_labels_vlines_width": "0px",
        # Table body borders
        "table_body_border_top_style": "solid",
        "table_body_border_top_width": "0px",  # Avoid double top border by default
        "table_body_border_top_color": "black",
        "table_body_border_bottom_style": "solid",
        "table_body_border_bottom_width": "2px",  # Bottom rule (bottomrule) - thickest
        "table_body_border_bottom_color": "black",
        "table_body_hlines_style": "none",
        "table_body_vlines_color": "white",
        "table_body_vlines_width": "0px",
        # Row group borders
        "row_group_border_top_style": "solid",
        "row_group_border_top_width": "0.25px",
        "row_group_border_top_color": "black",
        "row_group_border_bottom_style": "solid",
        "row_group_border_bottom_width": "0.25px",
        "row_group_border_bottom_color": "black",
        "row_group_border_left_color": "white",
        "row_group_border_right_color": "white",
    }

    def __init__(
        self,
        df: pd.DataFrame,
        notes: str = DEFAULT_NOTES,
        caption: str | None = DEFAULT_CAPTION,
        tab_label: str | None = DEFAULT_TAB_LABEL,
        rgroup_sep: str = DEFAULT_RGROUP_SEP,
        rgroup_display: bool = DEFAULT_RGROUP_DISPLAY,
        default_paths: None | str | dict = DEFAULT_SAVE_PATH,
        # Style parameters for auto-display in notebooks (applied when no type is specified)
        tex_style: dict[str, object] | None = None,
        typst_style: dict[str, object] | None = None,
        docx_style: dict[str, object] | None = None,
        gt_style: dict[str, object] | None = None,
        # No other style/render defaults here; handled in output methods
        **kwargs,
    ):
        assert isinstance(df, pd.DataFrame), "df must be a pandas DataFrame."
        assert not isinstance(df.index, pd.MultiIndex) or df.index.nlevels <= 2, (
            "Row index can have at most two levels."
        )
        self.df = df
        self.notes = notes
        self.caption = caption
        self.tab_label = tab_label
        self.rgroup_sep = rgroup_sep
        self.rgroup_display = rgroup_display
        if isinstance(default_paths, str):
            self.default_paths = dict.fromkeys(
                self.ADMISSIBLE_SAVE_TYPES, default_paths
            )
        elif isinstance(default_paths, dict):
            self.default_paths = default_paths.copy()
        else:
            self.default_paths = {}

        # Store style parameters for auto-display
        # When displayed automatically (__repr__), styles are used directly
        # since the default display shows both LaTeX (for Quarto) and HTML (for notebooks)
        self._display_styles = {
            "tex_style": tex_style or {},
            "typst_style": typst_style or {},
            "docx_style": docx_style or {},
            "gt_style": gt_style or {},
        }

    def _translate_symbols(self, text: str, output_format: str) -> str:
        """
        Translate special symbols in text for the specified output format.

        Args:
            text: Text containing symbols to translate
            output_format: Target format ('tex', 'html', 'docx', 'gt', 'plain')

        Returns
        -------
            Text with symbols translated to the target format
        """
        return translate_symbols(text, output_format)

    def make(self, type: str | None = None, **kwargs):
        r"""
        Create the output object of the table (either gt, tex, docx, or html).
        If type is None, displays both HTML and LaTeX outputs for compatibility
        with both notebook viewing and Quarto rendering.

        Parameters
        ----------
        type : str, optional
            The type of the output object ("gt", "tex", "typst", "docx", "html").
        **kwargs : Further arguments forwarded to the respective output method when type is specified.
            - For type="tex" (LaTeX):

              - tex_style: Dict[str, object] (default MTable.DEFAULT_TEX_STYLE)
                Per-table overrides for TeX rendering, e.g.:
                {first_col_width: "3cm", tab_width: r"0.8\\textwidth", texlocation: "htbp",
                 arraystretch: 1.15, tabcolsep: "4pt", data_align: "c", x_col_align: "left",
                 cmidrule_trim: "lr", first_row_addlinespace: "0.75ex", data_addlinespace: "0.25ex",
                group_header_format: r"\\bfseries %s", notes_intro: r"\\footnotesize"}
              Note: When tab_width is set, ensure your document loads
              the tabularx and array packages.
            - For type="typst" (Typst):

              - typst_style: Dict[str, object] (default MTable.DEFAULT_TYPST_STYLE)
                Per-table overrides for Typst rendering, e.g.:
                {first_col_width: "2.5cm", tab_width: "100%", data_align: "center",
                 group_header_format: "%s", notes_font_size: "9pt"}
            - For type="gt" (HTML via great-tables):

              - gt_style: Dict[str, object]
                Dictionary passed directly to GT.tab_options() plus special parameters:
                'align' (for cols_align), 'table_width' (controls table width),
                'first_col_width' (width for first column, e.g., "150px", "3cm", "20%"),
                and 'table_font_size_all' (default font size for all elements except notes; individual settings override).
                Supports ALL GT styling options - see GT documentation for complete list.
                Common examples: table_font_size, column_labels_font_size, source_notes_font_size,
                data_row_padding, column_labels_padding, border styles/colors/widths, etc.
            - For type="docx" (Word):

              - docx_style: Dict[str, object]
                Overrides keys from MTable.DEFAULT_DOCX_STYLE such as:
                first_col_width, font_name, font_color_rgb, font_size_pt, notes_font_size_pt,
                caption_font_name, caption_font_size_pt, caption_align, notes_align,
                align_center_cells, border_*_rule_sz, cell_margins_dxa, table_style_name.

        Returns
        -------
        output : object or None
            - If type is specified: returns the backend output object.
            - If type is None: displays dual output in notebooks (HTML + LaTeX) and returns None.
        """
        if type is None:
            # If no type is specified, directly display dual output
            # Create dual output object for notebook/Quarto compatibility
            class DualOutput:
                """Display different outputs in notebook vs Quarto rendering."""

                def __init__(self, notebook_html, quarto_latex):
                    self.notebook_html = notebook_html
                    self.quarto_latex = quarto_latex

                def _repr_mimebundle_(self, include=None, exclude=None):
                    return {
                        "text/html": self.notebook_html,
                        "text/latex": self.quarto_latex,
                    }

            # Generate both HTML and LaTeX outputs
            html_output = self._output_gt().as_raw_html()
            tex_output = self._output_tex()

            # Add CSS to remove zebra striping if desired
            html_output = (
                """
            <style>
            table tr:nth-child(even) {
                background-color: transparent !important;
            }
            </style>
            """
                + html_output
            )
            # Create and display the dual output object
            dual_output = DualOutput(html_output, tex_output)
            display(dual_output)
            return None

        # For explicitly specified types
        assert type in self.ADMISSIBLE_TYPES, "types must be either " + ", ".join(
            self.ADMISSIBLE_TYPES
        )
        if type == "gt":
            return self._output_gt(**kwargs)
        elif type == "tex":
            return self._output_tex(**kwargs)
        elif type == "typst":
            return self._output_typst(**kwargs)
        elif type == "docx":
            return self._output_docx(**kwargs)
        else:
            return self._output_gt(**kwargs).as_raw_html()

    def save(
        self,
        type: str = DEFAULT_SAVE_TYPE,
        file_name: str | None = None,
        show: bool = True,
        replace: bool = DEFAULT_REPLACE,
        **kwargs,
    ):
        """
        Save the output object of the table to a file.

        Parameters
        ----------
        type : str, optional
            Output type to save ("tex", "docx", "html"). Default is 'html'.
        file_name : str, optional
            Path to save the file. If None, uses DEFAULT_SAVE_PATH[type] + tab_label.
        show : bool, optional
            If True, also returns the table as a GT object for display. Default True.
        replace : bool, optional
            If False and file exists, raises unless DEFAULT_REPLACE or replace=True.
        **kwargs : Arguments forwarded to the respective output method:
            - type="tex": tex_style (see make()).
            - type="docx": docx_style (see make()).
            - type="html": gt_style via _output_gt (e.g., gt_style with table_width).

        Returns
        -------
        output : GT
            When show=True, returns a GT object for display (HTML).
        """
        # No instance default injection; defaults resolved in output methods

        assert type in self.ADMISSIBLE_SAVE_TYPES, "types must be either " + ", ".join(
            self.ADMISSIBLE_SAVE_TYPES
        )
        if file_name is None:
            if self.tab_label is None:
                raise ValueError("tab_label must be provided if file_name is None")
            if self.default_paths.get(type) is None:
                raise ValueError(
                    f"Default path for type {type} has to be set if file_name is None"
                )
            # file name will be default path and tab_label:
            file_name = self.default_paths.get(type) + self.tab_label
        elif not os.path.splitext(file_name)[1]:
            # if file_name does not have an extension, add the extension
            file_name += f".{type}"
        if self.default_paths.get(type) is not None and not os.path.isabs(file_name):
            # if file_name is not an absolute path, and default path is set, then add default path to file_name
            file_name = os.path.join(self.default_paths.get(type, ""), file_name)
        if not replace and file_name is not None and os.path.exists(file_name):
            # when replace is False, check if file exists & abort if it does
            raise ValueError(
                f"File {file_name} already exists. Set replace=True or use class parameter DEFAULT_REPLACE=True to replace the file."
            )
        assert isinstance(file_name, str) and os.path.isdir(
            os.path.dirname(file_name)
        ), f"{file_name} is not a valid path."
        if type == "tex":
            with open(file_name, "w") as f:
                f.write(self._output_tex(**kwargs))  # Write the latex code to a file
        elif type == "typst":
            with open(file_name, "w") as f:
                f.write(self._output_typst(**kwargs))  # Write the typst code to a file
        elif type == "docx":
            document = self._output_docx(file_name=file_name, **kwargs)
            document.save(file_name)
        else:
            with open(file_name, "w") as f:
                f.write(self._output_gt(**kwargs).as_raw_html())
        if show:
            return self._output_gt(**kwargs)

    def update_docx(
        self,
        file_name: str | None = None,
        tab_num: int | None = None,
        show: bool = False,
        docx_style: dict[str, object] | None = None,
        **kwargs,
    ):
        """
        Update an existing DOCX file with the output object of the table.

        Parameters
        ----------
        file_name : str
            Path to the DOCX file. If relative and DEFAULT_SAVE_PATH['docx'] is set,
            that path is prepended. Must end with .docx (or no extension to auto-append).
        tab_num : int, optional
            1-based index of the table to replace. If None or out of range, appends a new table.
        show : bool, optional
            If True, also returns a GT object for display (HTML). Default False.
        docx_style : Dict[str, object], optional
            Per-call overrides for MTable.DEFAULT_DOCX_STYLE (see make()).
            Can include first_col_width for column width control.
        **kwargs : dict.

        Returns
        -------
        output : GT
            When show=True, returns a GT object for display (HTML).
        """
        assert file_name is not None, "file_name must be provided"
        # Resolve DOCX style (per-call -> class default)
        s = dict(self.DEFAULT_DOCX_STYLE)
        if docx_style:
            s.update(docx_style)

        # check if file_name is an absolute path, if not add default path
        if self.default_paths.get("docx") is not None and not os.path.isabs(file_name):
            file_name = os.path.join(self.default_paths.get("docx", ""), file_name)
        # check if file has no extension and if yes append docx extension
        if not os.path.splitext(file_name)[1]:
            file_name += ".docx"
        elif os.path.splitext(file_name)[1] != ".docx":
            raise ValueError("file_name must have .docx extension")
        assert isinstance(file_name, str) and os.path.isdir(
            os.path.dirname(file_name)
        ), f"{file_name} is not a valid path."
        # Check if the document exists
        if file_name and os.path.exists(file_name):
            document = Document(file_name)
            # Determine the number of tables in the document
            n_tables = len(document.tables)
        else:
            # if the document does not yet exist, create a new one
            document = Document()
            n_tables = 0

        # Check whether existing table should be replaced
        if n_tables > 0 and tab_num is not None and tab_num <= n_tables:
            # Replace the table at position tab_num
            table = document.tables[tab_num - 1]
            # Replace the caption before the table
            if self.caption is not None:
                # Find the paragraph before the table
                table_idx = list(document._body._body).index(table._element)
                if table_idx > 0:
                    prev_par_element = document._body._element[table_idx - 1]
                    if (
                        prev_par_element.tag.endswith("p")
                        and "Table" in prev_par_element.text
                    ):
                        # replace text in last subelement of prev_par_element (this should be the old caption)
                        prev_par_element[-1].text = f": {self.caption}"
            # Delete all rows in the old table
            for row in table.rows:
                table._element.remove(row._element)
            # Build the new table in the existing document
            self._build_docx_table(table, s)
        else:
            # Add a caption if specified
            if self.caption is not None:
                paragraph = document.add_paragraph("Table ", style="Caption")
                self._build_docx_caption(self.caption, paragraph, s)
            # Add a new table to the document
            table = document.add_table(rows=0, cols=self.df.shape[1] + 1)
            table.style = s.get("table_style_name") or "Table Grid"
            self._build_docx_table(table, s)
            # Add a line break after the table
            document.add_paragraph()

        # Save the document
        document.save(file_name)

        # return gt table if show is True
        if show:
            return self._output_gt(**kwargs)

    def update_typst(
        self,
        file_name: str | None = None,
        tab_label: str | None = None,
        show: bool = False,
        typst_style: dict[str, object] | None = None,
        **kwargs,
    ):
        """
        Update an existing Typst document by replacing a figure containing a table
        with the same tab_label. If no matching figure is found, append a new one.

        Parameters
        ----------
        file_name : str, optional
            Path to the Typst file. If relative and DEFAULT_SAVE_PATH['typst'] is set,
            that path is prepended. If omitted, defaults to <default_path>/<tab_label>.typ.
        tab_label : str, optional
            Label to search for. Defaults to self.tab_label.
        show : bool, optional
            If True, also returns a GT object for display (HTML). Default False.
        typst_style : Dict[str, object], optional
            Per-call overrides for MTable.DEFAULT_TYPST_STYLE (see make()).
        **kwargs : dict
            Forwarded to _output_typst (e.g., symbol translation overrides).
        """
        label = tab_label or self.tab_label
        assert label is not None, "tab_label must be provided"

        # Resolve file path
        if file_name is None:
            if self.default_paths.get("typst") is None:
                raise ValueError(
                    "Default path for typst has to be set if file_name is None"
                )
            file_name = os.path.join(self.default_paths.get("typst", ""), label)
        elif self.default_paths.get("typst") is not None and not os.path.isabs(
            file_name
        ):
            file_name = os.path.join(self.default_paths.get("typst", ""), file_name)

        root, ext = os.path.splitext(file_name)
        if not ext:
            file_name += ".typ"
        elif ext not in (".typ", ".typst"):
            raise ValueError("file_name must have .typ or .typst extension")

        assert isinstance(file_name, str) and os.path.isdir(
            os.path.dirname(file_name)
        ), f"{file_name} is not a valid path."

        # Render the new Typst figure (preserving original tab_label)
        original_label = self.tab_label
        self.tab_label = label
        try:
            new_figure = self._output_typst(typst_style=typst_style, **kwargs)
        finally:
            self.tab_label = original_label

        # Helper: find the matching closing parenthesis for the figure call
        def _find_matching_paren(text: str, open_idx: int) -> int:
            depth = 0
            for idx in range(open_idx, len(text)):
                ch = text[idx]
                if ch == "(":
                    depth += 1
                elif ch == ")":
                    depth -= 1
                    if depth == 0:
                        return idx
            return -1

        # Load existing content (if any)
        if os.path.exists(file_name):
            with open(file_name, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            content = ""

        marker = f"<{label}>"
        replaced = False

        if content and marker in content:
            label_pos = content.find(marker)
            figure_start = content.rfind("#figure", 0, label_pos)
            if figure_start != -1:
                open_paren = content.find("(", figure_start)
                if open_paren != -1:
                    close_paren = _find_matching_paren(content, open_paren)
                    label_end = content.find(">", label_pos)
                    if close_paren != -1 and label_end != -1:
                        label_end += 1
                        while label_end < len(content) and content[label_end] in "\r\n":
                            label_end += 1
                        figure_block = content[figure_start:label_end]
                        if "#table" in figure_block:
                            content = content[:figure_start] + new_figure + content[label_end:]
                            replaced = True

        if not replaced:
            if content and not content.endswith("\n"):
                content += "\n\n"
            elif content and not content.endswith("\n\n"):
                content += "\n"
            content = content + new_figure

        with open(file_name, "w", encoding="utf-8") as f:
            f.write(content)

        if show:
            return self._output_gt(**kwargs)

    def update_tex(
        self,
        file_name: str | None = None,
        tab_label: str | None = None,
        show: bool = False,
        tex_style: dict[str, object] | None = None,
        **kwargs,
    ):
        """
        Update an existing LaTeX document by replacing a table with the same tab_label.
        If no matching table is found, append a new one.

        Parameters
        ----------
        file_name : str, optional
            Path to the LaTeX file. If relative and DEFAULT_SAVE_PATH['tex'] is set,
            that path is prepended. If omitted, defaults to <default_path>/<tab_label>.tex.
        tab_label : str, optional
            Label to search for. Defaults to self.tab_label.
        show : bool, optional
            If True, also returns a GT object for display (HTML). Default False.
        tex_style : Dict[str, object], optional
            Per-call overrides for MTable.DEFAULT_TEX_STYLE (see make()).
        **kwargs : dict
            Forwarded to _output_tex (e.g., symbol translation overrides).
        """
        label = tab_label or self.tab_label
        assert label is not None, "tab_label must be provided"

        # Resolve file path
        if file_name is None:
            if self.default_paths.get("tex") is None:
                raise ValueError(
                    "Default path for tex has to be set if file_name is None"
                )
            file_name = os.path.join(self.default_paths.get("tex", ""), label)
        elif self.default_paths.get("tex") is not None and not os.path.isabs(
            file_name
        ):
            file_name = os.path.join(self.default_paths.get("tex", ""), file_name)

        root, ext = os.path.splitext(file_name)
        if not ext:
            file_name += ".tex"
        elif ext != ".tex":
            raise ValueError("file_name must have .tex extension")

        assert isinstance(file_name, str) and os.path.isdir(
            os.path.dirname(file_name)
        ), f"{file_name} is not a valid path."

        # Render the new LaTeX table (preserving original tab_label)
        original_label = self.tab_label
        self.tab_label = label
        try:
            new_table = self._output_tex(tex_style=tex_style, **kwargs)
        finally:
            self.tab_label = original_label

        # Load existing content (if any)
        if os.path.exists(file_name):
            with open(file_name, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            content = ""

        # Search for existing table with matching label
        label_marker = f"\\label{{{label}}}"
        replaced = False

        if content and label_marker in content:
            label_pos = content.find(label_marker)
            # Find the start of the table environment (search backwards for \begin{table})
            table_start = content.rfind("\\begin{table", 0, label_pos)
            if table_start != -1:
                # Find the end of the table environment
                table_end = content.find("\\end{table}", label_pos)
                if table_end != -1:
                    table_end = content.find("\n", table_end) + 1
                    # Replace the entire table environment
                    content = content[:table_start] + new_table + content[table_end:]
                    replaced = True

        if not replaced:
            # Insert before \end{document} if it exists, otherwise append
            end_doc_pos = content.rfind("\\end{document}")
            if end_doc_pos != -1:
                # Insert before \end{document}
                before = content[:end_doc_pos]
                after = content[end_doc_pos:]
                if not before.endswith("\n"):
                    before += "\n\n"
                elif not before.endswith("\n\n"):
                    before += "\n"
                content = before + new_table + "\n\n" + after
            else:
                # No \end{document} found, append to end
                if content and not content.endswith("\n"):
                    content += "\n\n"
                elif content and not content.endswith("\n\n"):
                    content += "\n"
                content = content + new_table

        with open(file_name, "w", encoding="utf-8") as f:
            f.write(content)

        if show:
            return self._output_gt(**kwargs)

    def _output_docx(self, docx_style: dict[str, object] | None = None, **kwargs):
        # Create a new Document
        document = Document()

        # Resolve DOCX style (per-call -> class default)
        s = dict(self.DEFAULT_DOCX_STYLE)
        if docx_style:
            s.update(docx_style)

        # Add caption if specified
        if self.caption is not None:
            paragraph = document.add_paragraph("Table ", style="Caption")
            # Apply symbol translation to caption
            translated_caption = self._translate_symbols(self.caption, "docx")
            self._build_docx_caption(translated_caption, paragraph, s)

        # Add table
        table = document.add_table(rows=0, cols=self.df.shape[1] + 1)
        table.style = s.get("table_style_name") or "Table Grid"
        self._build_docx_table(table, s)

        return document

    def _build_docx_caption(self, caption: str, paragraph, s: dict[str, object]):
        run = paragraph.add_run()
        r = run._r
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        r.append(fldChar)
        instrText = OxmlElement("w:instrText")
        instrText.text = r"SEQ Table \* ARABIC"
        r.append(instrText)
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        r.append(fldChar)
        bold_run = paragraph.add_run(f": {caption}")
        bold_run.bold = False
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        paragraph.alignment = align_map.get(
            str(s.get("caption_align", "center")).lower(), WD_ALIGN_PARAGRAPH.CENTER
        )
        # Font settings
        rgb = tuple(s.get("font_color_rgb", (0, 0, 0)))
        cap_font_name = str(
            s.get("caption_font_name", s.get("font_name", "Times New Roman"))
        )
        cap_font_size = Pt(int(s.get("caption_font_size_pt", 11)))
        for r_ in paragraph.runs:
            r_.font.name = cap_font_name
            r_.font.color.rgb = RGBColor(*rgb)
            r_.font.size = cap_font_size

        # Apply "Keep with next" to caption paragraph if page break prevention is enabled
        if s.get("prevent_page_breaks", True):
            pPr = paragraph._element.get_or_add_pPr()
            keepNext = OxmlElement("w:keepNext")
            pPr.append(keepNext)

    def _build_docx_table(self, table, s: dict[str, object]):
        # Make a copy of the DataFrame to avoid modifying the original
        dfs = self.df.copy()

        # Number of headline levels
        headline_levels = dfs.columns.nlevels
        # Are there row groups: is the case when dfs.index.nlevels > 1
        row_groups = dfs.index.nlevels > 1
        # Number of columns
        ncols = dfs.shape[1] + 1

        # Add column headers
        if isinstance(dfs.columns, pd.MultiIndex):
            # Add multiple headline rows for MultiIndex columns
            for level in range(headline_levels):
                hdr_cells = table.add_row().cells
                prev_col = None
                prev_cell_index = None
                for i, col in enumerate(dfs.columns.get_level_values(level)):
                    cell_index = i + 1
                    if col != prev_col:
                        hdr_cells[cell_index].text = self._translate_symbols(
                            str(col), "docx"
                        )
                        prev_col = col
                        prev_cell_index = cell_index
                    else:
                        # Only merge if prev_cell_index is not None and cell_index is valid
                        if prev_cell_index is not None and cell_index < len(hdr_cells):
                            hdr_cells[prev_cell_index].merge(hdr_cells[cell_index])
        else:
            hdr_cells = table.add_row().cells
            for i, col in enumerate(dfs.columns):
                hdr_cells[i + 1].text = self._translate_symbols(str(col), "docx")

        # Add row names and data
        row_group_rows = []
        if row_groups:
            current_group = None
            for idx, row in dfs.iterrows():
                if idx[0] != current_group:
                    # New row group
                    current_group = idx[0]
                    # append row number to row_group_rows
                    row_group_rows.append(len(table.rows))
                    if self.rgroup_display:
                        # Add a row for the group name
                        group_row_cells = table.add_row().cells
                        # add row group name
                        group_row_cells[0].text = self._translate_symbols(
                            str(current_group), "docx"
                        )
                        # make this cell slightly taller
                        for paragraph in group_row_cells[0].paragraphs:
                            paragraph.paragraph_format.space_after = Pt(3)
                            paragraph.paragraph_format.space_before = Pt(3)
                        for cell in group_row_cells[1:]:
                            cell.text = ""
                row_cells = table.add_row().cells
                row_cells[0].text = self._translate_symbols(str(idx[1]), "docx")
                for i, val in enumerate(row):
                    row_cells[i + 1].text = self._translate_symbols(str(val), "docx")
        else:
            for idx, row in dfs.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = self._translate_symbols(str(idx), "docx")
                for i, val in enumerate(row):
                    row_cells[i + 1].text = self._translate_symbols(str(val), "docx")

        # Set first column width if specified
        if s.get("first_col_width") is not None:
            first_col_width_str = str(s["first_col_width"]).strip().lower()
            try:
                # Parse width specification
                if first_col_width_str.endswith("in"):
                    width_val = float(first_col_width_str[:-2])
                    width = Inches(width_val)
                elif first_col_width_str.endswith("cm"):
                    width_val = float(first_col_width_str[:-2])
                    width = Cm(width_val)
                elif first_col_width_str.endswith("pt"):
                    width_val = float(first_col_width_str[:-2])
                    width = Pt(width_val)
                else:
                    # Try to parse as points if no unit specified
                    width_val = float(first_col_width_str)
                    width = Pt(width_val)

                # Set the width for the first column in all rows
                for row in table.rows:
                    if len(row.cells) > 0:
                        row.cells[0].width = width
            except (ValueError, IndexError):
                # If parsing fails, ignore the width setting
                pass

        # Center all columns except the first one
        if s.get("align_center_cells", True):
            for row in table.rows:
                for cell in row.cells[1:]:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add notes (Note: we alsways add notes, even if empty)
        # Add row to the table that consists only of one cell with the notes
        notes_row = table.add_row().cells
        notes_row[0].text = self._translate_symbols(self.notes, "docx")
        # Merge the cell with the notes
        table.cell(-1, 0).merge(table.cell(-1, ncols - 1))
        # Set alignment and font size for the notes
        for paragraph in notes_row[0].paragraphs:
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            paragraph.alignment = align_map.get(
                str(s.get("notes_align", "justify")).lower(), WD_ALIGN_PARAGRAPH.JUSTIFY
            )
            for run in paragraph.runs:
                run.font.name = str(s.get("font_name", "Times New Roman"))
                run.font.size = Pt(int(s.get("notes_font_size_pt", 9)))
                rgb = tuple(s.get("font_color_rgb", (0, 0, 0)))
                run.font.color.rgb = RGBColor(*rgb)

        # Apply font to all table cells
        rgb_all = tuple(s.get("font_color_rgb", (0, 0, 0)))
        base_size = Pt(int(s.get("font_size_pt", 11)))
        notes_size = Pt(int(s.get("notes_font_size_pt", 9)))
        for ridx, row in enumerate(table.rows):
            is_notes_row = ridx == len(table.rows) - 1
            size = notes_size if is_notes_row else base_size
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = str(s.get("font_name", "Times New Roman"))
                        run.font.color.rgb = RGBColor(*rgb_all)
                        run.font.size = size

        # First hide all borders
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                top_border = OxmlElement("w:top")
                top_border.set(qn("w:val"), "nil")
                borders.append(top_border)
                bottom_border = OxmlElement("w:bottom")
                bottom_border.set(qn("w:val"), "nil")
                borders.append(bottom_border)
                left_border = OxmlElement("w:left")
                left_border.set(qn("w:val"), "nil")
                borders.append(left_border)
                right_border = OxmlElement("w:right")
                right_border.set(qn("w:val"), "nil")
                borders.append(right_border)
                tcPr.append(borders)

        # Add a thicker line above the top row
        for cell in table.rows[0].cells:
            tcPr = cell._element.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            top_border = OxmlElement("w:top")
            top_border.set(qn("w:val"), "single")
            top_border.set(qn("w:sz"), str(int(s.get("border_top_rule_sz", 8))))
            borders.append(top_border)
            tcPr.append(borders)

        # Add a line to the last headline row
        for cell in table.rows[headline_levels - 1].cells:
            tcPr = cell._element.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            bottom_border = OxmlElement("w:bottom")
            bottom_border.set(qn("w:val"), "single")
            bottom_border.set(qn("w:sz"), str(int(s.get("border_header_rule_sz", 4))))
            borders.append(bottom_border)
            tcPr.append(borders)

        # If the column index has more than one level, add a line above the last headline
        # row that spans all but the first cell
        if headline_levels > 1:
            for cell in table.rows[headline_levels - 1].cells[1:]:
                tcPr = cell._element.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                top_border = OxmlElement("w:top")
                top_border.set(qn("w:val"), "single")
                top_border.set(qn("w:sz"), str(int(s.get("border_header_rule_sz", 4))))
                borders.append(top_border)
                tcPr.append(borders)

        # Row group lines
        for row in row_group_rows:
            if "t" in self.rgroup_sep:
                for cell in table.rows[row].cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    borders = OxmlElement("w:tcBorders")
                    top_border = OxmlElement("w:top")
                    top_border.set(qn("w:val"), "single")
                    top_border.set(
                        qn("w:sz"), str(int(s.get("border_group_rule_sz", 4)))
                    )
                    borders.append(top_border)
                    tcPr.append(borders)
            if self.rgroup_display and "b" in self.rgroup_sep:
                for cell in table.rows[row].cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    borders = OxmlElement("w:tcBorders")
                    bottom_border = OxmlElement("w:bottom")
                    bottom_border.set(qn("w:val"), "single")
                    bottom_border.set(
                        qn("w:sz"), str(int(s.get("border_group_rule_sz", 4)))
                    )
                    borders.append(bottom_border)
                    tcPr.append(borders)

        # Add a thicker line below the last row
        for cell in table.rows[-2].cells:
            tcPr = cell._element.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            bottom_border = OxmlElement("w:bottom")
            bottom_border.set(qn("w:val"), "single")
            bottom_border.set(qn("w:sz"), str(int(s.get("border_bottom_rule_sz", 8))))
            borders.append(bottom_border)
            tcPr.append(borders)

        # Adapt cell margins
        tc = table._element
        tblPr = tc.tblPr
        tblCellMar = OxmlElement("w:tblCellMar")
        _margins = s.get(
            "cell_margins_dxa", {"left": 0, "right": 0, "top": 60, "bottom": 60}
        )
        for m in ("left", "right", "top", "bottom"):
            node = OxmlElement(f"w:{m}")
            node.set(qn("w:w"), str(int(_margins.get(m, 0))))
            node.set(qn("w:type"), "dxa")
            tblCellMar.append(node)
        tblPr.append(tblCellMar)

        # Prevent page breaks within table (keep table together)
        if s.get("prevent_page_breaks", True):
            # Apply "Keep lines together" and "Keep with next" to all paragraphs in the table
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # Get paragraph properties
                        pPr = paragraph._element.get_or_add_pPr()

                        # Add "Keep lines together" property
                        keepLines = OxmlElement("w:keepLines")
                        pPr.append(keepLines)

                        # Add "Keep with next" property
                        keepNext = OxmlElement("w:keepNext")
                        pPr.append(keepNext)

    def _output_tex(
        self,
        tex_style: dict[str, object] | None = None,
        **kwargs,
    ):
        # Make a copy of the DataFrame to avoid modifying the original
        dfs = self.df.copy()

        # Resolve TeX style (per-call -> class default)
        s = dict(getattr(self, "DEFAULT_TEX_STYLE", {}))
        if tex_style:
            s.update(tex_style)

        # Resolve TeX defaults from style dictionary
        _fcw = s.get("first_col_width")

        # Normalize tab_width (only these two keywords are mapped)
        def _normalize_width(w: str | None) -> str | None:
            if w is None:
                return None
            v = str(w).strip()
            low = v.lower()
            if low == "linewidth":
                return r"\linewidth"
            if low == "textwidth":
                return r"\textwidth"
            return v  # e.g., "4cm", r"0.5\textwidth"

        _tw = _normalize_width(s.get("tab_width"))
        use_tabularx = _tw is not None

        # Replace newlines and wrap cells with makecell if needed
        def _prep_cell(x):
            if isinstance(x, float) and np.isnan(x):
                return ""
            if isinstance(x, str):
                x = x.replace("\n", r"\\")
                if r"\\" in x:
                    return f"\\makecell{{{x}}}"
                return x
            return str(x)

        # Element-wise conversion; prefer DataFrame.map (pandas >= 2.1), fallback to applymap
        dfs = dfs.map(_prep_cell) if hasattr(dfs, "map") else dfs.applymap(_prep_cell)  # type: ignore[attr-defined]

        # Determine row groups (if MultiIndex on rows)
        row_levels = dfs.index.nlevels
        row_groups_present = row_levels > 1
        if row_groups_present:
            top_row_id = dfs.index.get_level_values(0).to_list()
            row_groups = list(dict.fromkeys(top_row_id))
            row_groups_len = [top_row_id.count(group) for group in row_groups]
            # Show only the inner row labels in the stub
            dfs.index = dfs.index.droplevel(0)

        stub_cols = dfs.index.nlevels
        data_cols = dfs.shape[1]

        # Column spec
        if use_tabularx:
            align_map = {
                "left": r">{\raggedright\arraybackslash}X",
                "center": r">{\centering\arraybackslash}X",
                "right": r">{\raggedleft\arraybackslash}X",
            }
            x_align = align_map.get(
                str(s.get("x_col_align", "center")).lower(), align_map["center"]
            )
            # Default: first column auto-sizes to content (uses left-aligned), data columns share space (X)
            if _fcw:
                first_spec = f"p{{{_fcw}}}"
            else:
                first_spec = r">{\raggedright\arraybackslash}l"  # Auto-size to content
            # Data columns and any additional stub columns use flexible space
            n_flex = max(0, stub_cols - 1) + data_cols
            rest_spec = x_align * n_flex
            # Start flush left by removing left padding with @{}
            colspec = "@{}" + first_spec + rest_spec
        else:
            # Default: first column auto-sizes to content (l), other stubs left-aligned (l), data aligned per setting
            first_stub = f"p{{{_fcw}}}" if _fcw else "l"
            other_stubs = "l" * max(0, stub_cols - 1)
            data_align = str(s.get("data_align", "c")).lower()
            if data_align not in {"l", "c", "r"}:
                data_align = "c"
            data_spec = data_align * data_cols
            # Start flush left by removing left padding with @{}
            colspec = "@{}" + first_stub + other_stubs + data_spec

        # Build header rows (MultiIndex columns -> spanners + cmidrules)
        header_lines = []

        def _make_spanner_row(level_labels):
            # level_labels is a list (len = n data columns) of labels at a given level
            cells = []
            spans = []
            i = 0
            n = len(level_labels)
            while i < n:
                label = level_labels[i]
                span = 1
                j = i + 1
                while j < n and level_labels[j] == label:
                    span += 1
                    j += 1
                cells.append(str(label))
                spans.append(span)
                i = j
            return cells, spans

        # Prepare column labels by levels
        if isinstance(dfs.columns, pd.MultiIndex):
            col_levels = dfs.columns.nlevels
            # For each level except the last, create a spanner row
            for lvl in range(col_levels - 1):
                labels_lvl = [dfs.columns[i][lvl] for i in range(len(dfs.columns))]
                row_cells, row_spans = _make_spanner_row(labels_lvl)
                # Build the LaTeX row: prepend stub blanks, then multicolumns
                parts = [""] * stub_cols
                cmid_ranges = []
                left = stub_cols + 1
                for cell, span in zip(row_cells, row_spans, strict=False):
                    parts.append(f"\\multicolumn{{{span}}}{{c}}{{{cell}}}")
                    cmid_ranges.append((left, left + span - 1))
                    left += span
                # Add the spanner row
                header_lines.append(" & ".join(parts) + r" \\")
                # Place the cmidrule(s) directly under the spanner row
                if cmid_ranges:
                    trim = str(s.get("cmidrule_trim", "lr"))
                    opt = f"({trim})" if trim else ""
                    cmids = " ".join(
                        [rf"\cmidrule{opt}{{{L}-{R}}}" for (L, R) in cmid_ranges]
                    )
                    header_lines.append(cmids)
            # Last level: the actual column names
            last_labels = [dfs.columns[i][-1] for i in range(len(dfs.columns))]
            last_parts = [""] * stub_cols + [str(x) for x in last_labels]
            header_lines.append(" & ".join(last_parts) + r" \\")
        else:
            # Single-level columns: one header row with the column names
            last_parts = [""] * stub_cols + [str(c) for c in dfs.columns]
            header_lines.append(" & ".join(last_parts) + r" \\")

        # Build body rows
        body_lines = []

        if row_groups_present:
            start = 0
            for gi, (gname, glen) in enumerate(
                zip(row_groups, row_groups_len, strict=False)
            ):
                if self.rgroup_display:
                    fmt = str(s.get("group_header_format", r"\emph{%s}"))
                    body_lines.append((fmt % str(gname)) + r" \\")
                    # Only add space after group header if data_addlinespace is set
                    if s.get("data_addlinespace") is not None:
                        body_lines.append(rf"\addlinespace[{s['data_addlinespace']}]")
                # Add first_row_addlinespace before the first data row of this group
                if s.get("first_row_addlinespace") is not None:
                    body_lines.append(rf"\addlinespace[{s['first_row_addlinespace']}]")
                # Rows for this group
                end = start + glen
                for ridx in range(start, end):
                    # Add space before data row (except for the very first row of the group)
                    if s.get("data_addlinespace") is not None and ridx > start:
                        body_lines.append(rf"\addlinespace[{s['data_addlinespace']}]")

                    row_label = str(dfs.index[ridx])
                    vals = [dfs.iloc[ridx, j] for j in range(data_cols)]
                    row_parts = [row_label] + [str(v) for v in vals]
                    body_lines.append(" & ".join(row_parts) + r" \\")

                    # Add space after data row
                    if s.get("data_addlinespace") is not None:
                        body_lines.append(rf"\addlinespace[{s['data_addlinespace']}]")
                # Group separator - always add spacing if set, optionally add midrule
                if gi < len(row_groups) - 1:
                    rg_space = s.get("rgroup_addlinespace")
                    if rg_space is not None:
                        body_lines.append(rf"\addlinespace[{rg_space}]")

                    # Only add midrule if rgroup_sep is not empty
                    if self.rgroup_sep != "":
                        body_lines.append(r"\midrule")
                        if rg_space is not None:
                            body_lines.append(rf"\addlinespace[{rg_space}]")
                start = end
        else:
            for ridx in range(dfs.shape[0]):
                # Add first_row_addlinespace before the very first row
                if s.get("first_row_addlinespace") is not None and ridx == 0:
                    body_lines.append(rf"\addlinespace[{s['first_row_addlinespace']}]")
                # Add space before data row (except for the very first row)
                elif s.get("data_addlinespace") is not None and ridx > 0:
                    body_lines.append(rf"\addlinespace[{s['data_addlinespace']}]")

                row_label = str(dfs.index[ridx])
                vals = [dfs.iloc[ridx, j] for j in range(data_cols)]
                row_parts = [row_label] + [str(v) for v in vals]
                body_lines.append(" & ".join(row_parts) + r" \\")

                # Add space after data row
                if s.get("data_addlinespace") is not None:
                    body_lines.append(rf"\addlinespace[{s['data_addlinespace']}]")
            # Spacing now controlled symmetrically by data_addlinespace setting

        # Assemble tabular/tabularx content
        tab_env = "tabularx" if use_tabularx else "tabular"
        width_arg = f"{{{_tw}}}" if use_tabularx else ""
        lines = []
        # Scope style changes to this table
        lines.append(r"\begingroup")
        # Inject optional TeX commands scoped to this table
        if s.get("group_intro"):
            lines.append(str(s["group_intro"]))
        # Top-align makecell content
        lines.append(r"\renewcommand\cellalign{t}")
        if s.get("arraystretch") is not None:
            lines.append(rf"\renewcommand\arraystretch{{{s['arraystretch']}}}")
        if s.get("tabcolsep"):
            lines.append(rf"\setlength{{\tabcolsep}}{{{s['tabcolsep']}}}")
        lines.append(rf"\begin{{{tab_env}}}{width_arg}{{{colspec}}}")
        lines.append(r"\toprule")
        lines.extend(header_lines)
        lines.append(r"\midrule")
        lines.extend(body_lines)
        lines.append(r"\bottomrule")
        lines.append(rf"\end{{{tab_env}}}")
        lines.append(r"\endgroup")

        latex_res = "\n".join(lines)

        # Wrap threeparttable
        if self.notes is not None:
            notes_intro = str(s.get("notes_intro", r"\footnotesize"))
            latex_res = (
                "\\begin{threeparttable}\n"
                + latex_res
                + "\n\\noindent\\begin{minipage}{\\linewidth}\\smallskip" + notes_intro + "\n"
                + self.notes
                + "\\end{minipage}\n"
                + "\n\\end{threeparttable}"
            )
        else:
            latex_res = (
                "\\begin{threeparttable}\n" + latex_res + "\n\\end{threeparttable}"
            )

        # Optional table float wrapper
        if (self.caption is not None) or (self.tab_label is not None):
            latex_res = (
                "\\begin{table}["
                + str(s.get("texlocation", "htbp"))
                + "]\n"
                + "\\centering\n"
                + (
                    "\\caption{" + self.caption + "}\n"
                    if self.caption is not None
                    else ""
                )
                + (
                    "\\label{" + self.tab_label + "}\n"
                    if self.tab_label is not None
                    else ""
                )
                + "\\smallskip\n"  # Add space between caption and table
                + latex_res
                + "\n\\end{table}"
            )

        # Apply symbol translation to the final LaTeX output
        latex_res = self._translate_symbols(latex_res, "tex")

        return latex_res

    def _escape_typst(self, text: str, escape_asterisks: bool = True) -> str:
        """
        Escape special characters for Typst syntax.
        
        Parameters
        ----------
        text : str
            Text to escape
        escape_asterisks : bool, default True
            Whether to escape asterisks (*, **, etc.). Set to False for header cells
            that intentionally use *text* for bold formatting.
        
        Returns
        -------
        str
            Escaped text safe for Typst
        """
        text = str(text)
        
        # Temporarily preserve actual newlines by replacing them with a placeholder
        newline_placeholder = "\x00NEWLINE\x00"
        text = text.replace("\n", newline_placeholder)
        
        # Escape backslashes (must be first!)
        text = text.replace("\\", "\\\\")
        # Escape bracket characters that interfere with Typst table syntax
        text = text.replace("[", "\\[")
        text = text.replace("]", "\\]")
        # Escape other special characters
        text = text.replace("#", "\\#")
        if escape_asterisks:
            # Only escape asterisks that are NOT already preceded by backslash
            # (to preserve intentional escapes like \* from user input)
            import re
            text = re.sub(r'(?<!\\)\*', r'\\*', text)
        text = text.replace("<", "\\<")
        text = text.replace(">", "\\>")
        text = text.replace("_", "\\_")
        
        # Restore actual newlines
        text = text.replace(newline_placeholder, "\n")
        
        return text

    def _output_typst(
        self,
        typst_style: dict[str, object] | None = None,
        **kwargs,
    ):
        """
        Generate Typst table code.
        
        Tables are created using Typst's native table syntax with configurable styling.
        """
        # Make a copy of the DataFrame to avoid modifying the original
        dfs = self.df.copy()

        # Resolve Typst style (per-call -> class default)
        s = dict(getattr(self, "DEFAULT_TYPST_STYLE", {}))
        if typst_style:
            s.update(typst_style)
        notes_fontsize = s.get("notes_font_size", "10pt")

        # Prepare cell content (escape special Typst characters)
        def _prep_cell(x):
            if isinstance(x, float) and np.isnan(x):
                return ""
            # Keep newlines as-is for Typst (they work natively in cell content)
            # Don't escape yet - we'll handle that after formatting for multi-line cells
            return str(x)

        # Element-wise conversion
        dfs = dfs.map(_prep_cell) if hasattr(dfs, "map") else dfs.applymap(_prep_cell)  # type: ignore[attr-defined]

        # Determine row groups (if MultiIndex on rows)
        row_levels = dfs.index.nlevels
        row_groups_present = row_levels > 1
        if row_groups_present:
            top_row_id = dfs.index.get_level_values(0).to_list()
            row_groups = list(dict.fromkeys(top_row_id))
            row_groups_len = [top_row_id.count(group) for group in row_groups]
            # Show only the inner row labels in the stub
            dfs.index = dfs.index.droplevel(0)

        stub_cols = dfs.index.nlevels
        data_cols = dfs.shape[1]
        total_cols = stub_cols + data_cols

        # Helper to collapse repeated labels into spans (same as TeX logic)
        def _make_spanner_row(level_labels):
            cells = []
            spans = []
            i = 0
            n = len(level_labels)
            while i < n:
                label = level_labels[i]
                span = 1
                j = i + 1
                while j < n and level_labels[j] == label:
                    span += 1
                    j += 1
                cells.append(str(label))
                spans.append(span)
                i = j
            return cells, spans

        # Build the table command with booktabs-style formatting
        lines = []
        lines.append("#table(")
        
        # Build column alignment: first column left-aligned, data columns per style (default center)
        data_align = s.get("data_align", "center")
        align_specs = ["left"] * stub_cols + [data_align] * data_cols
        align_str = ", ".join(align_specs)
        
        # Set column widths - first column auto-sizes to content, others use 1fr for equal distribution
        if s.get("first_col_width"):
            first_width = str(s["first_col_width"])
            width_specs = [first_width] + ["1fr"] * (total_cols - 1)
        else:
            # Default: first column auto-sizes to widest content, others share remaining space equally
            width_specs = ["auto"] + ["1fr"] * (total_cols - 1)
        
        width_str = ", ".join(width_specs)
        lines.append(f"  columns: ({width_str}), align: ({align_str}),")
        
        # Add column-gutter for spacing between column groups
        if isinstance(dfs.columns, pd.MultiIndex):
            # Create gutter pattern: 0.5em where second-level index changes, 0em elsewhere
            gutter_specs = ["0em"]  # After stub column
            # Check for changes in the second-level index (index level -2)
            if dfs.columns.nlevels >= 2:
                second_level_idx = dfs.columns.get_level_values(-2)
                for i in range(data_cols - 1):
                    # If the second-level value changes between columns i and i+1, add spacing
                    if second_level_idx[i] != second_level_idx[i + 1]:
                        gutter_specs.append("0.5em")
                    else:
                        gutter_specs.append("0em")
            else:
                # If no second level, no spacing
                gutter_specs.extend(["0em"] * (data_cols - 1))
            gutter_str = ", ".join(gutter_specs)
            lines.append(f"  column-gutter: ({gutter_str}),")
        
        lines.append(f"  stroke: none, row-gutter: 0.2em,")
        lines.append(f"  table.hline(stroke: 0.08em),")

        # Header rows with spanners (MultiIndex columns)
        if isinstance(dfs.columns, pd.MultiIndex):
            col_levels = dfs.columns.nlevels
            for lvl in range(col_levels - 1):
                labels_lvl = [dfs.columns[i][lvl] for i in range(len(dfs.columns))]
                row_cells, row_spans = _make_spanner_row(labels_lvl)

                row_parts: list[str] = ["[]"] * stub_cols
                for cell_label, span in zip(row_cells, row_spans, strict=False):
                    lbl = self._escape_typst(cell_label, escape_asterisks=False)
                    row_parts.append(f"[#table.cell(colspan: {span})[{lbl}]]")
                lines.append("  " + ", ".join(row_parts) + ",")

                # cmidrule-like partial lines under each spanner
                # Typst uses 1-based column indexing, and hlines span the full width of each spanner
                start_col = 1  # First column (including stub)
                for span in row_spans:
                    end_col = start_col + span
                    lines.append(f"  table.hline(stroke: 0.03em, start: {start_col}, end: {end_col}),")
                    start_col = end_col  # Next spanner starts where this one ends

            # Last level: the actual column names
            last_labels = [dfs.columns[i][-1] for i in range(len(dfs.columns))]
        else:
            last_labels = list(dfs.columns)

        # Bottom header row (always present)
        last_parts: list[str] = ["[]"] * stub_cols
        for label in last_labels:
            lbl = self._escape_typst(label, escape_asterisks=False)
            last_parts.append(f"[{lbl}]")
        lines.append("  " + ", ".join(last_parts) + ",")

        # Midrule after headers
        lines.append(f"  table.hline(stroke: 0.05em),")

        # Data rows (with optional row groups)
        def _row_line(cells_list: list[str]) -> str:
            return "  " + ", ".join(cells_list) + ","
        
        def _format_typst_cell(content: str) -> str:
            """Format cell content for Typst, handling multi-line content."""
            # ETable produces actual newline characters (\n as single char 0x0A)
            if "\n" in content:
                # Multi-line content: split by newline, escape each line, then format
                lines_content = content.split("\n")
                escaped_lines = [self._escape_typst(line.strip(), escape_asterisks=True) for line in lines_content if line.strip()]
                # Use Typst's backslash syntax for line breaks within a cell
                # Join lines with " \ " (backslash is Typst's line break operator)
                formatted_content = " \\ ".join(escaped_lines)
                return f"[{formatted_content}]"
            else:
                # Single-line content: escape and keep on one line
                escaped_content = self._escape_typst(content, escape_asterisks=True)
                return f"[{escaped_content}]"

        if row_groups_present:
            start = 0
            for gi, (gname, glen) in enumerate(zip(row_groups, row_groups_len, strict=False)):
                # Add top separator for row group (if not first group or if 't' in rgroup_sep)
                if gi > 0 and "t" in self.rgroup_sep:
                    lines.append(f"  table.hline(stroke: 0.03em),")
                
                if self.rgroup_display:
                    fmt = str(s.get("group_header_format", "%s"))
                    gtext = fmt % str(gname)
                    gtext = self._escape_typst(gtext, escape_asterisks=False)
                    # Use single cell spanning all columns to avoid line breaks in long headers
                    lines.append(f"  [#table.cell(colspan: {total_cols})[{gtext}]],")

                # Add bottom separator for row group (if 'b' in rgroup_sep)
                if "b" in self.rgroup_sep:
                    lines.append(f"  table.hline(stroke: 0.03em),")

                end = start + glen
                for ridx in range(start, end):
                    row_label = str(dfs.index[ridx])
                    row_parts = [_format_typst_cell(row_label)]
                    for j in range(data_cols):
                        cell_val = str(dfs.iloc[ridx, j])
                        row_parts.append(_format_typst_cell(cell_val))
                    lines.append(_row_line(row_parts))
                
                start = end
        else:
            for ridx in range(dfs.shape[0]):
                row_label = str(dfs.index[ridx])
                row_parts = [_format_typst_cell(row_label)]
                for j in range(data_cols):
                    cell_val = str(dfs.iloc[ridx, j])
                    row_parts.append(_format_typst_cell(cell_val))
                lines.append(_row_line(row_parts))

        # Add bottom rule, then optional notes row below the rule, then close table
        lines.append(f"  table.hline(stroke: 0.08em),")
        if self.notes is not None:
            notes_escaped = self._escape_typst(self.notes, escape_asterisks=True)
            lines.append(f"  [#table.cell(colspan: {total_cols})[#text(size: {notes_fontsize})[{notes_escaped}]]],")
        lines.append(")")

        typst_table = "\n".join(lines)

        # Add caption and/or reference label if present
        if self.caption is not None or self.tab_label is not None:
            # Build figure wrapper with caption and optional label for references
            caption_line = ""
            if self.caption is not None:
                # Escape special characters in caption (but preserve * for bold formatting)
                caption_escaped = self._escape_typst(self.caption, escape_asterisks=False)
                caption_line = f", caption: [{caption_escaped}]"
            
            label_line = ""
            if self.tab_label is not None:
                label_line = f"\n<{self.tab_label}>"
            
            # Wrap table in figure() context with caption and label
            # Syntax: #figure([content]{caption_line}){label_line}
            typst_table = f"#figure([\n{typst_table}\n]{caption_line}){label_line}"
        
        # Apply symbol translation
        typst_table = self._translate_symbols(typst_table, "typst")

        return typst_table

    def _output_gt(
        self,
        gt_style: dict[str, object] | None = None,
        **kwargs,
    ):
        # Make a copy of the DataFrame to avoid modifying the original
        dfs = self.df.copy()
        # Resolve GT style (per-call -> class default)
        s = dict(self.DEFAULT_GT_STYLE)
        if gt_style:
            s.update(gt_style)

        # In all cells replace line breaks with <br>
        dfs = dfs.replace(r"\n", "<br>", regex=True)

        # GT does not support MultiIndex columns, so we need to flatten the columns
        if isinstance(dfs.columns, pd.MultiIndex):
            # Store labels of the last level of the column index (to use as column names)
            col_names = dfs.columns.get_level_values(-1)
            nlevels = dfs.columns.nlevels

            # Assign column numbers to the lowest index level
            col_numbers = list(map(str, range(len(dfs.columns))))
            # Save the whole column index in order to generate table spanner labels later
            dfcols = dfs.columns.to_list()
            # Flatten the column index just numbering the columns
            dfs.columns = pd.Index(col_numbers)
            # Store the mapping of column numbers to column names
            col_dict = dict(zip(col_numbers, col_names, strict=False))
            # Modify the last elements in each tuple in dfcols
            dfcols = [(t[:-1] + (col_numbers[i],)) for i, t in enumerate(dfcols)]
        else:
            nlevels = 1

        # store row index and then reset to have the index as columns to be displayed in the table
        rowindex = dfs.index

        # Handle potential name conflicts when resetting index
        # Find safe names for index columns that don't conflict with existing columns
        if isinstance(rowindex, pd.MultiIndex):
            index_names = []
            for i, name in enumerate(rowindex.names):
                if name is None or name in dfs.columns:
                    # Use a safe default name
                    safe_name = f"__index_level_{i}__"
                    while safe_name in dfs.columns:
                        safe_name = f"__index_level_{i}_{hash(safe_name) % 1000}__"
                    index_names.append(safe_name)
                else:
                    index_names.append(name)

            # Temporarily rename index levels to avoid conflicts
            dfs.index.names = index_names
        else:
            if rowindex.name is None or rowindex.name in dfs.columns:
                # Use a safe default name
                safe_name = "__index__"
                while safe_name in dfs.columns:
                    safe_name = f"__index_{hash(safe_name) % 1000}__"
                dfs.index.name = safe_name

        dfs.reset_index(inplace=True)

        # Specify the rowname_col and groupname_col
        if isinstance(rowindex, pd.MultiIndex):
            rowname_col = dfs.columns[1]
            groupname_col = dfs.columns[0]
        else:
            rowname_col = dfs.columns[0]
            groupname_col = None

        # Generate the table with GT
        gt = GT(dfs, auto_align=False)

        # When caption is provided, add it to the table
        if self.caption is not None:
            gt = gt.tab_header(title=self.caption).tab_options(
                table_border_top_style="hidden"
            )

        # Add column spanners based on multiindex
        if nlevels > 1:
            for i in range(nlevels - 1):
                col_spanners = {}
                # Iterate over columns and group them by the labels in the respective level
                for c in dfcols:
                    key = c[i]
                    if key not in col_spanners:
                        col_spanners[key] = []
                    col_spanners[key].append(c[-1])
                for label, columns in col_spanners.items():
                    if label != "":
                        gt = gt.tab_spanner(
                            label=label, columns=columns, level=nlevels - 1 - i
                        )
                # Restore column names
                gt = gt.cols_label(**col_dict)

        # Customize the table layout using GT style defaults
        gt = gt.tab_source_note(self.notes).tab_stub(
            rowname_col=rowname_col, groupname_col=groupname_col
        )

        # Handle table_font_size_all - applies to all font sizes except source_notes_font_size
        # Individual font sizes override the all setting if explicitly provided
        if s.get("table_font_size_all"):
            font_size = s["table_font_size_all"]
            font_params = {
                "table_font_size": font_size,
                "column_labels_font_size": font_size,
                "row_group_font_size": font_size,
                "stub_font_size": font_size,
                "heading_title_font_size": font_size,
            }
            for param, default_value in font_params.items():
                if (
                    param not in s
                ):  # Only set if not already explicitly provided by user
                    s[param] = default_value

        # Prepare style options, excluding special parameters handled separately
        special_params = {
            "align",
            "table_width",
            "first_col_width",
            "table_font_size_all",
        }
        tab_options_dict = {k: v for k, v in s.items() if k not in special_params}

        # Add required defaults
        tab_options_dict.update(
            {
                "table_border_bottom_style": "hidden",
                "stub_border_style": "hidden",
            }
        )

        # Handle table width setting
        if s.get("table_width"):
            tab_options_dict["table_width"] = str(s["table_width"])

        # Apply all styling options at once
        gt = gt.tab_options(**tab_options_dict).cols_align(
            align=s.get("align", "center")
        )

        # Handle first column width setting
        if s.get("first_col_width"):
            # Apply width to the first column (stub column)
            gt = gt.cols_width(cases={rowname_col: str(s["first_col_width"])})

        # Customize row group display
        if "t" not in self.rgroup_sep:
            gt = gt.tab_options(row_group_border_top_style="none")
        if "b" not in self.rgroup_sep:
            gt = gt.tab_options(row_group_border_bottom_style="none")
        if not self.rgroup_display:
            gt = gt.tab_options(row_group_font_size="0px", row_group_padding="0px")

        # Apply symbol translation to the final HTML output
        html_output = gt.as_raw_html()
        translated_html = self._translate_symbols(html_output, "html")

        # Create a new GT object with the translated HTML
        # Since GT doesn't have a direct way to modify HTML, we'll monkey-patch it
        gt._repr_html_ = lambda: translated_html
        gt.as_raw_html = lambda: translated_html

        return gt

    def __repr__(self):
        """
        Return a representation of the table.

        In notebook environments, this will automatically display the table
        with dual output format (HTML in notebooks, LaTeX in Quarto) without
        requiring an explicit call to make().

        Returns
        -------
        str
            An empty string
        """

        # For dual output, we need to handle parameters differently
        # Create dual output object for notebook/Quarto compatibility
        class DualOutput:
            """Display different outputs in notebook vs Quarto rendering."""

            def __init__(self, notebook_html, quarto_latex):
                self.notebook_html = notebook_html
                self.quarto_latex = quarto_latex

            def _repr_mimebundle_(self, include=None, exclude=None):
                return {
                    "text/html": self.notebook_html,
                    "text/latex": self.quarto_latex,
                }

        # Generate both HTML and LaTeX outputs with their specific styles
        gt_style = self._display_styles.get("gt_style", {})
        tex_style = self._display_styles.get("tex_style", {})

        html_output = self._output_gt(gt_style=gt_style).as_raw_html()
        tex_output = self._output_tex(tex_style=tex_style)

        # Add CSS to remove zebra striping if desired
        html_output = (
            """
        <style>
        table tr:nth-child(even) {
            background-color: transparent !important;
        }
        </style>
        """
            + html_output
        )
        # Create and display the dual output object
        from IPython.display import display

        dual_output = DualOutput(html_output, tex_output)
        display(dual_output)
        return ""

    def __call__(self, type=None, **kwargs):
        """
        Make this object callable, equivalent to calling make().

        Parameters
        ----------
        type : str, optional
            The output type to create. If None, displays dual output.
        **kwargs : dict
            Additional parameters to pass to make().

        Returns
        -------
        output : object
            The output object returned by make().
        """
        return self.make(type=type, **kwargs)

